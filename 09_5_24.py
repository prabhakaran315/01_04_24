#--------------------Imported libraries--------------------#
from tkinter import *
from tkinter import ttk
from tkinter import messagebox
import math

import numpy as np
from PIL import ImageTk, Image
from PySimpleGUI import MsgBox
from docx import Document
from docx.shared import RGBColor
from docx2pdf import convert
import os
from tkinter import filedialog
from docx.shared import Pt
from docx.shared import Cm, Inches
from docxcompose.composer import Composer
from docx.enum.section import WD_ORIENT
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import time
#from docx.oxml.xmlchemy import OxmlElement
from docx.oxml.parser import OxmlElement
from docx.oxml.shared import qn
from tkinter import font
import contextlib
import sys
import subprocess
import math
from cryptography.hazmat.backends import default_backend
from cryptography.hazmat.primitives import hashes
from cryptography.hazmat.primitives.kdf.pbkdf2 import PBKDF2HMAC
from cryptography.fernet import Fernet
import requests
from datetime import datetime, timedelta
import mysql.connector
import uuid
#------------------------------------#

#------------------Libraries ended---------------------#
# while compling to excel uncomment this :
'''if sys.platform.startswith('win'):
    # Redirect stdout and stderr to null device
    # to avoid console window popup
    startupinfo = subprocess.STARTUPINFO()
    startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW
    null_device = open(os.devnull, 'w')
    sys.stdout = null_device
    sys.stderr = null_device


@contextlib.contextmanager
def stdout_redirector(stream):
    old_stdout = sys.stdout
    sys.stdout = stream
    try:
        yield
    finally:
        sys.stdout = old_stdout'''

#-----------------Creating windows and assign empty values for variables-------------#

root = Tk()
root.title("Astra Sizing Tool")
root.geometry("730x600")
# root.iconbitmap("astra.ico")
# root.geometry("1500x650")
root.configure(bg="white")
# root.wm_iconbitmap("")

icon = Image.open('inphase.ico')
photo = ImageTk.PhotoImage(icon)
root.wm_iconphoto(False, photo) #This line is used to change the icon
root.iconbitmap('inphase.ico')

root.resizable(False, False)
harmonicsentries = []
valueerrorflag = 0
fault = 0
page1_frequency = 0
number_of_entries = 0
adminflag = 0
active_user = 0
selected_directory = ''
selected_file_name = ''
export_dir_name = ''
export_file_name = ''
modified_flag = 0
save_flag = 0
export_progress_flag = 0
import_progress_flag = 0
# Page - 2 global variable
p2_harmonicsentries = []
p2_valueerrorflag = 0
p2_fault = 0
p2_number_of_entries = 0
page2_frequency = 0
# Page - 3 global variable
p3_harmonicsentries = []
p3_valueerrorflag = 0
p3_fault = 0
p3_number_of_entries = 0
page3_frequency = 0
page1_comment_box_message = ""
page2_comment_box_message = ""
page3_comment_box_message = ""

# Enable modified flag
def enable_modified_flag(*args):
    # page -1 enable modfied tags
    global modified_flag
    modified_flag = 1
    if (astranotebook.index(astranotebook.select()) == 0):
        astranotebook.tab(0, text="Astra 3P,3W*")
        modified_indication.config(text="*")
        # Clearing Astra Rating
        AHFsizeentry.config(borderwidth=2, state="normal")
        AHFsizeentry.delete(0, "end")
        AHFsizeentry.insert(0, str(empty_entry.get(1.0, "end-1c")))
        AHFsizeentry.config(borderwidth=2, state="disable")
        # Clearing Astra Rating (ambient)
        AHFsize1entry.config(borderwidth=2, state="normal")
        AHFsize1entry.delete(0, "end")
        AHFsize1entry.insert(0, str(empty_entry.get(1.0, "end-1c")))
        AHFsize1entry.config(borderwidth=2, state="disable")
        location_field.config(text="location :" + str(selected_directory))
        p2_location_field.config(text="location :" + str(selected_directory))
        p3_location_field.config(text="location :" + str(selected_directory))
        p4_location_field.config(text="location :" + str(selected_directory))
        print(modified_flag)
    if (astranotebook.index(astranotebook.select()) == 1):
        astranotebook.tab(1, text="Astra 3P,4W*")
        p2_modified_indication.config(text="*")
        # Clearing Astra Rating
        p2_AHFsizeentry.config(borderwidth=2, state="normal")
        p2_AHFsizeentry.delete(0, "end")
        p2_AHFsizeentry.insert(0, str(empty_entry.get(1.0, "end-1c")))
        p2_AHFsizeentry.config(borderwidth=2, state="disable")
        # Clearing Astra Rating (ambient)
        p2_AHFsize1entry.config(borderwidth=2, state="normal")
        p2_AHFsize1entry.delete(0, "end")
        p2_AHFsize1entry.insert(0, str(empty_entry.get(1.0, "end-1c")))
        p2_AHFsize1entry.config(borderwidth=2, state="disable")
        location_field.config(text="location :" + str(selected_directory))
        p2_location_field.config(text="location :" + str(selected_directory))
        p3_location_field.config(text="location :" + str(selected_directory))
        p4_location_field.config(text="location :" + str(selected_directory))
        print(modified_flag)
    if (astranotebook.index(astranotebook.select()) == 2):
        astranotebook.tab(2, text="Astra 3P, N*")
        p3_modified_indication.config(text="*")
        # Clearing Astra Rating
        p3_AHFsizeentry.config(borderwidth=2, state="normal")
        p3_AHFsizeentry.delete(0, "end")
        p3_AHFsizeentry.insert(0, str(empty_entry.get(1.0, "end-1c")))
        p3_AHFsizeentry.config(borderwidth=2, state="disable")
        # Clearing Astra Rating (ambient)
        p3_AHFsize1entry.config(borderwidth=2, state="normal")
        p3_AHFsize1entry.delete(0, "end")
        p3_AHFsize1entry.insert(0, str(empty_entry.get(1.0, "end-1c")))
        p3_AHFsize1entry.config(borderwidth=2, state="disable")
        location_field.config(text="location :" + str(selected_directory))
        p2_location_field.config(text="location :" + str(selected_directory))
        p3_location_field.config(text="location :" + str(selected_directory))
        p4_location_field.config(text="location :" + str(selected_directory))
        print(modified_flag)
    #----------- 4th notebook of True Power Factor ---------------#
    if (astranotebook.index(astranotebook.select()) == 3):
        astranotebook.tab(3, text="True Power Factor*")
        head_notebook.config(text="True Power Factor Performance Calculator *")
        location_field.config(text="location :" + str(selected_directory))
        p2_location_field.config(text="location :" + str(selected_directory))
        p3_location_field.config(text="location :" + str(selected_directory))
        p4_location_field.config(text="location :" + str(selected_directory))
        print('modified flag : ', modified_flag)



# disable modified flag
def disable_modified_flag(*args):
    # page -1 disable modfied tags
    global modified_flag
    modified_flag = 0

    if (astranotebook.index(astranotebook.select()) == 0):
        astranotebook.tab(0, text="Astra 3P,3W")
        modified_indication.config(text="")
        location_field.config(text="location :" + str(selected_directory))
        p2_location_field.config(text="location :" + str(selected_directory))
        p3_location_field.config(text="location :" + str(selected_directory))
        p4_location_field.config(text="location :" + str(selected_directory))
        print(modified_flag)
    if (astranotebook.index(astranotebook.select()) == 1):
        astranotebook.tab(1, text="Astra 3P,4W")
        p2_modified_indication.config(text="")
        location_field.config(text="location :" + str(selected_directory))
        p2_location_field.config(text="location :" + str(selected_directory))
        p3_location_field.config(text="location :" + str(selected_directory))
        p4_location_field.config(text="location :" + str(selected_directory))
        print(modified_flag)
    if (astranotebook.index(astranotebook.select()) == 2):
        astranotebook.tab(2, text="Astra 3P, N")
        p3_modified_indication.config(text="")
        location_field.config(text="location :" + str(selected_directory))
        p2_location_field.config(text="location :" + str(selected_directory))
        p3_location_field.config(text="location :" + str(selected_directory))
        p4_location_field.config(text="location :" + str(selected_directory))
        print(modified_flag)

    if (astranotebook.index(astranotebook.select()) == 3):
        astranotebook.tab(3, text="True Power Factor")
        head_notebook.config(text="True Power Factor Performance Calculator  ")
        location_field.config(text="location :" + str(selected_directory))
        p2_location_field.config(text="location :" + str(selected_directory))
        p3_location_field.config(text="location :" + str(selected_directory))
        p4_location_field.config(text="location :" + str(selected_directory))
        print('modified flag : ', modified_flag)

#------------make rows bold--------------#
def make_rows_bold(*rows):
    # make the rows bold
    for row in rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.name = 'Calibri'
                    run.font.size = Pt(14)  # 14-14

#----------------------Export to PDF for Astra--------------------#
def export_to_pdf():
    # page -1 export to pdf function
    global export_progress_flag
    global table, export_ini
    fetch_comment_box_message()
    # print(astranotebook.index(astranotebook.select()))
    global number_of_entries, export_dir_name, export_file_name, selected_directory, selected_file_name
    global p2_number_of_entries, p3_number_of_entries, page1_comment_box_message, page2_comment_box_message, page3_comment_box_message
    if (astranotebook.index(astranotebook.select()) == 0):
        process()
        try:
            if (number_of_entries > 0):
                # print("number_of_entries", number_of_entries)
                initial_table_rows = number_of_entries + 1
                # dir_name = filedialog.askdirectory()  # asks user to choose a directory
                status_p1_entry.config(borderwidth=2, state="normal")
                status_p1_entry.delete(0, "end")
                status_p1_entry.insert(0, "Export initiated...")
                status_p1_entry.config(borderwidth=2, state="disable")
                if (export_dir_name == ''):
                    export_dir_name = filedialog.asksaveasfilename(title="Export As",
                                                                   filetypes=(("PDF", "*.pdf"), ("All Files", "*.*")),
                                                                   initialfile="Design Document")  # asks user to choose a directory
                    try:
                        os.chdir(os.path.dirname(export_dir_name))
                        # file_name = os.path.basename(dir_name).split('.', 1)[0]
                        export_file_name = os.path.basename(export_dir_name).split('.', 1)[0]
                        # copying the user file and directory to the save location
                        selected_directory = export_dir_name
                        selected_file_name = export_file_name
                    except:
                        pass
                        print("Path Nor defined")
                else:
                    export_dir_name = selected_directory
                    export_file_name = selected_file_name

                doc = Document()
                ######################################################
                # Adding A Header to the Document

                # Choosing the top most section of the page
                section = doc.sections[0]

                # Calling the header
                header = section.header

                # Calling the paragraph already present in
                # the header section
                header_para = header.paragraphs[0]

                # Adding text in the header
                header_para.text = "InPhase Power Technologies - ASTRA Sizing Document"
                header_para.alignment = 1
                ######################################################
                # Add a Title to the document
                doc.add_heading('Astra Rating', 1)

                # Creating a table object
                table = doc.add_table(rows=initial_table_rows, cols=2)
                # aliging table to center
                table.alignment = WD_TABLE_ALIGNMENT.CENTER

                # Creating table Heading
                table.cell(0, 0).text = 'Harmonics (n)'
                table.cell(0, 1).text = 'Current (I)'

                # vetrical centering

                table.cell(0, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.cell(0, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

                for x in range(0, 2): make_rows_bold(table.columns[x])  # making bold and font size change

                # Align table column in table
                for cell in table.columns[0].cells: cell.width = Inches(2)
                for cell in table.columns[1].cells: cell.width = Inches(2)

                for row in table.rows:
                    row.height = Cm(.75)

                table.style = 'Light Grid Accent 5'

                for x in range(0, number_of_entries):
                    table.cell(x + 1, 0).text = str(harmonicsentries[x].get())
                    table.cell(x + 1, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                for x in range(0, number_of_entries):
                    table.cell(x + 1, 1).text = str(harmonicsentries[x + 10].get()) + ' A'
                    table.cell(x + 1, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

                # making Bold of the both the column
                for x in range(0, 2): make_rows_bold(table.columns[x])
                # doc.save('rating.docx')
                doc.add_heading("SETTINGS", 1)

                settings_table = doc.add_table(rows=5, cols=2)
                # aliging table to center
                settings_table.alignment = WD_TABLE_ALIGNMENT.CENTER

                # Creating table Heading
                settings_table.cell(0, 0).text = 'Settings'
                settings_table.cell(1, 0).text = 'Frequency'
                if (active_user == 0):
                    settings_table.cell(2, 0).text = 'Notch Profile'
                elif (active_user == 1):
                    settings_table.cell(2, 0).text = 'Base (n)'
                settings_table.cell(3, 0).text = 'Ambient Temperature (°C)'
                settings_table.cell(4, 0).text = 'Amplification Factor '
                settings_table.cell(0, 1).text = 'Value'
                if (page1_frequency == 50):
                    settings_table.cell(1, 1).text = "50 Hz"
                elif (page1_frequency == 60):
                    settings_table.cell(1, 1).text = "60 Hz"
                if (active_user == 0):
                    if (str(basenentry.get()) == "10"): settings_table.cell(2, 1).text = "LOW"
                    if (str(basenentry.get()) == "5"): settings_table.cell(2, 1).text = "MEDIUM"
                    if (str(basenentry.get()) == "3"): settings_table.cell(2, 1).text = "HIGH"
                    settings_table.cell(3, 1).text = str(ambtempentry.get()) + ' °C'
                    settings_table.cell(4, 1).text = str(ambfactorentry.get())
                elif (active_user == 1):
                    settings_table.cell(2, 1).text = str(basenentry.get())
                    settings_table.cell(3, 1).text = str(ambtempentry.get()) + ' °C'
                    settings_table.cell(4, 1).text = str(ambfactorentry.get())

                # vetrical centering

                settings_table.cell(0, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                settings_table.cell(0, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

                for x in range(0, 2): make_rows_bold(settings_table.columns[x])  # making bold and font size change

                # Align table column in table
                for cell in settings_table.columns[0].cells: cell.width = Inches(3)
                for cell in settings_table.columns[1].cells: cell.width = Inches(3)

                for row in settings_table.rows:
                    row.height = Cm(.75)

                settings_table.style = 'Light Grid Accent 5'

                #####################################################
                # Creating Result Table
                doc.add_heading("ASTRA RATING", 1)

                rating_table = doc.add_table(rows=3, cols=2)
                # aliging table to center
                rating_table.alignment = WD_TABLE_ALIGNMENT.CENTER

                # Creating table Heading
                rating_table.cell(0, 0).text = 'Description'
                rating_table.cell(1, 0).text = 'ASTRA RATING'
                rating_table.cell(2, 0).text = 'ASTRA RATING @' + str(ambtempentry.get()) + '(°C)'
                rating_table.cell(0, 1).text = 'Current Rating'
                rating_table.cell(1, 1).text = str(astrarating) + ' A'
                rating_table.cell(2, 1).text = str(ambastrarating).lstrip('0') + ' A'

                # for x in range(1,4):
                rating_table.cell(0, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                rating_table.cell(0, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

                # vetrical centering

                rating_table.cell(0, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                rating_table.cell(0, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

                for x in range(0, 2): make_rows_bold(rating_table.columns[x])  # making bold and font size change

                # Align table column in table
                for cell in rating_table.columns[0].cells: cell.width = Inches(3)
                for cell in rating_table.columns[1].cells: cell.width = Inches(3)

                for row in rating_table.rows:
                    row.height = Cm(.75)

                rating_table.style = 'Light Grid Accent 5'

                #####################################################
                # adding a page break
                doc.add_page_break()
                #####################################################
                ###########################################
                # Reactive power Calaculation
                if formula_combo.get() == "V, I, IPF, TPF":
                    doc.add_heading("Reactive Current details", 1)

                    reactive_table = doc.add_table(rows=9, cols=2)
                    # aliging table to center
                    reactive_table.alignment = WD_TABLE_ALIGNMENT.CENTER

                    # Creating table Heading
                    reactive_table.cell(0, 0).text = 'Description'
                    reactive_table.cell(1, 0).text = 'Voltage*'
                    reactive_table.cell(2, 0).text = 'Current*'
                    reactive_table.cell(3, 0).text = 'Active Power'
                    reactive_table.cell(4, 0).text = 'Apparent Power'
                    reactive_table.cell(5, 0).text = 'Reactive Power'
                    reactive_table.cell(6, 0).text = 'IPF*'
                    reactive_table.cell(7, 0).text = 'TPF*'
                    reactive_table.cell(8, 0).text = 'Reactive Current'
                    reactive_table.cell(0, 1).text = 'Value'
                    reactive_table.cell(1, 1).text = str(volentry.get()) + ' V'
                    reactive_table.cell(2, 1).text = str(curentry.get()) + ' A'
                    reactive_table.cell(3, 1).text = str(kwentry.get()) + ' kW'
                    reactive_table.cell(4, 1).text = str(kvaentry.get()) + ' kVA'
                    reactive_table.cell(5, 1).text = str(kvarentry.get()) + ' kVAr'
                    reactive_table.cell(6, 1).text = str(format(float(IPFentry.get()), ".2f"))
                    reactive_table.cell(7, 1).text = str(format(float(TPFentry.get()), ".2f"))
                    reactive_table.cell(8, 1).text = str(IQentry.get()) + ' A'

                    reactive_table.cell(0, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                    reactive_table.cell(0, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

                    # vetrical centering

                    reactive_table.cell(0, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                    reactive_table.cell(0, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

                    for x in range(0, 2): make_rows_bold(reactive_table.columns[x])  # making bold and font size change

                    # Align table column in table
                    for cell in reactive_table.columns[0].cells: cell.width = Inches(3)
                    for cell in reactive_table.columns[1].cells: cell.width = Inches(3)

                    for row in reactive_table.rows:
                        row.height = Cm(.75)

                    reactive_table.style = 'Light Grid Accent 5'
                    doc.add_paragraph(''' ''')
                    word = "* - Represents User Inputs"
                    p = doc.add_paragraph()
                    runner = p.add_run(word)
                    runner.bold = True
                    runner.italic = True

                elif formula_combo.get() == "V, kW, IPF, TPF":
                    doc.add_heading("Reactive Current details", 1)

                    reactive_table = doc.add_table(rows=9, cols=2)
                    # aliging table to center
                    reactive_table.alignment = WD_TABLE_ALIGNMENT.CENTER

                    # Creating table Heading
                    reactive_table.cell(0, 0).text = 'Description'
                    reactive_table.cell(1, 0).text = 'Voltage*'
                    reactive_table.cell(2, 0).text = 'Current'
                    reactive_table.cell(3, 0).text = 'Active Power*'
                    reactive_table.cell(4, 0).text = 'Apparent Power'
                    reactive_table.cell(5, 0).text = 'Reactive Power'
                    reactive_table.cell(6, 0).text = 'IPF*'
                    reactive_table.cell(7, 0).text = 'TPF*'
                    reactive_table.cell(8, 0).text = 'Reactive Current'
                    reactive_table.cell(0, 1).text = 'Value'
                    reactive_table.cell(1, 1).text = str(volentry.get()) + ' V'
                    reactive_table.cell(2, 1).text = str(curentry.get()) + ' A'
                    reactive_table.cell(3, 1).text = str(kwentry.get()) + ' kW'
                    reactive_table.cell(4, 1).text = str(kvaentry.get()) + ' kVA'
                    reactive_table.cell(5, 1).text = str(kvarentry.get()) + ' kVAr'
                    reactive_table.cell(6, 1).text = str(format(float(IPFentry.get()), ".2f"))
                    reactive_table.cell(7, 1).text = str(format(float(TPFentry.get()), ".2f"))
                    reactive_table.cell(8, 1).text = str(IQentry.get()) + ' A'

                    reactive_table.cell(0, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                    reactive_table.cell(0, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

                    # vetrical centering

                    reactive_table.cell(0, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                    reactive_table.cell(0, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

                    for x in range(0, 2): make_rows_bold(reactive_table.columns[x])  # making bold and font size change

                    # Align table column in table
                    for cell in reactive_table.columns[0].cells: cell.width = Inches(3)
                    for cell in reactive_table.columns[1].cells: cell.width = Inches(3)

                    for row in reactive_table.rows:
                        row.height = Cm(.75)

                    reactive_table.style = 'Light Grid Accent 5'
                    doc.add_paragraph(''' ''')
                    word = "* - Represents User Inputs"
                    p = doc.add_paragraph()
                    runner = p.add_run(word)
                    runner.bold = True
                    runner.italic = True

                elif formula_combo.get() == "V, kW, kVA, TPF":
                    doc.add_heading("Reactive Current details", 1)

                    reactive_table = doc.add_table(rows=9, cols=2)
                    # aliging table to center
                    reactive_table.alignment = WD_TABLE_ALIGNMENT.CENTER

                    # Creating table Heading
                    reactive_table.cell(0, 0).text = 'Description'
                    reactive_table.cell(1, 0).text = 'Voltage*'
                    reactive_table.cell(2, 0).text = 'Current'
                    reactive_table.cell(3, 0).text = 'Active Power*'
                    reactive_table.cell(4, 0).text = 'Apparent Power*'
                    reactive_table.cell(5, 0).text = 'Reactive Power'
                    reactive_table.cell(6, 0).text = 'IPF'
                    reactive_table.cell(7, 0).text = 'TPF*'
                    reactive_table.cell(8, 0).text = 'Reactive Current'
                    reactive_table.cell(0, 1).text = 'Value'
                    reactive_table.cell(1, 1).text = str(volentry.get()) + ' V'
                    reactive_table.cell(2, 1).text = str(curentry.get()) + ' A'
                    reactive_table.cell(3, 1).text = str(kwentry.get()) + ' kW'
                    reactive_table.cell(4, 1).text = str(kvaentry.get()) + ' kVA'
                    reactive_table.cell(5, 1).text = str(kvarentry.get()) + ' kVAr'
                    reactive_table.cell(6, 1).text = str(format(float(IPFentry.get()), ".2f"))
                    reactive_table.cell(7, 1).text = str(format(float(TPFentry.get()), ".2f"))
                    reactive_table.cell(8, 1).text = str(IQentry.get()) + ' A'

                    reactive_table.cell(0, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                    reactive_table.cell(0, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

                    # vetrical centering

                    reactive_table.cell(0, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                    reactive_table.cell(0, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

                    for x in range(0, 2): make_rows_bold(reactive_table.columns[x])  # making bold and font size change

                    # Align table column in table
                    for cell in reactive_table.columns[0].cells: cell.width = Inches(3)
                    for cell in reactive_table.columns[1].cells: cell.width = Inches(3)

                    for row in reactive_table.rows:
                        row.height = Cm(.75)

                    reactive_table.style = 'Light Grid Accent 5'
                    doc.add_paragraph(''' ''')
                    word = "* - Represents User Inputs"
                    p = doc.add_paragraph()
                    runner = p.add_run(word)
                    runner.bold = True
                    runner.italic = True

                elif formula_combo.get() == "V, kVA, IPF, TPF":
                    doc.add_heading("Reactive Current details", 1)

                    reactive_table = doc.add_table(rows=9, cols=2)
                    # aliging table to center
                    reactive_table.alignment = WD_TABLE_ALIGNMENT.CENTER

                    # Creating table Heading
                    reactive_table.cell(0, 0).text = 'Description'
                    reactive_table.cell(1, 0).text = 'Voltage*'
                    reactive_table.cell(2, 0).text = 'Current'
                    reactive_table.cell(3, 0).text = 'Active Power'
                    reactive_table.cell(4, 0).text = 'Apparent Power*'
                    reactive_table.cell(5, 0).text = 'Reactive Power'
                    reactive_table.cell(6, 0).text = 'IPF*'
                    reactive_table.cell(7, 0).text = 'TPF*'
                    reactive_table.cell(8, 0).text = 'Reactive Current'
                    reactive_table.cell(0, 1).text = 'Value'
                    reactive_table.cell(1, 1).text = str(volentry.get()) + ' V'
                    reactive_table.cell(2, 1).text = str(curentry.get()) + ' A'
                    reactive_table.cell(3, 1).text = str(kwentry.get()) + ' kW'
                    reactive_table.cell(4, 1).text = str(kvaentry.get()) + ' kVA'
                    reactive_table.cell(5, 1).text = str(kvarentry.get()) + ' kVAr'
                    reactive_table.cell(6, 1).text = str(format(float(IPFentry.get()), ".2f"))
                    reactive_table.cell(7, 1).text = str(format(float(TPFentry.get()), ".2f"))
                    reactive_table.cell(8, 1).text = str(IQentry.get()) + ' A'

                    reactive_table.cell(0, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                    reactive_table.cell(0, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

                    # vetrical centering

                    reactive_table.cell(0, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                    reactive_table.cell(0, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

                    for x in range(0, 2): make_rows_bold(reactive_table.columns[x])  # making bold and font size change

                    # Align table column in table
                    for cell in reactive_table.columns[0].cells: cell.width = Inches(3)
                    for cell in reactive_table.columns[1].cells: cell.width = Inches(3)

                    for row in reactive_table.rows:
                        row.height = Cm(.75)

                    reactive_table.style = 'Light Grid Accent 5'
                    doc.add_paragraph(''' ''')
                    word = "* - Represents User Inputs"
                    p = doc.add_paragraph()
                    runner = p.add_run(word)
                    runner.bold = True
                    runner.italic = True

                elif formula_combo.get() == "V, kVAR":
                    doc.add_heading("Reactive Current details", 1)

                    reactive_table = doc.add_table(rows=4, cols=2)
                    # aliging table to center
                    reactive_table.alignment = WD_TABLE_ALIGNMENT.CENTER

                    # Creating table Heading
                    reactive_table.cell(0, 0).text = 'Description'
                    reactive_table.cell(1, 0).text = 'Voltage*'
                    reactive_table.cell(2, 0).text = 'Reactive Power*'
                    reactive_table.cell(3, 0).text = 'Reactive Current'
                    reactive_table.cell(0, 1).text = 'Value'
                    reactive_table.cell(1, 1).text = str(volentry.get()) + ' V'
                    reactive_table.cell(2, 1).text = str(kvarentry.get()) + ' kVAr'
                    reactive_table.cell(3, 1).text = str(IQentry.get()) + ' A'

                    reactive_table.cell(0, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                    reactive_table.cell(0, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

                    # vetrical centering

                    reactive_table.cell(0, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                    reactive_table.cell(0, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

                    for x in range(0, 2): make_rows_bold(reactive_table.columns[x])  # making bold and font size change

                    # Align table column in table
                    for cell in reactive_table.columns[0].cells: cell.width = Inches(3)
                    for cell in reactive_table.columns[1].cells: cell.width = Inches(3)

                    for row in reactive_table.rows:
                        row.height = Cm(.75)

                    reactive_table.style = 'Light Grid Accent 5'
                    doc.add_paragraph(''' ''')
                    word = "* - Represents User Inputs"
                    p = doc.add_paragraph()
                    runner = p.add_run(word)
                    runner.bold = True
                    runner.italic = True

                elif formula_combo.get() == "IQ":
                    doc.add_heading("Reactive Current details", 1)

                    reactive_table = doc.add_table(rows=2, cols=2)
                    # aliging table to center
                    reactive_table.alignment = WD_TABLE_ALIGNMENT.CENTER

                    # Creating table Heading
                    reactive_table.cell(0, 0).text = 'Description'
                    reactive_table.cell(1, 0).text = 'Reactive Current*'
                    reactive_table.cell(0, 1).text = 'Value'
                    reactive_table.cell(1, 1).text = str(IQentry.get()) + ' A'

                    reactive_table.cell(0, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                    reactive_table.cell(0, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

                    # vetrical centering

                    reactive_table.cell(0, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                    reactive_table.cell(0, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

                    for x in range(0, 2): make_rows_bold(reactive_table.columns[x])  # making bold and font size change

                    # Align table column in table
                    for cell in reactive_table.columns[0].cells: cell.width = Inches(3)
                    for cell in reactive_table.columns[1].cells: cell.width = Inches(3)

                    for row in reactive_table.rows:
                        row.height = Cm(.75)

                    reactive_table.style = 'Light Grid Accent 5'
                    doc.add_paragraph(''' ''')
                    word = "* - Represents User Inputs"
                    p = doc.add_paragraph()
                    runner = p.add_run(word)
                    runner.bold = True
                    runner.italic = True
                ###########################################
                #####################################################
                # Deleting the place holder
                if comment_box_message.get('1.0', 'end-1c') == placeholder_text:
                    pass
                else:
                    doc.add_heading("Comments", 1)
                    para = doc.add_paragraph(str(page1_comment_box_message))
                # bold_para = para.add_run('''\n '''+str(page1_comment_box_message))
                ############################################

                #####################################################
                doc.add_heading("Prepared By", 1)
                para = doc.add_paragraph()
                bold_para = para.add_run(''' \n Advanced Power Quality''')
                ############################################

                # Adding a border to a Page
                sec_pr = doc.sections[0]._sectPr  # get the section properties el
                # create new borders el
                pg_borders = OxmlElement('w:pgBorders')
                # specifies how the relative positioning of the borders should be calculated
                pg_borders.set(qn('w:offsetFrom'), 'page')
                for border_name in ('top', 'left', 'bottom', 'right',):  # set all borders
                    border_el = OxmlElement(f'w:{border_name}')
                    border_el.set(qn('w:val'), 'single')  # a single line
                    border_el.set(qn('w:sz'), '4')  # for meaning of  remaining attrs please look docs
                    border_el.set(qn('w:space'), '24')
                    border_el.set(qn('w:color'), 'auto')
                    pg_borders.append(border_el)  # register single border to border el
                sec_pr.append(pg_borders)  # apply border changes to section
                ############################################
                doc.save(str(export_file_name) + ".docx")
                convert(str(export_file_name) + ".docx")
                os.remove(str(export_file_name) + ".docx")
                export_progress_flag = 1
                save_nfo()
                export_progress_flag = 0
                status_p1_entry.config(borderwidth=2, state="normal")
                status_p1_entry.delete(0, "end")
                status_p1_entry.config(borderwidth=2, state="disable")
                disable_modified_flag()
                messagebox.showinfo("Export Information", export_file_name + ".pdf" + " Exported")

            else:
                mandatory_entries_check()
                status_p1_entry.config(borderwidth=2, state="normal")
                status_p1_entry.delete(0, "end")
                status_p1_entry.config(borderwidth=2, state="disable")
        except:
            pass
    if (astranotebook.index(astranotebook.select()) == 1):
        process()
        try:
            if (p2_number_of_entries > 0):
                # print("number_of_entries", number_of_entries)
                initial_table_rows = p2_number_of_entries + 1
                # dir_name = filedialog.askdirectory()  # asks user to choose a directory
                status_p2_entry.config(borderwidth=2, state="normal")
                status_p2_entry.delete(0, "end")
                status_p2_entry.insert(0, "Export initiated...")
                status_p2_entry.config(borderwidth=2, state="disable")
                if (export_dir_name == ''):
                    export_dir_name = filedialog.asksaveasfilename(title="Export As", filetypes=(("PDF", "*.pdf"), ("All Files", "*.*")),
                                                                   initialfile="Design Document")  # asks user to choose a directory
                    try:
                        os.chdir(os.path.dirname(export_dir_name))
                        export_file_name = os.path.basename(export_dir_name).split('.', 1)[0]
                        # copying the user file and directory to the save location
                        selected_directory = export_dir_name
                        selected_file_name = export_file_name
                    except:
                        pass
                        print("Path Nor defined")
                else:
                    export_dir_name = selected_directory
                    export_file_name = selected_file_name

                doc = Document()
                ######################################################
                # Adding A Header to the Document

                # Choosing the top most section of the page
                section = doc.sections[0]

                # Calling the header
                header = section.header

                # Calling the paragraph already present in
                # the header section
                header_para = header.paragraphs[0]

                # Adding text in the header
                header_para.text = "InPhase Power Technologies - ASTRA Sizing Document"
                header_para.alignment = 1
                ######################################################
                # Add a Title to the document
                doc.add_heading('Astra Rating', 1)

                # Creating a table object
                table = doc.add_table(rows=initial_table_rows, cols=3)
                # aliging table to center
                table.alignment = WD_TABLE_ALIGNMENT.CENTER

                # Creating table Heading
                table.cell(0, 0).text = 'Harmonics (n)'
                table.cell(0, 1).text = 'Current (P)'
                table.cell(0, 2).text = 'Current (N)'

                # vetrical centering

                table.cell(0, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.cell(0, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.cell(0, 2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

                for x in range(0, 3): make_rows_bold(table.columns[x])  # making bold and font size change

                # Align table column in table
                for cell in table.columns[0].cells: cell.width = Inches(2)
                for cell in table.columns[1].cells: cell.width = Inches(2)
                for cell in table.columns[2].cells: cell.width = Inches(2)

                for row in table.rows:
                    row.height = Cm(.75)

                table.style = 'Light Grid Accent 5'

                for x in range(0, p2_number_of_entries):
                    table.cell(x + 1, 0).text = str(p2_harmonicsentries[x].get())
                    table.cell(x + 1, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                for x in range(0, p2_number_of_entries):
                    table.cell(x + 1, 1).text = str(p2_harmonicsentries[x + 10].get()) + ' A'
                    table.cell(x + 1, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                for x in range(0, p2_number_of_entries):
                    table.cell(x + 1, 2).text = str(p2_harmonicsentries[x + 20].get()) + ' A'
                    table.cell(x + 1, 2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

                # making Bold of the both the column
                for x in range(0, 3): make_rows_bold(table.columns[x])
                # doc.save('rating.docx')
                doc.add_heading("SETTINGS", 1)
                # Adding A blank Line
                # para = doc.add_paragraph()
                # bold_para = para.add_run(''' ''')

                settings_table = doc.add_table(rows=5, cols=2)
                # aliging table to center
                settings_table.alignment = WD_TABLE_ALIGNMENT.CENTER

                # Creating table Heading
                settings_table.cell(0, 0).text = 'Settings'
                settings_table.cell(1, 0).text = 'Frequency'
                if (active_user == 0):
                    settings_table.cell(2, 0).text = 'Notch Profile'
                elif (active_user == 1):
                    settings_table.cell(2, 0).text = 'Base (n)'
                settings_table.cell(3, 0).text = 'Ambient Temperature (°C)'
                settings_table.cell(4, 0).text = 'Amplification Factor '
                settings_table.cell(0, 1).text = 'Value'
                if (page2_frequency == 50):
                    settings_table.cell(1, 1).text = "50 Hz"
                elif (page2_frequency == 60):
                    settings_table.cell(1, 1).text = "60 Hz"
                if (active_user == 0):
                    if (str(p2_basenentry.get()) == "6"): settings_table.cell(2, 1).text = "LOW"
                    if (str(p2_basenentry.get()) == "3"): settings_table.cell(2, 1).text = "MEDIUM"
                    settings_table.cell(3, 1).text = str(p2_ambtempentry.get()) + ' °C'
                    settings_table.cell(4, 1).text = str(p2_ambfactorentry.get())
                elif (active_user == 1):
                    settings_table.cell(2, 1).text = str(p2_basenentry.get())
                    settings_table.cell(3, 1).text = str(p2_ambtempentry.get()) + ' °C'
                    settings_table.cell(4, 1).text = str(p2_ambfactorentry.get())

                # vetrical centering

                settings_table.cell(0, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                settings_table.cell(0, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

                for x in range(0, 2): make_rows_bold(settings_table.columns[x])  # making bold and font size change

                # Align table column in table
                for cell in settings_table.columns[0].cells: cell.width = Inches(3)
                for cell in settings_table.columns[1].cells: cell.width = Inches(3)

                for row in settings_table.rows:
                    row.height = Cm(.75)

                settings_table.style = 'Light Grid Accent 5'

                #####################################################
                # Creating Result Table
                doc.add_heading("ASTRA RATING", 1)

                rating_table = doc.add_table(rows=3, cols=2)
                # aliging table to center
                rating_table.alignment = WD_TABLE_ALIGNMENT.CENTER

                # Creating table Heading
                rating_table.cell(0, 0).text = 'Description'
                rating_table.cell(1, 0).text = 'ASTRA RATING'
                rating_table.cell(2, 0).text = 'ASTRA RATING @' + str(p2_ambtempentry.get()) + '(°C)'
                rating_table.cell(0, 1).text = 'Current Rating'
                rating_table.cell(1, 1).text = str(p2_astrarating) + ' A'
                rating_table.cell(2, 1).text = str(p2_ambastrarating).lstrip('0') + ' A'

                # for x in range(1,4):
                rating_table.cell(0, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                rating_table.cell(0, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

                # vetrical centering

                rating_table.cell(0, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                rating_table.cell(0, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

                for x in range(0, 2): make_rows_bold(rating_table.columns[x])  # making bold and font size change

                # Align table column in table
                for cell in rating_table.columns[0].cells: cell.width = Inches(3)
                for cell in rating_table.columns[1].cells: cell.width = Inches(3)

                for row in rating_table.rows:
                    row.height = Cm(.75)

                rating_table.style = 'Light Grid Accent 5'

                #####################################################
                # adding a page break
                doc.add_page_break()
                #####################################################
                ###########################################
                # Reactive power Calaculation
                if p2_formula_combo.get() == "V, I, IPF, TPF":
                    doc.add_heading("Reactive Current details", 1)

                    reactive_table = doc.add_table(rows=9, cols=2)
                    # aliging table to center
                    reactive_table.alignment = WD_TABLE_ALIGNMENT.CENTER

                    # Creating table Heading
                    reactive_table.cell(0, 0).text = 'Description'
                    reactive_table.cell(1, 0).text = 'Voltage*'
                    reactive_table.cell(2, 0).text = 'Current*'
                    reactive_table.cell(3, 0).text = 'Active Power'
                    reactive_table.cell(4, 0).text = 'Apparent Power'
                    reactive_table.cell(5, 0).text = 'Reactive Power'
                    reactive_table.cell(6, 0).text = 'IPF*'
                    reactive_table.cell(7, 0).text = 'TPF*'
                    reactive_table.cell(8, 0).text = 'Reactive Current'
                    reactive_table.cell(0, 1).text = 'Value'
                    reactive_table.cell(1, 1).text = str(p2_volentry.get()) + ' V'
                    reactive_table.cell(2, 1).text = str(p2_curentry.get()) + ' A'
                    reactive_table.cell(3, 1).text = str(p2_kwentry.get()) + ' kW'
                    reactive_table.cell(4, 1).text = str(p2_kvaentry.get()) + ' kVA'
                    reactive_table.cell(5, 1).text = str(p2_kvarentry.get()) + ' kVAr'
                    reactive_table.cell(6, 1).text = str(format(float(p2_IPFentry.get()), ".2f"))
                    reactive_table.cell(7, 1).text = str(format(float(p2_TPFentry.get()), ".2f"))
                    reactive_table.cell(8, 1).text = str(p2_IQentry.get()) + ' A'

                    reactive_table.cell(0, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                    reactive_table.cell(0, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

                    # vetrical centering

                    reactive_table.cell(0, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                    reactive_table.cell(0, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

                    for x in range(0, 2): make_rows_bold(reactive_table.columns[x])  # making bold and font size change

                    # Align table column in table
                    for cell in reactive_table.columns[0].cells: cell.width = Inches(3)
                    for cell in reactive_table.columns[1].cells: cell.width = Inches(3)

                    for row in reactive_table.rows:
                        row.height = Cm(.75)

                    reactive_table.style = 'Light Grid Accent 5'
                    doc.add_paragraph(''' ''')
                    word = "* - Represents User Inputs"
                    p = doc.add_paragraph()
                    runner = p.add_run(word)
                    runner.bold = True
                    runner.italic = True

                elif p2_formula_combo.get() == "V, kW, IPF, TPF":
                    doc.add_heading("Reactive Current details", 1)

                    reactive_table = doc.add_table(rows=9, cols=2)
                    # aliging table to center
                    reactive_table.alignment = WD_TABLE_ALIGNMENT.CENTER

                    # Creating table Heading
                    reactive_table.cell(0, 0).text = 'Description'
                    reactive_table.cell(1, 0).text = 'Voltage*'
                    reactive_table.cell(2, 0).text = 'Current'
                    reactive_table.cell(3, 0).text = 'Active Power*'
                    reactive_table.cell(4, 0).text = 'Apparent Power'
                    reactive_table.cell(5, 0).text = 'Reactive Power'
                    reactive_table.cell(6, 0).text = 'IPF*'
                    reactive_table.cell(7, 0).text = 'TPF*'
                    reactive_table.cell(8, 0).text = 'Reactive Current'
                    reactive_table.cell(0, 1).text = 'Value'
                    reactive_table.cell(1, 1).text = str(p2_volentry.get()) + ' V'
                    reactive_table.cell(2, 1).text = str(p2_curentry.get()) + ' A'
                    reactive_table.cell(3, 1).text = str(p2_kwentry.get()) + ' kW'
                    reactive_table.cell(4, 1).text = str(p2_kvaentry.get()) + ' kVA'
                    reactive_table.cell(5, 1).text = str(p2_kvarentry.get()) + ' kVAr'
                    reactive_table.cell(6, 1).text = str(format(float(p2_IPFentry.get()), ".2f"))
                    reactive_table.cell(7, 1).text = str(format(float(p2_TPFentry.get()), ".2f"))
                    reactive_table.cell(8, 1).text = str(p2_IQentry.get()) + ' A'

                    reactive_table.cell(0, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                    reactive_table.cell(0, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

                    # vetrical centering

                    reactive_table.cell(0, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                    reactive_table.cell(0, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

                    for x in range(0, 2): make_rows_bold(reactive_table.columns[x])  # making bold and font size change

                    # Align table column in table
                    for cell in reactive_table.columns[0].cells: cell.width = Inches(3)
                    for cell in reactive_table.columns[1].cells: cell.width = Inches(3)

                    for row in reactive_table.rows:
                        row.height = Cm(.75)

                    reactive_table.style = 'Light Grid Accent 5'
                    doc.add_paragraph(''' ''')
                    word = "* - Represents User Inputs"
                    p = doc.add_paragraph()
                    runner = p.add_run(word)
                    runner.bold = True
                    runner.italic = True

                elif p2_formula_combo.get() == "V, kW, kVA, TPF":
                    doc.add_heading("Reactive Current details", 1)

                    reactive_table = doc.add_table(rows=9, cols=2)
                    # aliging table to center
                    reactive_table.alignment = WD_TABLE_ALIGNMENT.CENTER

                    # Creating table Heading
                    reactive_table.cell(0, 0).text = 'Description'
                    reactive_table.cell(1, 0).text = 'Voltage*'
                    reactive_table.cell(2, 0).text = 'Current'
                    reactive_table.cell(3, 0).text = 'Active Power*'
                    reactive_table.cell(4, 0).text = 'Apparent Power*'
                    reactive_table.cell(5, 0).text = 'Reactive Power'
                    reactive_table.cell(6, 0).text = 'IPF'
                    reactive_table.cell(7, 0).text = 'TPF*'
                    reactive_table.cell(8, 0).text = 'Reactive Current'
                    reactive_table.cell(0, 1).text = 'Value'
                    reactive_table.cell(1, 1).text = str(p2_volentry.get()) + ' V'
                    reactive_table.cell(2, 1).text = str(p2_curentry.get()) + ' A'
                    reactive_table.cell(3, 1).text = str(p2_kwentry.get()) + ' kW'
                    reactive_table.cell(4, 1).text = str(p2_kvaentry.get()) + ' kVA'
                    reactive_table.cell(5, 1).text = str(p2_kvarentry.get()) + ' kVAr'
                    reactive_table.cell(6, 1).text = str(format(float(p2_IPFentry.get()), ".2f"))
                    reactive_table.cell(7, 1).text = str(format(float(p2_TPFentry.get()), ".2f"))
                    reactive_table.cell(8, 1).text = str(p2_IQentry.get()) + ' A'

                    reactive_table.cell(0, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                    reactive_table.cell(0, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

                    # vetrical centering

                    reactive_table.cell(0, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                    reactive_table.cell(0, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

                    for x in range(0, 2): make_rows_bold(reactive_table.columns[x])  # making bold and font size change

                    # Align table column in table
                    for cell in reactive_table.columns[0].cells: cell.width = Inches(3)
                    for cell in reactive_table.columns[1].cells: cell.width = Inches(3)

                    for row in reactive_table.rows:
                        row.height = Cm(.75)

                    reactive_table.style = 'Light Grid Accent 5'
                    doc.add_paragraph(''' ''')
                    word = "* - Represents User Inputs"
                    p = doc.add_paragraph()
                    runner = p.add_run(word)
                    runner.bold = True
                    runner.italic = True

                elif p2_formula_combo.get() == "V, kVA, IPF, TPF":
                    doc.add_heading("Reactive Current details", 1)

                    reactive_table = doc.add_table(rows=9, cols=2)
                    # aliging table to center
                    reactive_table.alignment = WD_TABLE_ALIGNMENT.CENTER

                    # Creating table Heading
                    reactive_table.cell(0, 0).text = 'Description'
                    reactive_table.cell(1, 0).text = 'Voltage*'
                    reactive_table.cell(2, 0).text = 'Current'
                    reactive_table.cell(3, 0).text = 'Active Power'
                    reactive_table.cell(4, 0).text = 'Apparent Power*'
                    reactive_table.cell(5, 0).text = 'Reactive Power'
                    reactive_table.cell(6, 0).text = 'IPF*'
                    reactive_table.cell(7, 0).text = 'TPF*'
                    reactive_table.cell(8, 0).text = 'Reactive Current'
                    reactive_table.cell(0, 1).text = 'Value'
                    reactive_table.cell(1, 1).text = str(p2_volentry.get()) + ' V'
                    reactive_table.cell(2, 1).text = str(p2_curentry.get()) + ' A'
                    reactive_table.cell(3, 1).text = str(p2_kwentry.get()) + ' kW'
                    reactive_table.cell(4, 1).text = str(p2_kvaentry.get()) + ' kVA'
                    reactive_table.cell(5, 1).text = str(p2_kvarentry.get()) + ' kVAr'
                    reactive_table.cell(6, 1).text = str(format(float(p2_IPFentry.get()), ".2f"))
                    reactive_table.cell(7, 1).text = str(format(float(p2_TPFentry.get()), ".2f"))
                    reactive_table.cell(8, 1).text = str(p2_IQentry.get()) + ' A'

                    reactive_table.cell(0, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                    reactive_table.cell(0, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

                    # vetrical centering

                    reactive_table.cell(0, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                    reactive_table.cell(0, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

                    for x in range(0, 2): make_rows_bold(reactive_table.columns[x])  # making bold and font size change

                    # Align table column in table
                    for cell in reactive_table.columns[0].cells: cell.width = Inches(3)
                    for cell in reactive_table.columns[1].cells: cell.width = Inches(3)

                    for row in reactive_table.rows:
                        row.height = Cm(.75)

                    reactive_table.style = 'Light Grid Accent 5'
                    doc.add_paragraph(''' ''')
                    word = "* - Represents User Inputs"
                    p = doc.add_paragraph()
                    runner = p.add_run(word)
                    runner.bold = True
                    runner.italic = True

                elif p2_formula_combo.get() == "V, kVAR":
                    doc.add_heading("Reactive Current details", 1)

                    reactive_table = doc.add_table(rows=4, cols=2)
                    # aliging table to center
                    reactive_table.alignment = WD_TABLE_ALIGNMENT.CENTER

                    # Creating table Heading
                    reactive_table.cell(0, 0).text = 'Description'
                    reactive_table.cell(1, 0).text = 'Voltage*'
                    reactive_table.cell(2, 0).text = 'Reactive Power*'
                    reactive_table.cell(3, 0).text = 'Reactive Current'
                    reactive_table.cell(0, 1).text = 'Value'
                    reactive_table.cell(1, 1).text = str(p2_volentry.get()) + ' V'
                    reactive_table.cell(2, 1).text = str(p2_kvarentry.get()) + ' kVAr'
                    reactive_table.cell(3, 1).text = str(p2_IQentry.get()) + ' A'

                    reactive_table.cell(0, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                    reactive_table.cell(0, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

                    # vetrical centering

                    reactive_table.cell(0, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                    reactive_table.cell(0, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

                    for x in range(0, 2): make_rows_bold(reactive_table.columns[x])  # making bold and font size change

                    # Align table column in table
                    for cell in reactive_table.columns[0].cells: cell.width = Inches(3)
                    for cell in reactive_table.columns[1].cells: cell.width = Inches(3)

                    for row in reactive_table.rows:
                        row.height = Cm(.75)

                    reactive_table.style = 'Light Grid Accent 5'
                    doc.add_paragraph(''' ''')
                    word = "* - Represents User Inputs"
                    p = doc.add_paragraph()
                    runner = p.add_run(word)
                    runner.bold = True
                    runner.italic = True

                elif p2_formula_combo.get() == "IQ":
                    doc.add_heading("Reactive Current details", 1)

                    reactive_table = doc.add_table(rows=2, cols=2)
                    # aliging table to center
                    reactive_table.alignment = WD_TABLE_ALIGNMENT.CENTER

                    # Creating table Heading
                    reactive_table.cell(0, 0).text = 'Description'
                    reactive_table.cell(1, 0).text = 'Reactive Current*'
                    reactive_table.cell(0, 1).text = 'Value'
                    reactive_table.cell(1, 1).text = str(p2_IQentry.get()) + ' A'

                    reactive_table.cell(0, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                    reactive_table.cell(0, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

                    # vetrical centering

                    reactive_table.cell(0, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                    reactive_table.cell(0, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

                    for x in range(0, 2): make_rows_bold(reactive_table.columns[x])  # making bold and font size change

                    # Align table column in table
                    for cell in reactive_table.columns[0].cells: cell.width = Inches(3)
                    for cell in reactive_table.columns[1].cells: cell.width = Inches(3)

                    for row in reactive_table.rows:
                        row.height = Cm(.75)

                    reactive_table.style = 'Light Grid Accent 5'
                    doc.add_paragraph(''' ''')
                    word = "* - Represents User Inputs"
                    p = doc.add_paragraph()
                    runner = p.add_run(word)
                    runner.bold = True
                    runner.italic = True
                ###########################################
                #####################################################
                # Deleting the place holder
                if p2_comment_box_message.get('1.0', 'end-1c') == placeholder_text:
                    pass
                else:
                    doc.add_heading("Comments", 1)
                    para = doc.add_paragraph(str(page2_comment_box_message))
                # bold_para = para.add_run('''\n '''+str(page1_comment_box_message))
                ############################################

                #####################################################
                doc.add_heading("Prepared By", 1)
                para = doc.add_paragraph()
                bold_para = para.add_run(''' \n Advanced Power Quality''')
                ############################################

                # Adding a border to a Page
                sec_pr = doc.sections[0]._sectPr  # get the section properties el
                # create new borders el
                pg_borders = OxmlElement('w:pgBorders')
                # specifies how the relative positioning of the borders should be calculated
                pg_borders.set(qn('w:offsetFrom'), 'page')
                for border_name in ('top', 'left', 'bottom', 'right',):  # set all borders
                    border_el = OxmlElement(f'w:{border_name}')
                    border_el.set(qn('w:val'), 'single')  # a single line
                    border_el.set(qn('w:sz'), '4')  # for meaning of  remaining attrs please look docs
                    border_el.set(qn('w:space'), '24')
                    border_el.set(qn('w:color'), 'auto')
                    pg_borders.append(border_el)  # register single border to border el
                sec_pr.append(pg_borders)  # apply border changes to section
                ############################################
                doc.save(str(export_file_name) + ".docx")
                convert(str(export_file_name) + ".docx")
                os.remove(str(export_file_name) + ".docx")
                export_progress_flag = 1
                save_nfo()
                export_progress_flag = 0
                status_p2_entry.config(borderwidth=2, state="normal")
                status_p2_entry.delete(0, "end")
                status_p2_entry.config(borderwidth=2, state="disable")
                disable_modified_flag()
                messagebox.showinfo("Export Information", export_file_name + ".pdf" + " Exported")

            else:
                mandatory_entries_check()
                status_p2_entry.config(borderwidth=2, state="normal")
                status_p2_entry.delete(0, "end")
                status_p2_entry.config(borderwidth=2, state="disable")
        except:
            pass
    if (astranotebook.index(astranotebook.select()) == 2):
        process()
        try:
            if (p3_number_of_entries > 0):
                # print("number_of_entries", number_of_entries)
                initial_table_rows = p3_number_of_entries + 1
                # dir_name = filedialog.askdirectory()  # asks user to choose a directory
                status_p3_entry.config(borderwidth=2, state="normal")
                status_p3_entry.delete(0, "end")
                status_p3_entry.insert(0, "Export initiated...")
                status_p3_entry.config(borderwidth=2, state="disable")
                if (export_dir_name == ''):
                    export_dir_name = filedialog.asksaveasfilename(title="Export As",
                                                                   filetypes=(("PDF", "*.pdf"), ("All Files", "*.*")),
                                                                   initialfile="Design Document")  # asks user to choose a directory
                    try:
                        os.chdir(os.path.dirname(export_dir_name))
                        export_file_name = os.path.basename(export_dir_name).split('.', 1)[0]
                        # copying the user file and directory to the save location
                        selected_directory = export_dir_name
                        selected_file_name = export_file_name
                    except:
                        pass
                        print("Path Nor defined")
                else:
                    export_dir_name = selected_directory
                    export_file_name = selected_file_name

                doc = Document()
                ######################################################
                # Adding A Header to the Document
                # Choosing the top most section of the page
                section = doc.sections[0]

                # Calling the header
                header = section.header

                # Calling the paragraph already present in
                # the header section
                header_para = header.paragraphs[0]

                # Adding text in the header
                header_para.text = "InPhase Power Technologies - ASTRA Sizing Document"
                header_para.alignment = 1
                ######################################################
                # Add a Title to the document
                doc.add_heading('Astra Rating', 1)

                # Creating a table object
                table = doc.add_table(rows=initial_table_rows, cols=2)

                # aliging table to center
                table.alignment = WD_TABLE_ALIGNMENT.CENTER

                # Creating table Heading
                table.cell(0, 0).text = 'Harmonics (n)'
                table.cell(0, 1).text = 'Current (N)'

                # vetrical centering

                table.cell(0, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.cell(0, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

                for x in range(0, 2): make_rows_bold(table.columns[x])  # making bold and font size change

                # Align table column in table
                for cell in table.columns[0].cells: cell.width = Inches(2)
                for cell in table.columns[1].cells: cell.width = Inches(2)

                for row in table.rows:
                    row.height = Cm(.75)

                table.style = 'Light Grid Accent 5'

                for x in range(0, p3_number_of_entries):
                    table.cell(x + 1, 0).text = str(p3_harmonicsentries[x].get())
                    table.cell(x + 1, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                for x in range(0, p3_number_of_entries):
                    table.cell(x + 1, 1).text = str(p3_harmonicsentries[x + 10].get()) + ' A'
                    table.cell(x + 1, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

                # making Bold of the both the column
                for x in range(0, 2): make_rows_bold(table.columns[x])
                # doc.save('rating.docx')
                doc.add_heading("SETTINGS", 1)

                settings_table = doc.add_table(rows=5, cols=2)
                # aliging table to center
                settings_table.alignment = WD_TABLE_ALIGNMENT.CENTER

                # Creating table Heading
                settings_table.cell(0, 0).text = 'Settings'
                settings_table.cell(1, 0).text = 'Frequency'
                if (active_user == 0):
                    settings_table.cell(2, 0).text = 'Notch Profile'
                elif (active_user == 1):
                    settings_table.cell(2, 0).text = 'Base (n)'
                settings_table.cell(3, 0).text = 'Ambient Temperature (°C)'
                settings_table.cell(4, 0).text = 'Amplification Factor '
                settings_table.cell(0, 1).text = 'Value'
                if (page3_frequency == 50):
                    settings_table.cell(1, 1).text = "50 Hz"
                elif (page3_frequency == 60):
                    settings_table.cell(1, 1).text = "60 Hz"
                if (active_user == 0):
                    if (str(p3_basenentry.get()) == "6"): settings_table.cell(2, 1).text = "LOW"
                    if (str(p3_basenentry.get()) == "3"): settings_table.cell(2, 1).text = "MEDIUM"
                    settings_table.cell(3, 1).text = str(p3_ambtempentry.get()) + ' °C'
                    settings_table.cell(4, 1).text = str(p3_ambfactorentry.get())
                elif (active_user == 1):
                    settings_table.cell(2, 1).text = str(p3_basenentry.get())
                    settings_table.cell(3, 1).text = str(p3_ambtempentry.get()) + ' °C'
                    settings_table.cell(4, 1).text = str(p3_ambfactorentry.get())

                # vetrical centering

                settings_table.cell(0, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                settings_table.cell(0, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

                for x in range(0, 2): make_rows_bold(settings_table.columns[x])  # making bold and font size change

                # Align table column in table
                for cell in settings_table.columns[0].cells: cell.width = Inches(3)
                for cell in settings_table.columns[1].cells: cell.width = Inches(3)

                for row in settings_table.rows:
                    row.height = Cm(.75)

                settings_table.style = 'Light Grid Accent 5'

                #####################################################
                # Creating Result Table
                doc.add_heading("ASTRA RATING", 1)

                rating_table = doc.add_table(rows=3, cols=2)
                # aliging table to center
                rating_table.alignment = WD_TABLE_ALIGNMENT.CENTER

                # Creating table Heading
                rating_table.cell(0, 0).text = 'Description'
                rating_table.cell(1, 0).text = 'ASTRA RATING'
                rating_table.cell(2, 0).text = 'ASTRA RATING @' + str(p3_ambtempentry.get()) + '(°C)'
                rating_table.cell(0, 1).text = 'Current Rating'
                rating_table.cell(1, 1).text = str(p3_astrarating) + ' A'
                rating_table.cell(2, 1).text = str(p3_ambastrarating).lstrip('0') + ' A'

                # for x in range(1,4):
                rating_table.cell(0, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                rating_table.cell(0, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

                # vetrical centering

                rating_table.cell(0, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                rating_table.cell(0, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

                for x in range(0, 2): make_rows_bold(rating_table.columns[x])  # making bold and font size change

                # Align table column in table
                for cell in rating_table.columns[0].cells: cell.width = Inches(3)
                for cell in rating_table.columns[1].cells: cell.width = Inches(3)

                for row in rating_table.rows:
                    row.height = Cm(.75)

                rating_table.style = 'Light Grid Accent 5'

                #####################################################
                # adding a page break
                doc.add_page_break()
                #####################################################
                #####################################################
                # Deleting the place holder
                if p3_comment_box_message.get('1.0', 'end-1c') == placeholder_text:
                    pass
                else:
                    doc.add_heading("Comments", 1)
                    para = doc.add_paragraph(str(page3_comment_box_message))
                # bold_para = para.add_run('''\n '''+str(page1_comment_box_message))
                ############################################

                #####################################################
                doc.add_heading("Prepared By", 1)
                para = doc.add_paragraph()
                bold_para = para.add_run(''' \n Advanced Power Quality''')
                ############################################

                # Adding a border to a Page
                sec_pr = doc.sections[0]._sectPr  # get the section properties el
                # create new borders el
                pg_borders = OxmlElement('w:pgBorders')
                # specifies how the relative positioning of the borders should be calculated
                pg_borders.set(qn('w:offsetFrom'), 'page')
                for border_name in ('top', 'left', 'bottom', 'right',):  # set all borders
                    border_el = OxmlElement(f'w:{border_name}')
                    border_el.set(qn('w:val'), 'single')  # a single line
                    border_el.set(qn('w:sz'), '4')  # for meaning of  remaining attrs please look docs
                    border_el.set(qn('w:space'), '24')
                    border_el.set(qn('w:color'), 'auto')
                    pg_borders.append(border_el)  # register single border to border el
                sec_pr.append(pg_borders)  # apply border changes to section
                ############################################
                doc.save(str(export_file_name) + ".docx")
                convert(str(export_file_name) + ".docx")
                os.remove(str(export_file_name) + ".docx")
                export_progress_flag = 1
                save_nfo()
                export_progress_flag = 0
                status_p3_entry.config(borderwidth=2, state="normal")
                status_p3_entry.delete(0, "end")
                status_p3_entry.config(borderwidth=2, state="disable")
                disable_modified_flag()
                messagebox.showinfo("Export Information", export_file_name + ".pdf" + " Exported")

            else:
                mandatory_entries_check()
                status_p3_entry.config(borderwidth=2, state="normal")
                status_p3_entry.delete(0, "end")
                status_p3_entry.config(borderwidth=2, state="disable")
        except:
            pass

#-------------------- True power factor export to PDF ------------------#
    if (astranotebook.index(astranotebook.select()) == 3):
        submit()
        try:
            if primary_entry.get() and secondary_entry.get():
                pri_less_sec.config(text="")
                #astranotebook.tab(3, text="True Power Factor")
                #p4_modified_indication.config(text="")
                #head_1.config(text="True Power Factor Performance Calculator")
                status_p4_entry.config(text="Export initiated...")

                if export_dir_name == '':
                    export_dir_name = filedialog.asksaveasfilename(title="Export As",filetypes=(("PDF", "*.docx"), ("All Files", "*.*")),initialfile="Design Document")

                    try:
                        os.chdir(os.path.dirname(export_dir_name))
                        export_file_name = os.path.basename(export_dir_name).split('.', 1)[0]
                        # Set selected directory and file name
                        selected_directory = export_dir_name
                        selected_file_name = export_file_name
                    except OSError as e:
                        print("Error changing directory:", e)
                else:
                    export_dir_name = selected_directory
                    export_file_name = selected_file_name

                doc = Document()
                # Adding Header to the Document
                header = doc.sections[0].header
                header_para = header.paragraphs[0]
                header_para.text = "InPhase Power Technologies - ASTRA Sizing Document"
                header_para.alignment = WD_TABLE_ALIGNMENT.CENTER

                # Add Title to the document
                doc.add_heading('True Power Factor Performance Calculator :', level=1)

                # Create table
                table = doc.add_table(rows=13, cols=6)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER

                headings = ['ID', 'Panel Rating', 'Optimal CT Ratio', 'Condition', 'Minimum kVAr', 'Minimum kW']
                #headings.extend(item[:4] + [item[4], item[5]] for item in data_2[1:])
                column_widths = [Inches(0.5), Inches(2), Inches(2), Inches(1.5), Inches(1.5), Inches(1.5)]
                print(headings)
                # Assuming table is initialized properly before this loop
                for i, (heading, width) in enumerate(zip(headings, column_widths)):
                    cell = table.cell(0, i)
                    cell.text = heading
                    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    cell.width = width

                    # Align text vertically and horizontally
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                # Set row height
                for row in table.rows:
                    row.height = Cm(0.75)
                # Set table style
                table.style = 'Light Grid Accent 5'

                # Update the cell text using loops
                for i in range(12):
                    table.cell(i + 1, 0).text = str(rec_tab_val[i + 1][0])
                    table.cell(i + 1, 1).text = str(rec_tab_val[i + 1][1])
                    table.cell(i + 1, 2).text = str(rec_tab_val[i + 1][2])
                    table.cell(i + 1, 3).text = str(rec_tab_val[i + 1][3])
                    table.cell(i + 1, 4).text = str(rec_tab_val[i + 1][4])
                    table.cell(i + 1, 5).text = str(rec_tab_val[i + 1][5])

                # Make table headers bold and adjust row height
                for row in table.rows:
                    for cell in row.cells:
                        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                # Add Title to the document
                doc.add_heading('Suggestion :', level=1)

                # Create second table
                table_1 = doc.add_table(rows=4, cols=3)
                table_1.alignment = WD_TABLE_ALIGNMENT.CENTER

                headings_1 = ['ID', 'Panel Rating', 'Minimum kW']
                column_widths_1 = [Inches(0.5), Inches(2), Inches(2)]

                # Set table headers and properties for second table
                for i, (heading, width) in enumerate(zip(headings_1, column_widths_1)):
                    cell = table_1.cell(0, i)
                    cell.text = heading
                    cell.width = width

                    # Align text vertically and horizontally
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                # Update the cell text using loops for second table
                for i in range(3):
                    panel_id = minimum_kw[i][0]
                    panel_rating = minimum_kw[i][1]
                    panel_kw = minimum_kw[i][2]

                    table_1.cell(i + 1, 0).text = str(panel_id)
                    table_1.cell(i + 1, 1).text = str(panel_rating)
                    table_1.cell(i + 1, 2).text = str(panel_kw)

                    # Align text vertically and horizontally
                    for j in range(3):
                        cell = table_1.cell(i + 1, j)
                        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                # Set row height
                for row_1 in table_1.rows:
                    row_1.height = Cm(0.75)
                # Set table style
                table_1.style = 'Light Grid Accent 5'

                doc.add_paragraph(''' ''')
                word = "* - Represents User Inputs"
                p = doc.add_paragraph()
                runner = p.add_run(word)
                runner.bold = True
                runner.italic = True

                doc.add_heading("Prepared By", 1)
                para = doc.add_paragraph()
                bold_para = para.add_run(''' \n Advanced Power Quality''')

                section = doc.sections[0]
                section.left_margin = Inches(0.7)
                section.right_margin = Inches(0.7)
                section.top_margin = Inches(1.0)
                section.bottom_margin = Inches(1.0)
                # Set border around the entire page with 0.5 inches margin
                # Adding a border to a Page
                sec_pr = doc.sections[0]._sectPr  # get the section properties el
                # create new borders el
                pg_borders = OxmlElement('w:pgBorders')
                # specifies how the relative positioning of the borders should be calculated
                pg_borders.set(qn('w:offsetFrom'), 'page')
                for border_name in ('top', 'left', 'bottom', 'right',):  # set all borders
                    border_el = OxmlElement(f'w:{border_name}')
                    border_el.set(qn('w:val'), 'single')  # a single line
                    border_el.set(qn('w:sz'), '4')  # for meaning of  remaining attrs please look docs
                    border_el.set(qn('w:space'), '24')
                    border_el.set(qn('w:color'), 'auto')
                    pg_borders.append(border_el)  # register single border to border el
                sec_pr.append(pg_borders)  # apply border changes to section
                doc.save(str(export_file_name) + ".docx")
                convert(str(export_file_name) + ".docx")
                os.remove(str(export_file_name) + ".docx")

                save_nfo()
                status_p4_entry.config(text="Export Completed...")
                astranotebook.tab(3, text="True Power Factor")
                head_notebook.config(text="True Power Factor Performance Calculator  ")
            else:
                messagebox.showerror('Input Error', "Fill the both entries")
                status_p4_entry.config(text="You need to enter the Primary and Secondary values")

        except:
            pass

#-------------------- True power factor export code end ---------------------#

#--------------------------Export PDF for astra is ended----------------------#

#--------------------------Astra clear screen code started--------------------------#
def clear_results():
    # clearing the page -1 results which will be called at the time of process

    if (astranotebook.index(astranotebook.select()) == 0):
        status_p1_entry.config(borderwidth=2, state="normal")
        status_p1_entry.delete(0, "end")
        status_p1_entry.insert(0, (""))
        status_p1_entry.config(borderwidth=2, state="disable")

        for x in range(0, 20):
            genrated_data[x].config(borderwidth=2, state="normal")
            genrated_data[x].delete(0, "end")
            genrated_data[x].insert(0, str(empty_entry.get(1.0, "end-1c")))
            genrated_data[x].config(borderwidth=2, state="disable")

        # Clearing sum of I*I
        currentsqrentry.config(borderwidth=2, state="normal")
        currentsqrentry.delete(0, "end")
        currentsqrentry.insert(0, str(empty_entry.get(1.0, "end-1c")))
        currentsqrentry.config(borderwidth=2, state="disable")

        # Clearing sum of ILin
        currentlinentry.config(borderwidth=2, state="normal")
        currentlinentry.delete(0, "end")
        currentlinentry.insert(0, str(empty_entry.get(1.0, "end-1c")))
        currentlinentry.config(borderwidth=2, state="disable")

        # Clearing Astra Rating
        AHFsizeentry.config(borderwidth=2, state="normal")
        AHFsizeentry.delete(0, "end")
        AHFsizeentry.insert(0, str(empty_entry.get(1.0, "end-1c")))
        AHFsizeentry.config(borderwidth=2, state="disable")
        # Clearing Astra Rating (ambient)
        AHFsize1entry.config(borderwidth=2, state="normal")
        AHFsize1entry.delete(0, "end")
        AHFsize1entry.insert(0, str(empty_entry.get(1.0, "end-1c")))
        AHFsize1entry.config(borderwidth=2, state="disable")

    elif (astranotebook.index(astranotebook.select()) == 1):

        status_p2_entry.config(borderwidth=2, state="normal")
        status_p2_entry.delete(0, "end")
        status_p2_entry.insert(0, (""))
        status_p2_entry.config(borderwidth=2, state="disable")

        for x in range(0, 40):
            p2_genrated_data[x].config(borderwidth=2, state="normal")
            p2_genrated_data[x].delete(0, "end")
            p2_genrated_data[x].insert(0, str(empty_entry.get(1.0, "end-1c")))
            p2_genrated_data[x].config(borderwidth=2, state="disable")

        # Clearing Astra Rating
        p2_AHFsizeentry.config(borderwidth=2, state="normal")
        p2_AHFsizeentry.delete(0, "end")
        p2_AHFsizeentry.insert(0, str(empty_entry.get(1.0, "end-1c")))
        p2_AHFsizeentry.config(borderwidth=2, state="disable")
        # Clearing Astra Rating (ambient)
        p2_AHFsize1entry.config(borderwidth=2, state="normal")
        p2_AHFsize1entry.delete(0, "end")
        p2_AHFsize1entry.insert(0, str(empty_entry.get(1.0, "end-1c")))
        p2_AHFsize1entry.config(borderwidth=2, state="disable")
    elif (astranotebook.index(astranotebook.select()) == 2):
        status_p3_entry.config(borderwidth=2, state="normal")
        status_p3_entry.delete(0, "end")
        status_p3_entry.insert(0, (""))
        status_p3_entry.config(borderwidth=2, state="disable")

        for x in range(0, 40):
            p2_genrated_data[x].config(borderwidth=2, state="normal")
            p2_genrated_data[x].delete(0, "end")
            p2_genrated_data[x].insert(0, str(empty_entry.get(1.0, "end-1c")))
            p2_genrated_data[x].config(borderwidth=2, state="disable")

        # Clearing Astra Rating
        p3_AHFsizeentry.config(borderwidth=2, state="normal")
        p3_AHFsizeentry.delete(0, "end")
        p3_AHFsizeentry.insert(0, str(empty_entry.get(1.0, "end-1c")))
        p3_AHFsizeentry.config(borderwidth=2, state="disable")
        # Clearing Astra Rating (ambient)
        p3_AHFsize1entry.config(borderwidth=2, state="normal")
        p3_AHFsize1entry.delete(0, "end")
        p3_AHFsize1entry.insert(0, str(empty_entry.get(1.0, "end-1c")))
        p3_AHFsize1entry.config(borderwidth=2, state="disable")
    '''elif (astranotebook.index(astranotebook.select()) == 3):
        primary_entry.config(borderwidth=2, state="normal")
        primary_entry.delete(0, "end")
        primary_entry.insert(0, (""))
        primary_entry.config(borderwidth=2, state="disable")

        for x in range(0, 2):
            p2_genrated_data[x].config(borderwidth=2, state="normal")
            p2_genrated_data[x].delete(0, "end")
            p2_genrated_data[x].insert(0, str(empty_entry.get(1.0, "end-1c")))
            p2_genrated_data[x].config(borderwidth=2, state="disable")'''

#-----------------------Astra clear screen code ended-----------------------#

#---------------------Mandatory enrtries indication code started--------------------#
def mandatory_entries_check():
    # mandatory enties check for the page - 1

    global fault, p2_fault, p3_fault
    global page1_frequency, page2_frequency, page3_frequency

    if (astranotebook.index(astranotebook.select()) == 0):
        try:
            if formula_combo.get() == "select":
                fault = fault + 1
            elif formula_combo.get() == "V, I, IPF, TPF":

                if (str(volentry.get()) == str(empty_entry.get(1.0, "end-1c")) or
                        str(curentry.get()) == str(empty_entry.get(1.0, "end-1c")) or
                        str(IPFentry.get()) == str(empty_entry.get(1.0, "end-1c")) or
                        str(TPFentry.get()) == str(empty_entry.get(1.0, "end-1c"))):
                    fault = fault + 1
            elif formula_combo.get() == "V, kW, IPF, TPF":
                if (str(volentry.get()) == str(empty_entry.get(1.0, "end-1c")) or
                        str(kwentry.get()) == str(empty_entry.get(1.0, "end-1c")) or
                        str(IPFentry.get()) == str(empty_entry.get(1.0, "end-1c")) or
                        str(TPFentry.get()) == str(empty_entry.get(1.0, "end-1c"))):
                    fault = fault + 1
            elif formula_combo.get() == "V, kW, kVA, TPF":
                if (str(volentry.get()) == str(empty_entry.get(1.0, "end-1c")) or
                        str(kwentry.get()) == str(empty_entry.get(1.0, "end-1c")) or
                        str(kvaentry.get()) == str(empty_entry.get(1.0, "end-1c")) or
                        str(TPFentry.get()) == str(empty_entry.get(1.0, "end-1c"))):
                    fault = fault + 1
            elif formula_combo.get() == "V, kVA, IPF, TPF":
                if (str(volentry.get()) == str(empty_entry.get(1.0, "end-1c")) or
                        str(kvaentry.get()) == str(empty_entry.get(1.0, "end-1c")) or
                        str(IPFentry.get()) == str(empty_entry.get(1.0, "end-1c")) or
                        str(TPFentry.get()) == str(empty_entry.get(1.0, "end-1c"))):
                    fault = fault + 1
            elif formula_combo.get() == "V, kVAR":
                if (str(volentry.get()) == str(empty_entry.get(1.0, "end-1c")) or
                        str(kvarentry.get()) == str(empty_entry.get(1.0, "end-1c"))):
                    fault = fault + 1
            elif formula_combo.get() == "IQ":
                if (str(IQentry.get()) == str(empty_entry.get(1.0, "end-1c"))):
                    fault = fault + 1
                    IQlable.config(fg="red")
                else:
                    IQlable.config(fg="black")
            if (page1_frequency == 50 or page1_frequency == 60):
                frequencylable.config(fg="black")
            else:
                frequencylable.config(fg="red")
                fault = fault + 1
        except:
            fault = 1000
            pass
        try:
            if (str(basenentry.get()) == str(empty_entry.get(1.0, "end-1c"))):
                basenlable.config(fg="red")
                fault = fault + 1
            else:
                int(str(basenentry.get())) + int(str(basenentry.get()))
                basenlable.config(fg="black")
                basenentry.config(fg="black")
        except ValueError:
            basenentry.config(fg="red")
            fault = fault + 1
            pass
        try:
            if (str(ambtempentry.get()) == str(empty_entry.get(1.0, "end-1c"))):
                ambtemplable.config(fg="red")
                fault = fault + 1
            else:
                int(str(ambtempentry.get())) + int(str(ambtempentry.get()))
                ambtemplable.config(fg="black")
                ambtempentry.config(fg="black")
        except ValueError:
            ambtempentry.config(fg="red")
            fault = fault + 1
            pass
        try:
            if (str(ambfactorentry.get()) == str(empty_entry.get(1.0, "end-1c"))):
                ambfactorlable.config(fg="red")
                fault = fault + 1
            else:
                float(str(ambfactorentry.get())) + float(str(ambfactorentry.get()))
                ambfactorlable.config(fg="black")
                ambfactorentry.config(fg="black")
        except ValueError:
            ambfactorentry.config(fg="red")
            fault = fault + 1
            pass
    elif (astranotebook.index(astranotebook.select()) == 1):
        try:
            if p2_formula_combo.get() == "select":
                p2_fault = p2_fault + 1
            elif p2_formula_combo.get() == "V, I, IPF, TPF":

                if (str(p2_volentry.get()) == str(p2_empty_entry.get(1.0, "end-1c")) or
                        str(p2_curentry.get()) == str(p2_empty_entry.get(1.0, "end-1c")) or
                        str(p2_IPFentry.get()) == str(p2_empty_entry.get(1.0, "end-1c")) or
                        str(p2_TPFentry.get()) == str(p2_empty_entry.get(1.0, "end-1c"))):
                    p2_fault = p2_fault + 1
            elif p2_formula_combo.get() == "V, kW, IPF, TPF":
                if (str(p2_volentry.get()) == str(p2_empty_entry.get(1.0, "end-1c")) or
                        str(p2_kwentry.get()) == str(p2_empty_entry.get(1.0, "end-1c")) or
                        str(p2_IPFentry.get()) == str(p2_empty_entry.get(1.0, "end-1c")) or
                        str(p2_TPFentry.get()) == str(p2_empty_entry.get(1.0, "end-1c"))):
                    p2_fault = p2_fault + 1
            elif p2_formula_combo.get() == "V, kW, kVA, TPF":
                if (str(p2_volentry.get()) == str(p2_empty_entry.get(1.0, "end-1c")) or
                        str(p2_kwentry.get()) == str(p2_empty_entry.get(1.0, "end-1c")) or
                        str(p2_kvaentry.get()) == str(p2_empty_entry.get(1.0, "end-1c")) or
                        str(p2_TPFentry.get()) == str(p2_empty_entry.get(1.0, "end-1c"))):
                    p2_fault = p2_fault + 1
            elif p2_formula_combo.get() == "V, kVA, IPF, TPF":
                if (str(p2_volentry.get()) == str(p2_empty_entry.get(1.0, "end-1c")) or
                        str(p2_kvaentry.get()) == str(p2_empty_entry.get(1.0, "end-1c")) or
                        str(p2_IPFentry.get()) == str(p2_empty_entry.get(1.0, "end-1c")) or
                        str(p2_TPFentry.get()) == str(p2_empty_entry.get(1.0, "end-1c"))):
                    p2_fault = p2_fault + 1
            elif p2_formula_combo.get() == "V, kVAR":
                if (str(p2_volentry.get()) == str(p2_empty_entry.get(1.0, "end-1c")) or
                        str(p2_kvarentry.get()) == str(p2_empty_entry.get(1.0, "end-1c"))):
                    p2_fault = p2_fault + 1
            elif p2_formula_combo.get() == "IQ":
                if (str(p2_IQentry.get()) == str(p2_empty_entry.get(1.0, "end-1c"))):
                    p2_fault = p2_fault + 1
                    p2_IQlable.config(fg="red")
                else:
                    p2_IQlable.config(fg="black")
            if (page2_frequency == 50 or page2_frequency == 60):
                p2_frequencylable.config(fg="black")
            else:
                p2_frequencylable.config(fg="red")
                p2_fault = p2_fault + 1
        except:
            p2_fault = 1000
            pass
        try:
            if (str(p2_basenentry.get()) == str(p2_empty_entry.get(1.0, "end-1c"))):
                p2_basenlable.config(fg="red")
                p2_fault = p2_fault + 1
            else:
                int(str(p2_basenentry.get())) + int(str(p2_basenentry.get()))
                p2_basenlable.config(fg="black")
                p2_basenentry.config(fg="black")
        except ValueError:
            basenentry.config(fg="red")
            p2_fault = p2_fault + 1
            pass
        try:
            if (str(p2_ambtempentry.get()) == str(p2_empty_entry.get(1.0, "end-1c"))):
                p2_ambtemplable.config(fg="red")
                p2_fault = p2_fault + 1
            else:
                int(str(p2_ambtempentry.get())) + int(str(p2_ambtempentry.get()))
                p2_ambtemplable.config(fg="black")
                p2_ambtempentry.config(fg="black")
        except ValueError:
            p2_ambtempentry.config(fg="red")
            p2_fault = p2_fault + 1
            pass
        try:
            if (str(p2_ambfactorentry.get()) == str(p2_empty_entry.get(1.0, "end-1c"))):
                p2_ambfactorlable.config(fg="red")
                p2_fault = p2_fault + 1
            else:
                float(str(p2_ambfactorentry.get())) + float(str(p2_ambfactorentry.get()))
                p2_ambfactorlable.config(fg="black")
                p2_ambfactorentry.config(fg="black")
        except ValueError:
            p2_ambfactorentry.config(fg="red")
            p2_fault = p2_fault + 1
            pass
    elif (astranotebook.index(astranotebook.select()) == 2):
        if (page3_frequency == 50 or page3_frequency == 60):
            p3_frequencylable.config(fg="black")
        else:
            p3_frequencylable.config(fg="red")
            p3_fault = p3_fault + 1
        try:
            if (str(p3_basenentry.get()) == str(p3_empty_entry.get(1.0, "end-1c"))):
                p3_basenlable.config(fg="red")
                p3_fault = p3_fault + 1
            else:
                int(str(p3_basenentry.get())) + int(str(p3_basenentry.get()))
                p3_basenlable.config(fg="black")
                p3_basenentry.config(fg="black")
        except ValueError:
            basenentry.config(fg="red")
            p3_fault = p3_fault + 1
            pass
        try:
            if (str(p3_ambtempentry.get()) == str(p3_empty_entry.get(1.0, "end-1c"))):
                p3_ambtemplable.config(fg="red")
                p3_fault = p3_fault + 1
            else:
                int(str(p3_ambtempentry.get())) + int(str(p3_ambtempentry.get()))
                p3_ambtemplable.config(fg="black")
                p3_ambtempentry.config(fg="black")
        except ValueError:
            p3_ambtempentry.config(fg="red")
            p3_fault = p3_fault + 1
            pass
        try:
            if (str(p3_ambfactorentry.get()) == str(p3_empty_entry.get(1.0, "end-1c"))):
                p3_ambfactorlable.config(fg="red")
                p3_fault = p3_fault + 1
            else:
                float(str(p3_ambfactorentry.get())) + float(str(p3_ambfactorentry.get()))
                p3_ambfactorlable.config(fg="black")
                p3_ambfactorentry.config(fg="black")
        except ValueError:
            p3_ambfactorentry.config(fg="red")
            p3_fault = p3_fault + 1
            pass
#------------------------Mandatory entries code ended-------------------#

#-----------Astra process button code started-----------------#
def process():
    # calculation function which will be processed after the page- 1 process click

    global fault, p2_fault, p3_fault
    fault = 0
    p2_fault = 0
    p3_fault = 0
    mandatory_entries_check()
    if (astranotebook.index(astranotebook.select()) == 0):
        global isquareentries, currentsqrresult, currentlinresult, astrarating, ambastrarating, harmonics_entry_list, current_entry_list
        global number_of_entries
        isquareentries = []
        harmonics_entry_list = []
        current_entry_list = []
        currentsqrresult = 0
        currentlinresult = 0
        number_of_entries = 0
        # I*I Calculation
        if (fault == 0):
            valueerrorflag = 0
            clear_results()
            # -->disable_modified_flag()
            for x in range(0, 10):
                if (str(harmonicsentries[x].get()) != str(empty_entry.get(1.0, "end-1c")) or str(
                        harmonicsentries[x + 10].get()) != str(empty_entry.get(1.0, "end-1c"))):
                    try:
                        harmonics_entry_list.append(int(harmonicsentries[x].get()))
                        current_entry_list.append(int(harmonicsentries[x + 10].get()))
                        number_of_entries = number_of_entries + 1
                        if (x == 0):
                            # isquareentries.append(int(harmonicsentries[x+10].get()) * int(harmonicsentries[x + 10].get()))
                            isquareentries.append(int(harmonicsentries[x + 10].get()) ** 2)
                            harmonicsentries[x + 10].config(fg="black")
                        else:
                            # isquareentries.append(int(harmonicsentries[x+10].get()) * int(harmonicsentries[x + 10].get()))
                            isquareentries.append(
                                (int(harmonicsentries[x + 10].get()) ** 2) * (float(ambfactorentry.get()) ** 2))
                            harmonicsentries[x + 10].config(fg="black")
                    except:
                        harmonicsentries[x + 10].config(fg="red")
                        valueerrorflag = 1
                        break
                else:
                    isquareentries.append(0)
            # if(valueerrorflag == 1): return
            # sum of I*I
            currentsqrresult = float(math.sqrt(sum(isquareentries)))
            # Ilin Calculation
            # base_dash = round(int(basenentry.get()) * (50/page1_frequency),2)
            base_dash = int(basenentry.get()) * (50 / page1_frequency)
            # print("Base_dash=",base_dash)
            for x in range(0, 10):
                if (str(harmonicsentries[x].get()) != str(empty_entry.get(1.0, "end-1c")) or str(
                        harmonicsentries[x + 10].get()) != str(empty_entry.get(1.0, "end-1c"))):
                    try:
                        if (x == 0):
                            # isquareentries.append(
                            #    round(((int(harmonicsentries[x].get()) * int(harmonicsentries[x + 10].get())) / (
                            #        int(basenentry.get()))), 2))
                            isquareentries.append(
                                ((int(harmonicsentries[x].get()) * int(harmonicsentries[x + 10].get())) / (
                                    base_dash)))

                            harmonicsentries[x].config(fg="black")
                            harmonicsentries[x + 10].config(fg="black")
                        else:
                            # isquareentries.append(
                            #    round(((int(harmonicsentries[x].get()) * int(harmonicsentries[x + 10].get()) * float(
                            #        ambfactorentry.get())) / (int(basenentry.get()))), 2))
                            isquareentries.append(
                                ((int(harmonicsentries[x].get()) * int(harmonicsentries[x + 10].get()) * float(
                                    ambfactorentry.get())) / (base_dash)))

                            harmonicsentries[x].config(fg="black")
                            harmonicsentries[x + 10].config(fg="black")
                    except ValueError:
                        harmonicsentries[x].config(fg="red")
                        harmonicsentries[x + 10].config(fg="red")
                        valueerrorflag = 1
                        break
                else:
                    isquareentries.append(0)
            if (valueerrorflag == 1): return
            # sum of ILin
            for x in range(10, 20):
                currentlinresult = currentlinresult + isquareentries[x]

            for x in range(0, len(isquareentries)):
                if (isquareentries[x] > 0):
                    genrated_data[x].config(borderwidth=2, state="normal")
                    genrated_data[x].delete(0, "end")
                    genrated_data[x].insert(0, str(isquareentries[x]))
                    genrated_data[x].config(borderwidth=2, state="disable")

            astrarating = max(round(currentsqrresult, 2), round(currentlinresult, 2))
            # Ambient Astra Rating Calculation
            if (int(ambtempentry.get()) < 40):
                ambastrarating = astrarating
            else:
                ambastrarating = astrarating / (1 - ((int(ambtempentry.get()) - 40) / 50))
            # Display of I*I Sum
            if (currentsqrresult != 0):
                currentsqrresult = round(currentsqrresult, 2)
                currentsqrentry.config(borderwidth=2, state="normal")
                currentsqrentry.delete(0, "end")
                currentsqrentry.insert(0, (str(currentsqrresult)))
                currentsqrentry.config(borderwidth=2, state="disable")

            # Display of Ilin Sum
            if (currentlinresult != 0):
                currentlinresult = round(currentlinresult, 2)
                currentlinentry.config(borderwidth=2, state="normal")
                currentlinentry.delete(0, "end")
                currentlinentry.insert(0, (str(currentlinresult)))
                currentlinentry.config(borderwidth=2, state="disable")

            # Display of Astra Rating

            if (astrarating != 0):
                astrarating = round(astrarating, 2)
                AHFsizeentry.config(borderwidth=2, state="normal")
                AHFsizeentry.delete(0, "end")
                # AHFsizeentry.insert(0, (str(astrarating)))
                AHFsizeentry.insert(0, ((str(astrarating).lstrip('0')) + ' A'))
                AHFsizeentry.config(borderwidth=2, state="disable")

            # Display of Ambient Astra Rating
            if (ambastrarating != 0):
                ambastrarating = round(ambastrarating, 2)
                AHFsize1entry.config(borderwidth=2, state="normal")
                AHFsize1entry.delete(0, "end")
                # AHFsize1entry.insert(0, ((str(ambastrarating)).zfill(5)))
                AHFsize1entry.insert(0, ((str(ambastrarating).lstrip('0')) + ' A'))
                AHFsize1entry.config(borderwidth=2, state="disable")
        else:
            if (fault < 1000):
                status_p1_entry.config(borderwidth=2, state="normal")
                status_p1_entry.delete(0, "end")
                status_p1_entry.insert(0, ("Insufficient inputs"))
                status_p1_entry.config(borderwidth=2, state="disable")
            else:
                status_p1_entry.config(borderwidth=2, state="normal")
                status_p1_entry.delete(0, "end")
                status_p1_entry.insert(0, ("contact EPIC"))
                status_p1_entry.config(borderwidth=2, state="disable")
    elif (astranotebook.index(astranotebook.select()) == 1):
        print("in process")
        global p2_isquareentries, p2_ioentries, p2_currentlinentries, p2_astrarating, p2_ambastrarating
        global p2_currentsqrresult, p2_currentlinresult, p2_idcentries, p2_idcresult
        global p2_harmonics_entry_list, p2_current_entry_list, p2_current_entry_list_n
        global p2_number_of_entries

        p2_isquareentries = []
        p2_currentlinentries = []
        p2_harmonics_entry_list = []
        p2_current_entry_list = []
        p2_current_entry_list_n = []
        p2_ioentries = []
        p2_idcentries = []
        p2_idcresult = 0
        p2_currentsqrresult = 0
        p2_currentlinresult = 0
        p2_number_of_entries = 0
        # I*I Calculation
        # print("fault = ", fault)
        # p2_fault = 0 #must be removed
        if (p2_fault == 0):
            p2_valueerrorflag = 0
            clear_results()
            # -->disable_modified_flag()
            p2_base_dash = int(p2_basenentry.get()) * (50 / page2_frequency)
            p2_base_zero = 4.5
            for x in range(0, 10):
                print("loop", x)
                if (str(p2_harmonicsentries[x].get()) != str(p2_empty_entry.get(1.0, "end-1c")) or str(
                        p2_harmonicsentries[x + 10].get()) != str(p2_empty_entry.get(1.0, "end-1c")) or str(
                    p2_harmonicsentries[x + 20].get()) != str(p2_empty_entry.get(1.0, "end-1c"))):

                    try:
                        p2_harmonics_entry_list.append(int(p2_harmonicsentries[x].get()))
                        p2_current_entry_list.append(int(p2_harmonicsentries[x + 10].get()))
                        p2_current_entry_list_n.append(int(p2_harmonicsentries[x + 20].get()))
                        p2_number_of_entries = p2_number_of_entries + 1
                        p2_harmonicsentries[x].config(fg="black")
                        p2_harmonicsentries[x + 10].config(fg="black")
                        p2_harmonicsentries[x + 20].config(fg="black")
                    except ValueError:
                        p2_harmonicsentries[x].config(fg="red")
                        p2_harmonicsentries[x + 10].config(fg="red")
                        p2_harmonicsentries[x + 20].config(fg="red")
                        p2_valueerrorflag = 1
                    try:
                        p2_ioentries.append(int(p2_current_entry_list_n[x]) / 3)
                        p2_harmonicsentries[x + 20].config(fg="black")
                    except Exception as e:
                        p2_harmonicsentries[x + 20].config(fg="red")
                        p2_valueerrorflag = 1

                    if (p2_valueerrorflag == 1): return
                    try:
                        if (x == 0):
                            p2_isquareentries.append(
                                ((int(p2_harmonicsentries[x + 10].get()) * int(p2_harmonicsentries[x + 10].get()))))

                            p2_harmonicsentries[x + 10].config(fg="black")
                        else:
                            p2_isquareentries.append(
                                ((int(p2_harmonicsentries[x + 10].get()) * int(
                                    p2_harmonicsentries[x + 10].get())) * float(p2_ambfactorentry.get()) * float(
                                    p2_ambfactorentry.get())))

                            p2_harmonicsentries[x + 10].config(fg="black")
                    except ValueError:
                        print("%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%")
                        p2_harmonicsentries[x + 10].config(fg="red")
                        p2_valueerrorflag = 1
                        break

                    try:
                        if (x == 0):
                            p2_currentlinentries.append(
                                ((int(p2_harmonicsentries[x].get()) * int(
                                    p2_harmonicsentries[x + 10].get()))) / p2_base_dash)

                            p2_harmonicsentries[x].config(fg="black")
                            p2_harmonicsentries[x + 10].config(fg="black")
                        else:
                            p2_currentlinentries.append(
                                ((int(p2_harmonicsentries[x].get()) * int(p2_harmonicsentries[x + 10].get())) * float(
                                    p2_ambfactorentry.get())) / p2_base_dash)

                            p2_harmonicsentries[x].config(fg="black")
                            p2_harmonicsentries[x + 10].config(fg="black")
                    except ValueError:
                        p2_harmonicsentries[x].config(fg="red")
                        p2_harmonicsentries[x + 10].config(fg="red")
                        p2_valueerrorflag = 1
                        break
                    try:
                        if (int(p2_harmonicsentries[x].get() == 0)):
                            p2_idcentries.append(0)
                            break
                        if (x == 0):
                            p2_idcentries.append(
                                (p2_ioentries[x] * p2_base_zero) / int(p2_harmonicsentries[x].get()))

                            p2_harmonicsentries[x].config(fg="black")
                            p2_harmonicsentries[x + 20].config(fg="black")
                        else:
                            p2_idcentries.append(
                                (p2_ioentries[x] * p2_base_zero * float(p2_ambfactorentry.get())) / int(
                                    p2_harmonicsentries[x].get()))

                            p2_harmonicsentries[x].config(fg="black")
                            p2_harmonicsentries[x + 20].config(fg="black")
                    except ValueError:
                        p2_harmonicsentries[x].config(fg="red")
                        p2_harmonicsentries[x + 20].config(fg="red")
                        p2_valueerrorflag = 1
                        break
                else:
                    # else case need to written
                    p2_ioentries.append(0)
                    p2_isquareentries.append(0)
                    p2_currentlinentries.append(0)
                    p2_idcentries.append(0)

            # result of I*I
            p2_currentsqrresult = math.sqrt((sum(p2_isquareentries)))
            p2_currentlinresult = sum(p2_currentlinentries)
            p2_idcresult = sum(p2_idcentries)
            print("p2_currentsqrresult = ", p2_currentsqrresult)
            print("p2_currentlinresult = ", p2_currentlinresult)
            print("p2_idcresult = ", p2_idcresult)
            # calculation of Astra rating
            p2_astrarating = max(round(p2_currentsqrresult, 2), round(p2_currentlinresult, 2), round(p2_idcresult, 2))
            # Ambient Astra Rating Calculation
            if (int(p2_ambtempentry.get()) < 40):
                p2_ambastrarating = p2_astrarating
            else:
                p2_ambastrarating = p2_astrarating / (1 - ((int(p2_ambtempentry.get()) - 40) / 50))

            # Display of Astra Rating

            if (p2_astrarating != 0):
                p2_astrarating = round(p2_astrarating, 2)
                p2_AHFsizeentry.config(borderwidth=2, state="normal")
                p2_AHFsizeentry.delete(0, "end")
                # AHFsizeentry.insert(0, (str(astrarating)))
                p2_AHFsizeentry.insert(0, ((str(p2_astrarating).lstrip('0')) + ' A'))
                p2_AHFsizeentry.config(borderwidth=2, state="disable")

            # Display of Ambient Astra Rating
            if (p2_ambastrarating != 0):
                p2_ambastrarating = round(p2_ambastrarating, 2)
                p2_AHFsize1entry.config(borderwidth=2, state="normal")
                p2_AHFsize1entry.delete(0, "end")
                # AHFsize1entry.insert(0, ((str(ambastrarating)).zfill(5)))
                p2_AHFsize1entry.insert(0, ((str(p2_ambastrarating).lstrip('0')) + ' A'))
                p2_AHFsize1entry.config(borderwidth=2, state="disable")
        else:
            if (p2_fault < 1000):
                status_p2_entry.config(borderwidth=2, state="normal")
                status_p2_entry.delete(0, "end")
                status_p2_entry.insert(0, ("Insufficient inputs"))
                status_p2_entry.config(borderwidth=2, disabledforeground="red", state="disable")
            else:
                status_p2_entry.config(borderwidth=2, state="normal")
                status_p2_entry.delete(0, "end")
                status_p2_entry.insert(0, ("contact EPIC"))
                status_p2_entry.config(borderwidth=2, disabledforeground="red", state="disable")
    elif (astranotebook.index(astranotebook.select()) == 2):
        print("in process")
        global p3_isquareentries, p3_ioentries, p3_currentlinentries, p3_astrarating, p3_ambastrarating
        global p3_currentsqrresult, p3_currentlinresult, p3_idcentries, p3_idcresult
        global p3_harmonics_entry_list, p3_current_entry_list_n
        global p3_number_of_entries

        p3_isquareentries = []
        p3_currentlinentries = []
        p3_harmonics_entry_list = []
        p3_current_entry_list_n = []
        p3_ioentries = []
        p3_idcentries = []
        p3_idcresult = 0
        p3_currentsqrresult = 0
        p3_currentlinresult = 0
        p3_number_of_entries = 0
        # I*I Calculation
        if (p3_fault == 0):
            p3_valueerrorflag = 0
            clear_results()
            # -->disable_modified_flag()
            print(str(p3_basenentry.get()))
            p3_base_dash = int(p3_basenentry.get()) * (50 / page3_frequency)
            p3_base_zero = 4.5
            for x in range(0, 10):
                print("loop", x)
                if (str(p3_harmonicsentries[x].get()) != str(p3_empty_entry.get(1.0, "end-1c")) or str(
                        p3_harmonicsentries[x + 10].get()) != str(p3_empty_entry.get(1.0, "end-1c"))):

                    try:
                        p3_harmonics_entry_list.append(int(p3_harmonicsentries[x].get()))
                        p3_current_entry_list_n.append(int(p3_harmonicsentries[x + 10].get()))
                        p3_number_of_entries = p3_number_of_entries + 1
                    except Exception as e:
                        print(e)
                        p3_harmonicsentries[x].config(fg="red")
                        p3_harmonicsentries[x + 10].config(fg="red")
                        p3_valueerrorflag = 1
                    try:
                        p3_ioentries.append(int(p3_current_entry_list_n[x]) / 3)
                    except Exception as e:
                        print(e)
                        p3_harmonicsentries[x + 10].config(fg="red")
                        p3_valueerrorflag = 1

                    if (p3_valueerrorflag == 1):
                        print("Return")
                        return
                    try:
                        if (x == 0):
                            p3_isquareentries.append(
                                ((int(p3_ioentries[x]) * int(p3_ioentries[x]))))

                            p3_harmonicsentries[x].config(fg="black")
                            p3_harmonicsentries[x + 10].config(fg="black")
                        else:
                            p3_isquareentries.append(
                                ((int(p3_ioentries[x]) * int(p3_ioentries[x])) * float(p3_ambfactorentry.get()) * float(
                                    p3_ambfactorentry.get())))

                            p3_harmonicsentries[x].config(fg="black")
                            p3_harmonicsentries[x + 10].config(fg="black")
                    except ValueError:
                        p3_harmonicsentries[x].config(fg="red")
                        p3_harmonicsentries[x + 10].config(fg="red")
                        p3_valueerrorflag = 1
                        break

                    try:
                        if (x == 0):
                            p3_currentlinentries.append(
                                ((int(p3_harmonicsentries[x].get()) * int(p3_ioentries[x]))) / p3_base_dash)

                            p3_harmonicsentries[x].config(fg="black")
                            p3_harmonicsentries[x + 10].config(fg="black")
                        else:
                            p3_currentlinentries.append(
                                ((int(p3_harmonicsentries[x].get()) * int(p3_ioentries[x])) * float(
                                    p3_ambfactorentry.get())) / p3_base_dash)

                            p3_harmonicsentries[x].config(fg="black")
                            p3_harmonicsentries[x + 10].config(fg="black")
                    except ValueError:
                        p3_harmonicsentries[x].config(fg="red")
                        p3_harmonicsentries[x + 10].config(fg="red")
                        p3_valueerrorflag = 1
                        break
                    try:
                        if (int(p3_harmonicsentries[x].get() == 0)):
                            p3_idcentries.append(0)
                            break
                        if (x == 0):
                            p3_idcentries.append(
                                (p3_ioentries[x] * p3_base_zero) / int(p3_harmonicsentries[x].get()))

                            p3_harmonicsentries[x].config(fg="black")
                            p3_harmonicsentries[x + 10].config(fg="black")
                        else:
                            p3_idcentries.append(
                                (p3_ioentries[x] * p3_base_zero * float(p3_ambfactorentry.get())) / int(
                                    p3_harmonicsentries[x].get()))

                            p3_harmonicsentries[x].config(fg="black")
                            p3_harmonicsentries[x + 10].config(fg="black")
                    except ValueError:
                        p3_harmonicsentries[x].config(fg="red")
                        p3_harmonicsentries[x + 10].config(fg="red")
                        p3_valueerrorflag = 1
                        break
                else:
                    print("else")
                    # else case need to written
                    p3_ioentries.append(0)
                    p3_isquareentries.append(0)
                    p3_currentlinentries.append(0)
                    p3_idcentries.append(0)

            # result of I*I
            p3_currentsqrresult = math.sqrt((sum(p3_isquareentries)))
            p3_currentlinresult = sum(p3_currentlinentries)
            p3_idcresult = sum(p3_idcentries)
            print("p3_currentsqrresult = ", p3_currentsqrresult)
            print("p3_currentlinresult = ", p3_currentlinresult)
            print("p3_idcresult = ", p3_idcresult)
            # calculation of Astra rating
            p3_astrarating = max(round(p3_currentsqrresult, 2), round(p3_currentlinresult, 2), round(p3_idcresult, 2))
            # Ambient Astra Rating Calculation
            if (int(p3_ambtempentry.get()) < 40):
                p3_ambastrarating = p3_astrarating
            else:
                p3_ambastrarating = p3_astrarating / (1 - ((int(p3_ambtempentry.get()) - 40) / 50))

            # Display of Astra Rating

            if (p3_astrarating != 0):
                p3_astrarating = round(p3_astrarating, 2)
                p3_AHFsizeentry.config(borderwidth=2, state="normal")
                p3_AHFsizeentry.delete(0, "end")
                # AHFsizeentry.insert(0, (str(astrarating)))
                p3_AHFsizeentry.insert(0, ((str(p3_astrarating).lstrip('0')) + ' A'))
                p3_AHFsizeentry.config(borderwidth=2, state="disable")

            # Display of Ambient Astra Rating
            if (p3_ambastrarating != 0):
                p3_ambastrarating = round(p3_ambastrarating, 2)
                p3_AHFsize1entry.config(borderwidth=2, state="normal")
                p3_AHFsize1entry.delete(0, "end")
                # AHFsize1entry.insert(0, ((str(ambastrarating)).zfill(5)))
                p3_AHFsize1entry.insert(0, ((str(p3_ambastrarating).lstrip('0')) + ' A'))
                p3_AHFsize1entry.config(borderwidth=2, state="disable")
        else:
            if (p3_fault < 1000):
                status_p3_entry.config(borderwidth=2, state="normal")
                status_p3_entry.delete(0, "end")
                status_p3_entry.insert(0, ("Insufficient inputs"))
                status_p3_entry.config(borderwidth=2, disabledforeground="red", state="disable")
            else:
                status_p3_entry.config(borderwidth=2, state="normal")
                status_p3_entry.delete(0, "end")
                status_p3_entry.insert(0, ("contact EPIC"))
                status_p3_entry.config(borderwidth=2, disabledforeground="red", state="disable")
#----------------Astra process button code ended--------------#

#--------------------Function High is started-----------------#
def high():
    global save_flag
    if (astranotebook.index(astranotebook.select()) == 0):
        # page - 1 notch profile high
        if (save_flag == 0): enable_modified_flag()
        print("save_flag", save_flag)
        highnotchbtn.config(bg="#05c1fa", fg="white", font=font.Font(family='Calibri', size=10, weight=font.BOLD))
        mildnotchbtn.config(bg="white", fg="black", font=font.Font(family='Calibri', size=10, weight=font.NORMAL))
        lownotchbtn.config(bg="white", fg="black", font=font.Font(family='Calibri', size=10, weight=font.NORMAL))
        basenentry.delete(0, "end")
        basenentry.insert(0, ("3"))
#-----------------Function high is ended-------------------------#

#-----------------Function mild is ended-------------------------#
def mild():
    global save_flag
    if (astranotebook.index(astranotebook.select()) == 0):
        # page - 1 notch profile medium
        if (save_flag == 0): enable_modified_flag()
        print("save_flag", save_flag)
        highnotchbtn.config(bg="white", fg="black", font=font.Font(family='Calibri', size=10, weight=font.NORMAL))
        mildnotchbtn.config(bg="#05c1fa", fg="white", font=font.Font(family='Calibri', size=10, weight=font.BOLD))
        lownotchbtn.config(bg="white", fg="black", font=font.Font(family='Calibri', size=10, weight=font.NORMAL))
        basenentry.delete(0, "end")
        basenentry.insert(0, ("5"))
    elif (astranotebook.index(astranotebook.select()) == 1):
        # page - 2 notch profile medium
        if (save_flag == 0): enable_modified_flag()
        print("save_flag", save_flag)
        p2_lownotchbtn.config(bg="white", fg="black", font=font.Font(family='Calibri', size=10, weight=font.NORMAL))
        p2_mildnotchbtn.config(bg="#05c1fa", fg="white", font=font.Font(family='Calibri', size=10, weight=font.BOLD))
        p2_basenentry.delete(0, "end")
        p2_basenentry.insert(0, ("3"))
    elif (astranotebook.index(astranotebook.select()) == 2):
        # page - 2 notch profile medium
        if (save_flag == 0): enable_modified_flag()
        print("save_flag", save_flag)
        p3_lownotchbtn.config(bg="white", fg="black", font=font.Font(family='Calibri', size=10, weight=font.NORMAL))
        p3_mildnotchbtn.config(bg="#05c1fa", fg="white", font=font.Font(family='Calibri', size=10, weight=font.BOLD))
        p3_basenentry.delete(0, "end")
        p3_basenentry.insert(0, ("3"))
    elif (astranotebook.index(astranotebook.select()) == 3):
        # page - 2 notch profile medium
        if (save_flag == 0): enable_modified_flag()
        print("save_flag", save_flag)
        p3_lownotchbtn.config(bg="white", fg="black", font=font.Font(family='Calibri', size=10, weight=font.NORMAL))
        p3_mildnotchbtn.config(bg="#05c1fa", fg="white", font=font.Font(family='Calibri', size=10, weight=font.BOLD))
        p3_basenentry.delete(0, "end")
        p3_basenentry.insert(0, ("3"))

#---------------------Function mild is ended---------------------#

#------------------------Function Nonotch is started---------------------#
def nonotch():
    # page - 1 notch profile No notch
    global save_flag
    if (astranotebook.index(astranotebook.select()) == 0):

        if (save_flag == 0): enable_modified_flag()
        print("save_flag", save_flag)
        highnotchbtn.config(bg="white", fg="black", font=font.Font(family='Calibri', size=10, weight=font.NORMAL))
        mildnotchbtn.config(bg="white", fg="black", font=font.Font(family='Calibri', size=10, weight=font.NORMAL))
        lownotchbtn.config(bg="#05c1fa", fg="white", font=font.Font(family='Calibri', size=10, weight=font.BOLD))
        basenentry.delete(0, "end")
        basenentry.insert(0, ("10"))
    elif (astranotebook.index(astranotebook.select()) == 1):
        # page - 2 notch profile high
        if (save_flag == 0): enable_modified_flag()
        print("save_flag", save_flag)
        p2_lownotchbtn.config(bg="#05c1fa", fg="white", font=font.Font(family='Calibri', size=10, weight=font.BOLD))
        p2_mildnotchbtn.config(bg="white", fg="black", font=font.Font(family='Calibri', size=10, weight=font.NORMAL))
        p2_basenentry.delete(0, "end")
        p2_basenentry.insert(0, ("6"))
    elif (astranotebook.index(astranotebook.select()) == 2):
        # page - 2 notch profile high
        if (save_flag == 0): enable_modified_flag()
        print("save_flag", save_flag)
        p3_lownotchbtn.config(bg="#05c1fa", fg="white", font=font.Font(family='Calibri', size=10, weight=font.BOLD))
        p3_mildnotchbtn.config(bg="white", fg="black", font=font.Font(family='Calibri', size=10, weight=font.NORMAL))
        p3_basenentry.delete(0, "end")
        p3_basenentry.insert(0, ("6"))
#------------------------Function Nonotch is ended--------------------------#

#-------------------------Function Fhz is strated---------------------------#
def fhz():
    global save_flag
    if (astranotebook.index(astranotebook.select()) == 0):
        # page - 1 50 Hz frequency selection
        if (save_flag == 0): enable_modified_flag()
        print("save_flag", save_flag)
        global page1_frequency
        fhzbtn.config(bg="#05c1fa", fg="white", font=font.Font(family='Calibri', size=10, weight=font.BOLD))
        shzbtn.config(bg="white", fg="black", font=font.Font(family='Calibri', size=10, weight=font.NORMAL))
        page1_frequency = 50
    elif (astranotebook.index(astranotebook.select()) == 1):
        # page - 2 50 Hz frequency selection
        if (save_flag == 0): enable_modified_flag()
        print("save_flag", save_flag)
        global page2_frequency
        p2_fhzbtn.config(bg="#05c1fa", fg="white", font=font.Font(family='Calibri', size=10, weight=font.BOLD))
        p2_shzbtn.config(bg="white", fg="black", font=font.Font(family='Calibri', size=10, weight=font.NORMAL))
        page2_frequency = 50
    elif (astranotebook.index(astranotebook.select()) == 2):
        # page - 2 50 Hz frequency selection
        if (save_flag == 0): enable_modified_flag()
        print("save_flag", save_flag)
        global page3_frequency
        p3_fhzbtn.config(bg="#05c1fa", fg="white", font=font.Font(family='Calibri', size=10, weight=font.BOLD))
        p3_shzbtn.config(bg="white", fg="black", font=font.Font(family='Calibri', size=10, weight=font.NORMAL))
        page3_frequency = 50
#-------------------------Function Fhz is ended----------------#

#------------------------Function Shz is started------------------#
def shz():
    global save_flag
    if (astranotebook.index(astranotebook.select()) == 0):
        # page - 1 60 Hz frequency selection
        if (save_flag == 0): enable_modified_flag()
        print("save_flag", save_flag)
        global page1_frequency
        fhzbtn.config(bg="white", fg="black", font=font.Font(family='Calibri', size=10, weight=font.NORMAL))
        shzbtn.config(bg="#05c1fa", fg="white", font=font.Font(family='Calibri', size=10, weight=font.BOLD))
        page1_frequency = 60
    if (astranotebook.index(astranotebook.select()) == 1):
        # page - 2 60 Hz frequency selection
        if (save_flag == 0): enable_modified_flag()
        print("save_flag", save_flag)
        global page2_frequency
        p2_fhzbtn.config(bg="white", fg="black", font=font.Font(family='Calibri', size=10, weight=font.NORMAL))
        p2_shzbtn.config(bg="#05c1fa", fg="white", font=font.Font(family='Calibri', size=10, weight=font.BOLD))
        page2_frequency = 60
    if (astranotebook.index(astranotebook.select()) == 2):
        # page - 2 60 Hz frequency selection
        if (save_flag == 0): enable_modified_flag()
        print("save_flag", save_flag)
        global page3_frequency
        p3_fhzbtn.config(bg="white", fg="black", font=font.Font(family='Calibri', size=10, weight=font.NORMAL))
        p3_shzbtn.config(bg="#05c1fa", fg="white", font=font.Font(family='Calibri', size=10, weight=font.BOLD))
        page3_frequency = 60
#------------------------Function Shz is ended------------------#

#------------------To clear the comment box----------------#
def clear_comment_box_message():
    if (astranotebook.index(astranotebook.select()) == 0):
        # clearing the page - 1 comment box
        comment_box_message.delete(1.0, "end")
        comment_box_message.focus_set()
    if (astranotebook.index(astranotebook.select()) == 1):
        # clearing the page - 2 comment box
        p2_comment_box_message.delete(1.0, "end")
        p2_comment_box_message.focus_set()
    if (astranotebook.index(astranotebook.select()) == 2):
        # clearing the page - 2 comment box
        p3_comment_box_message.delete(1.0, "end")
        p3_comment_box_message.focus_set()
#-----------------------Ended the clear comment box---------------#

#--------------------Bind the FocusIn and FocusOut events to simulate a placeholder----------------#
def on_focus_in(event):
    if (astranotebook.index(astranotebook.select()) == 0):
        # page -1 comment box focus event
        # enable_modified_flag()
        if comment_box_message.get('1.0', 'end-1c') == placeholder_text:
            comment_box_message.delete('1.0', 'end-1c')
            comment_box_message.config(fg='black')
            enable_modified_flag()
    if (astranotebook.index(astranotebook.select()) == 1):
        # page -2 comment box focus event
        # enable_modified_flag()
        if p2_comment_box_message.get('1.0', 'end-1c') == placeholder_text:
            p2_comment_box_message.delete('1.0', 'end-1c')
            p2_comment_box_message.config(fg='black')
            enable_modified_flag()
    if (astranotebook.index(astranotebook.select()) == 2):
        # page -3 comment box focus event
        # enable_modified_flag()
        if p3_comment_box_message.get('1.0', 'end-1c') == placeholder_text:
            p3_comment_box_message.delete('1.0', 'end-1c')
            p3_comment_box_message.config(fg='black')
            enable_modified_flag()


def on_focus_out(event):
    if (astranotebook.index(astranotebook.select()) == 0):
        # page -1 comment box focus event
        # enable_modified_flag()
        if not comment_box_message.get('1.0', 'end-1c'):
            comment_box_message.insert('1.0', placeholder_text)
            comment_box_message.config(fg='gray')
            enable_modified_flag()
    if (astranotebook.index(astranotebook.select()) == 1):
        # page -2 comment box focus event
        # enable_modified_flag()
        if not p2_comment_box_message.get('1.0', 'end-1c'):
            p2_comment_box_message.insert('1.0', placeholder_text)
            p2_comment_box_message.config(fg='gray')
            enable_modified_flag()
    if (astranotebook.index(astranotebook.select()) == 2):
        # page -3 comment box focus event
        # enable_modified_flag()
        if not p3_comment_box_message.get('1.0', 'end-1c'):
            p3_comment_box_message.insert('1.0', placeholder_text)
            p3_comment_box_message.config(fg='gray')
            enable_modified_flag()

#--------------------Bind the FocusIn and FocusOut events to simulate a placeholder is ended----------------#

#-----------------Fetch the command box------------------#
def fetch_comment_box_message():
    if (astranotebook.index(astranotebook.select()) == 0):
        # page -1 comment box fetch the mage event
        global page1_comment_box_message
        page1_comment_box_message = comment_box_message.get(1.0, "end")
        # print(page1_comment_box_message)
    if (astranotebook.index(astranotebook.select()) == 1):
        # page -1 comment box fetch the mage event
        global page2_comment_box_message
        page2_comment_box_message = p2_comment_box_message.get(1.0, "end")
        # print(page1_comment_box_message)
    if (astranotebook.index(astranotebook.select()) == 2):
        # page -1 comment box fetch the mage event
        global page3_comment_box_message
        page3_comment_box_message = p3_comment_box_message.get(1.0, "end")
        # print(page1_comment_box_message)
#--------------------Fetch the command box is ended-------------------#


# Creating Notebook
astranotebook = ttk.Notebook(root)
astranotebook.pack(fill='both', expand=1)

# creating frame for all the pages
# Set background color and gradient effect
# c5d5db , A0D2EB , E6E6E6 , F6F6F8 - OK ,
background_color = "#F6F6F8"
gradient_color = "#C9C9C9"  # #C9C9C9 - ok
# Page - 1 3 phase 3 Wire
astrap1_frame = Canvas(root, bg=background_color, highlightthickness=0)
astrap1_frame.pack(fill='both', expand=1)

height = astrap1_frame.winfo_height()
for i in range(height):
    astrap1_frame.create_line(0, i, 5000, i, fill=gradient_color)

# Page - 2 3 phase 4 Wire
astrap2_frame = Canvas(root, bg=background_color, highlightthickness=0)
astrap2_frame.pack(fill='both', expand=1)

height = astrap2_frame.winfo_height()
for i in range(height):
    astrap2_frame.create_line(0, i, 5000, i, fill=gradient_color)

# Page - 3 3 phase Neutral compensation
astrap3_frame = Canvas(root, bg=background_color, highlightthickness=0)
astrap3_frame.pack(fill='both', expand=1)

height = astrap3_frame.winfo_height()
for i in range(height):
    astrap3_frame.create_line(0, i, 5000, i, fill=gradient_color)

# Page - 4 True power factorn
astrap4_frame = Canvas(root, bg=background_color, highlightthickness=0)
astrap4_frame.pack(fill='both', expand=1)

height = astrap3_frame.winfo_height()
for i in range(height):
    astrap3_frame.create_line(0, i, 5000, i, fill=gradient_color)

# adding the pages to notebook
astranotebook.add(astrap1_frame, text="Astra 3P,3W")
astranotebook.add(astrap2_frame, text="Astra 3P,4W")
astranotebook.add(astrap3_frame, text="Astra 3P, N")
astranotebook.add(astrap4_frame, text="True Power Factor ")

# Create a Frame widget and place it in the bottom-right corner of the astrap1_frame window
logoframe1 = Frame(astrap1_frame, bg="#F6F6F8")
logoframe1.place(relx=1.0, rely=1.0, anchor="se")

# Create a text widget with some text
astramakep1 = Text(logoframe1, height=1, width=20, highlightthickness=0, relief="flat")
astramakep1.insert("end", "  Made with ❤ in India")
astramakep1.configure(font=('Verdana', 8), bg="#F6F6F8", fg="gray")
astramakep1.grid(row=0, column=0)

# Get the position of the 'W' character
posp1 = astramakep1.search("❤", "1.0")

# Add a tag to the 'W' character
astramakep1.tag_add("red", posp1, f"{posp1}+1c")

# Configure the tag to use a different color
astramakep1.tag_config("red", foreground="red", background="#F6F6F8", font=('Verdana', 8))

# Disable the text widget so it's read-only and non-editable
astramakep1.configure(state="disabled")

# Copy Right Lable - page - 1
astracoprightp1 = Label(logoframe1, text="© 2024, InPhase | All Rights Reserved")
astracoprightp1.configure(font=('Verdana', 8), bg="#F6F6F8")
astracoprightp1.grid(row=1, column=0)

# Create a Frame widget and place it in the bottom-right corner of the root window
logoframe2 = Frame(astrap2_frame, bg="#F6F6F8")
logoframe2.place(relx=1.0, rely=1.0, anchor="se")

# Create a text widget with some text
astramakep2 = Text(logoframe2, height=1, width=20, highlightthickness=0, relief="flat")
astramakep2.insert("end", "  Made with ❤ in India")
astramakep2.configure(font=('Verdana', 8), bg="#F6F6F8", fg="gray")
astramakep2.grid(row=0, column=0)

# Get the position of the 'W' character
posp2 = astramakep2.search("❤", "1.0")

# Add a tag to the 'W' character
astramakep2.tag_add("red", posp2, f"{posp2}+1c")

# Configure the tag to use a different color
astramakep2.tag_config("red", foreground="red", background="#F6F6F8", font=('Verdana', 8))

# Disable the text widget so it's read-only and non-editable
astramakep2.configure(state="disabled")

# Copy Right Lable - page - 2
astracoprightp2 = Label(logoframe2, text="© 2024, InPhase | All Rights Reserved")
astracoprightp2.configure(font=('Verdana', 8), bg="#F6F6F8")
astracoprightp2.grid(row=1, column=0)

# Create a Frame widget and place it in the bottom-right corner of the root window
logoframe3 = Frame(astrap3_frame, bg="#F6F6F8")
logoframe3.place(relx=1.0, rely=1.0, anchor="se")

# Create a text widget with some text
astramakep3 = Text(logoframe3, height=1, width=20, highlightthickness=0, relief="flat")
astramakep3.insert("end", "  Made with ❤ in India")
astramakep3.configure(font=('Verdana', 8), bg="#F6F6F8", fg="gray")
astramakep3.grid(row=0, column=0)

# Get the position of the 'W' character
posp3 = astramakep3.search("❤", "1.0")

# Add a tag to the 'W' character
astramakep3.tag_add("red", posp3, f"{posp3}+1c")

# Configure the tag to use a different color
astramakep3.tag_config("red", foreground="red", background="#F6F6F8", font=('Verdana', 8))

# Disable the text widget so it's read-only and non-editable
astramakep3.configure(state="disabled")

# Copy Right Lable - page - 3
astracoprightp3 = Label(logoframe3, text="© 2024, InPhase | All Rights Reserved")
astracoprightp3.configure(font=('Verdana', 8), bg="#F6F6F8")
astracoprightp3.grid(row=1, column=0)

# Create a Frame widget and place it in the bottom-right corner of the root window
logoframe4 = Frame(root, bg="#F6F6F8")
logoframe4.place(relx=1.0, rely=1.0, anchor="se")

''''# Create a text widget with some text
astramakep4 = Text(logoframe4, height=1, width=20, highlightthickness=0, relief="flat")
astramakep4.insert("end", "  Made with ❤ in India")
astramakep4.configure(font=('Verdana', 8), bg="#F6F6F8", fg="gray")
astramakep4.grid(row=0, column=0)

# Get the position of the 'W' character
posp4 = astramakep3.search("❤", "1.0")

# Add a tag to the 'W' character
astramakep4.tag_add("red", posp3, f"{posp3}+1c")

# Configure the tag to use a different color
astramakep4.tag_config("red", foreground="red", background="#F6F6F8", font=('Verdana', 8))

# Disable the text widget so it's read-only and non-editable
astramakep4.configure(state="disabled")

# Copy Right Lable - page - 4
astracoprightp4 = Label(logoframe4, text="© 2024, InPhase | All Rights Reserved")
astracoprightp4.configure(font=('Verdana', 8), bg="#F6F6F8")
astracoprightp4.grid(row=1, column=0)'''

# Astra Lable Text
astralable = Label(astrap1_frame, text="       Astra Rating Calculator - 3P3W")
astralable.configure(font=('Verdana', 16), bg="#F6F6F8")

# Harmonics Cloumn
harmonicslable = Label(astrap1_frame, text="Harmonics")
harmonicslable.configure(font=('Verdana', 14), bg="#F6F6F8")

# Curent Cloumn
currentlable = Label(astrap1_frame, text="Current")
currentlable.configure(font=('Verdana', 14), bg="#F6F6F8")

# I*I Cloumn
currentsqrlable = Label(astrap1_frame, text="I*I")
currentsqrlable.configure(font=('Verdana', 14), bg="#F6F6F8")

# I*I total Entry
currentsqrentry = Entry(astrap1_frame, width=15, font=('Verdana 12'), justify='center', relief=FLAT)
# currentsqrentry = Entry(root, width=15, font=('Verdana 12'), justify='center')
currentsqrentry.config(borderwidth=1, state=DISABLED, disabledbackground="#F6F6F8", disabledforeground="black")
# currentsqrentry.config(borderwidth=4, disabledbackground="white", disabledforeground="black")

# Ilin Cloumn
currentlinlable = Label(astrap1_frame, text="Ilin")
currentlinlable.configure(font=('Verdana', 14), bg="#F6F6F8")

# Ilin total Entry
currentlinentry = Entry(astrap1_frame, width=15, font=('Verdana 12'), justify='center', relief=FLAT)
currentlinentry.config(borderwidth=1, state=DISABLED, disabledbackground="#F6F6F8", disabledforeground="black")

# Status Entry
status_p1_entry = Entry(astrap1_frame, width=20, font=('Verdana 12'), justify='center', relief=FLAT)
status_p1_entry.config(borderwidth=1, state=DISABLED, disabledbackground="#F6F6F8", disabledforeground="red")
status_p1_entry.place(x=455, y=415)
# calculate  button
# bg="SystemButtonFace"
calculate = Button(astrap1_frame, text="Process", command=process)
calculate.config(height=2, width=10, bg="white", fg="black",
                 font=font.Font(family='Calibri', size=9, weight=font.NORMAL))
calculate.place(x=520, y=445)

# Frequency button
# 50 Hz Button
# bg="SystemButtonFace"
fhzbtn = Button(astrap1_frame, text="50 Hz", command=fhz)
fhzbtn.config(height=1, width=10, bg="white", fg="black",
              font=font.Font(family='Calibri', size=10, weight=font.NORMAL))

# 60 Hz Button
# bg="SystemButtonFace"
shzbtn = Button(astrap1_frame, text="60 Hz", command=shz)
shzbtn.config(height=1, width=10, bg="white", fg="black",
              font=font.Font(family='Calibri', size=10, weight=font.NORMAL))

# Notch button
# High notch Button
highnotchbtn = Button(astrap1_frame, text="HIGH", command=high)
highnotchbtn.config(height=1, width=10, bg="white", fg="black",
                    font=font.Font(family='Calibri', size=10, weight=font.NORMAL))

# Mild notch Button
mildnotchbtn = Button(astrap1_frame, text="MEDIUM", command=mild)
mildnotchbtn.config(height=1, width=10, bg="white", fg="black",
                    font=font.Font(family='Calibri', size=10, weight=font.NORMAL))

# Low notch Button
lownotchbtn = Button(astrap1_frame, text="LOW", command=nonotch)
lownotchbtn.config(height=1, width=10, bg="white", fg="black",
                   font=font.Font(family='Calibri', size=10, weight=font.NORMAL))

# Login Button
# loginbtn = Button(astrap1_frame, text="login", command=validate_credentials)
# loginbtn.config(height=2, width=10, bg="white", fg="black",
#                font=font.Font(family='Calibri', size=9, weight=font.NORMAL))

# AHf Size
AHFsizelable = Label(astrap1_frame, text="AHF size (A)")
AHFsizelable.configure(font=('Verdana', 14), bg="#F6F6F8")

# AHF size (A) Entry
AHFsizeentry = Entry(astrap1_frame, width=15, font=('Verdana 12'), justify='center', relief=FLAT)
AHFsizeentry.config(borderwidth=1, state=DISABLED, disabledbackground="#F6F6F8", disabledforeground="black")

# AHF size (A, @Ta)	lable
AHFsize1lable = Label(astrap1_frame, text="AHF size (A, @Ta)")
AHFsize1lable.configure(font=('Verdana', 14), bg="#F6F6F8")

# AHF size (A, @Ta)	Entry
AHFsize1entry = Entry(astrap1_frame, width=15, font=('Verdana 12'), justify='center', relief=FLAT)
AHFsize1entry.config(borderwidth=1, state=DISABLED, disabledbackground="#F6F6F8", disabledforeground="black")

# Frequency
frequencylable = Label(astrap1_frame, text="Frequency")
frequencylable.configure(font=('Verdana', 14), bg="#F6F6F8")

# Base n
basenlable = Label(astrap1_frame, text="Notch Profile")
basenlable.configure(font=('Verdana', 14), bg="#F6F6F8")

# Base n Entry
basenentry = Entry(astrap1_frame, width=20, font=('Verdana 12'), justify='center')
basenentry.config(borderwidth=2)

# Ambient temperature
ambtemplable = Label(astrap1_frame, text="Ambient temperature")
ambtemplable.configure(font=('Verdana', 14), bg="#F6F6F8")

# Ambient temperature Entry
ambtempentry = Entry(astrap1_frame, width=20, font=('Verdana 12'), justify='center')
ambtempentry.config(borderwidth=2)

# Amplification Factor
ambfactorlable = Label(astrap1_frame, text="Amplification Factor")
ambfactorlable.configure(font=('Verdana', 14), bg="#F6F6F8")

# Amplification Factor Entry
ambfactorentry = Entry(astrap1_frame, width=20, font=('Verdana 12'), justify='center')
ambfactorentry.config(borderwidth=2)

# Password Entry
# passwordentry = Entry(astrap1_frame, width=20, font=('Verdana 12'), justify='center', show='•')
# passwordentry.config(borderwidth=2)

# empty position
empty_entry = Text(astrap1_frame, height=1, width=1)
empty_entry.config(borderwidth=0, bg="#F6F6F8", state="disabled")

# clear comment box
# bg="SystemButtonFace"
clear_comment_box = Button(astrap1_frame, text="Clear", command=clear_comment_box_message)
clear_comment_box.config(height=1, width=10, bg="white", fg="black",
                         font=font.Font(family='Calibri', size=9, weight=font.NORMAL))
clear_comment_box.place(x=520, y=580)
# x=573
# comment box
comment_box_message = Text(astrap1_frame, height=5, width=48)
comment_box_message.config(borderwidth=2)
comment_box_message.place(x=380, y=490)
# Set the placeholder text
placeholder_text = 'Comments here...'
comment_box_message.insert('1.0', placeholder_text)
comment_box_message.config(fg='gray')

comment_box_message.bind('<FocusIn>', on_focus_in)
comment_box_message.bind('<FocusOut>', on_focus_out)

# Fundamental current lable
FIlable = Label(astrap1_frame, text="Reactive current (IQ)")
FIlable.configure(font=('Verdana', 14), bg="#F6F6F8")
FIlable.place(x=815, y=100)

# ----
# V lable
vollable = Label(astrap1_frame, text="V")
vollable.configure(font=('Verdana', 14), bg="#F6F6F8")
vollable.place(x=810, y=170)

# V Entry
volentry = Entry(astrap1_frame, width=10, font=('Verdana 12'), justify='center')
volentry.config(borderwidth=2)
volentry.place(x=885, y=170)

# I lable
curlable = Label(astrap1_frame, text="I")
curlable.configure(font=('Verdana', 14), bg="#F6F6F8")
curlable.place(x=810, y=210)

# I Entry
curentry = Entry(astrap1_frame, width=10, font=('Verdana 12'), justify='center')
curentry.config(borderwidth=2)
curentry.place(x=885, y=210)

# -----
# kw lable
kwlable = Label(astrap1_frame, text="kW")
kwlable.configure(font=('Verdana', 14), bg="#F6F6F8")
kwlable.place(x=810, y=250)

# kw Entry
kwentry = Entry(astrap1_frame, width=10, font=('Verdana 12'), justify='center')
kwentry.config(borderwidth=2)
kwentry.place(x=885, y=250)

# kVA lable
kvalable = Label(astrap1_frame, text="kVA")
kvalable.configure(font=('Verdana', 14), bg="#F6F6F8")
kvalable.place(x=810, y=290)

# kva Entry
kvaentry = Entry(astrap1_frame, width=10, font=('Verdana 12'), justify='center')
kvaentry.config(borderwidth=2)
kvaentry.place(x=885, y=290)

# kVAR lable
kvarlable = Label(astrap1_frame, text="kVAR")
kvarlable.configure(font=('Verdana', 14), bg="#F6F6F8")
kvarlable.place(x=810, y=330)

# kVAR Entry
kvarentry = Entry(astrap1_frame, width=10, font=('Verdana 12'), justify='center')
kvarentry.config(borderwidth=2)
kvarentry.place(x=885, y=330)

# ---
# IPF lable
IPFlable = Label(astrap1_frame, text="IPF")
IPFlable.configure(font=('Verdana', 14), bg="#F6F6F8")
IPFlable.place(x=810, y=370)

# IPF Entry
IPFentry = Entry(astrap1_frame, width=10, font=('Verdana 12'), justify='center')
IPFentry.config(borderwidth=2)
IPFentry.place(x=885, y=370)

# TPF lable
TPFlable = Label(astrap1_frame, text="TPF")
TPFlable.configure(font=('Verdana', 14), bg="#F6F6F8")
TPFlable.place(x=810, y=410)

# TPF Entry
TPFentry = Entry(astrap1_frame, width=10, font=('Verdana 12'), justify='center')
TPFentry.config(borderwidth=2)
TPFentry.place(x=885, y=410)

# ---

# IQ lable
IQlable = Label(astrap1_frame, text="IQ")
IQlable.configure(font=('Verdana', 14), bg="#F6F6F8")
IQlable.place(x=810, y=450)

# IQ Entry
IQentry = Entry(astrap1_frame, width=10, font=('Verdana 12'), justify='center')
IQentry.config(borderwidth=2)
IQentry.place(x=885, y=450)

# ---

#####
# Formula sections for page -1

formula_title = [
    "select",
    "V, I, IPF, TPF",
    "V, kW, IPF, TPF",
    "V, kW, kVA, TPF",
    "V, kVA, IPF, TPF",
    "V, kVAR",
    "IQ",
]

#--------------clear junk entries------------------#
def clear_junk_entires():
    global import_progress_flag
    if (astranotebook.index(astranotebook.select()) == 0):
        # clearing the page - 1 junk data's which will be trigered at the time of selection of page - 1 formula box

        if (import_progress_flag == 0):
            harmonicsentries[10].config(borderwidth=2, state="normal")
            harmonicsentries[10].delete(0, "end")
            harmonicsentries[10].config(borderwidth=2, state="disable")

        volentry.config(borderwidth=2, state="normal")
        curentry.config(borderwidth=2, state="normal")
        kwentry.config(borderwidth=2, state="normal")
        kvaentry.config(borderwidth=2, state="normal")
        kvarentry.config(borderwidth=2, state="normal")
        IPFentry.config(borderwidth=2, state="normal")
        TPFentry.config(borderwidth=2, state="normal")
        IQentry.config(borderwidth=2, state="normal")
        harmonicsentries[10].config(fg="black", disabledbackground="#F6F6F8")
        vollable.config(fg="black")
        curlable.config(fg="black")
        kwlable.config(fg="black")
        kvalable.config(fg="black")
        kvarlable.config(fg="black")
        IPFlable.config(fg="black")
        TPFlable.config(fg="black")
        IQlable.config(fg="black")

        volentry.delete(0, "end")
        curentry.delete(0, "end")
        kwentry.delete(0, "end")
        kvaentry.delete(0, "end")
        kvarentry.delete(0, "end")
        IPFentry.delete(0, "end")
        TPFentry.delete(0, "end")
        IQentry.delete(0, "end")
    elif (astranotebook.index(astranotebook.select()) == 1):
        # clearing the page - 1 junk data's which will be trigered at the time of selection of page - 1 formula box

        if (import_progress_flag == 0):
            p2_harmonicsentries[10].config(borderwidth=2, state="normal")
            p2_harmonicsentries[10].delete(0, "end")
            p2_harmonicsentries[10].config(borderwidth=2, state="disable")

        p2_volentry.config(borderwidth=2, state="normal")
        p2_curentry.config(borderwidth=2, state="normal")
        p2_kwentry.config(borderwidth=2, state="normal")
        p2_kvaentry.config(borderwidth=2, state="normal")
        p2_kvarentry.config(borderwidth=2, state="normal")
        p2_IPFentry.config(borderwidth=2, state="normal")
        p2_TPFentry.config(borderwidth=2, state="normal")
        p2_IQentry.config(borderwidth=2, state="normal")
        p2_harmonicsentries[10].config(fg="black", disabledbackground="#F6F6F8")
        p2_vollable.config(fg="black")
        p2_curlable.config(fg="black")
        p2_kwlable.config(fg="black")
        p2_kvalable.config(fg="black")
        p2_kvarlable.config(fg="black")
        p2_IPFlable.config(fg="black")
        p2_TPFlable.config(fg="black")
        p2_IQlable.config(fg="black")

        p2_volentry.delete(0, "end")
        p2_curentry.delete(0, "end")
        p2_kwentry.delete(0, "end")
        p2_kvaentry.delete(0, "end")
        p2_kvarentry.delete(0, "end")
        p2_IPFentry.delete(0, "end")
        p2_TPFentry.delete(0, "end")
        p2_IQentry.delete(0, "end")
#------Clear junk entries ended--------------------#

#----------------------Selection code started--------------#
def selection_process(*args):
    if (import_progress_flag == 0): enable_modified_flag()
    if (astranotebook.index(astranotebook.select()) == 0):
        # page - 1 formula selection box

        if formula_combo.get() == "V, I, IPF, TPF":
            clear_junk_entires()

            # Enabled entries
            volentry.config(borderwidth=2, state="normal")
            curentry.config(borderwidth=2, state="normal")
            IPFentry.config(borderwidth=2, state="normal")
            TPFentry.config(borderwidth=2, state="normal")

            # Disabled entires
            kwentry.config(borderwidth=2, state="disable")
            kvaentry.config(borderwidth=2, state="disable")
            kvarentry.config(borderwidth=2, state="disable")
            IQentry.config(borderwidth=2, state="disable")

        elif formula_combo.get() == "V, kW, IPF, TPF":
            clear_junk_entires()

            # Enabled entries
            volentry.config(borderwidth=2, state="normal")
            kwentry.config(borderwidth=2, state="normal")
            IPFentry.config(borderwidth=2, state="normal")
            TPFentry.config(borderwidth=2, state="normal")

            # Disabled entires
            curentry.config(borderwidth=2, state="disable")
            kvaentry.config(borderwidth=2, state="disable")
            kvarentry.config(borderwidth=2, state="disable")
            IQentry.config(borderwidth=2, state="disable")

        elif formula_combo.get() == "V, kW, kVA, TPF":
            clear_junk_entires()

            # Enabled entries
            volentry.config(borderwidth=2, state="normal")
            kwentry.config(borderwidth=2, state="normal")
            kvaentry.config(borderwidth=2, state="normal")
            TPFentry.config(borderwidth=2, state="normal")

            # Disabled entires
            curentry.config(borderwidth=2, state="disable")
            kvarentry.config(borderwidth=2, state="disable")
            IPFentry.config(borderwidth=2, state="disable")
            IQentry.config(borderwidth=2, state="disable")

        elif formula_combo.get() == "V, kVA, IPF, TPF":
            clear_junk_entires()

            # Enabled entries
            volentry.config(borderwidth=2, state="normal")
            kvaentry.config(borderwidth=2, state="normal")
            IPFentry.config(borderwidth=2, state="normal")
            TPFentry.config(borderwidth=2, state="normal")

            # Disabled entires
            curentry.config(borderwidth=2, state="disable")
            kwentry.config(borderwidth=2, state="disable")
            kvarentry.config(borderwidth=2, state="disable")
            IQentry.config(borderwidth=2, state="disable")

        elif formula_combo.get() == "V, kVAR":
            clear_junk_entires()

            # Enabled entries
            volentry.config(borderwidth=2, state="normal")
            kvarentry.config(borderwidth=2, state="normal")

            # Disabled entires
            curentry.config(borderwidth=2, state="disable")
            kwentry.config(borderwidth=2, state="disable")
            kvaentry.config(borderwidth=2, state="disable")
            IPFentry.config(borderwidth=2, state="disable")
            TPFentry.config(borderwidth=2, state="disable")
            IQentry.config(borderwidth=2, state="disable")

        elif formula_combo.get() == "IQ":
            clear_junk_entires()

            # Enabled entries
            IQentry.config(borderwidth=2, state="normal")

            # Disabled entires
            volentry.config(borderwidth=2, state="disable")
            curentry.config(borderwidth=2, state="disable")
            kwentry.config(borderwidth=2, state="disable")
            kvaentry.config(borderwidth=2, state="disable")
            kvarentry.config(borderwidth=2, state="disable")
            IPFentry.config(borderwidth=2, state="disable")
            TPFentry.config(borderwidth=2, state="disable")
        elif formula_combo.get() == "select":
            clear_junk_entires()
            print("inside page -1 select")
            # Disabled entires
            volentry.config(borderwidth=2, state="disable")
            curentry.config(borderwidth=2, state="disable")
            kwentry.config(borderwidth=2, state="disable")
            kvaentry.config(borderwidth=2, state="disable")
            kvarentry.config(borderwidth=2, state="disable")
            IPFentry.config(borderwidth=2, state="disable")
            TPFentry.config(borderwidth=2, state="disable")
            IQentry.config(borderwidth=2, state="disable")
            print("inside page -1 disabled")

    elif (astranotebook.index(astranotebook.select()) == 1):
        # page - 2 formula selection box
        if p2_formula_combo.get() == "V, I, IPF, TPF":
            clear_junk_entires()

            # Enabled entries
            p2_volentry.config(borderwidth=2, state="normal")
            p2_curentry.config(borderwidth=2, state="normal")
            p2_IPFentry.config(borderwidth=2, state="normal")
            p2_TPFentry.config(borderwidth=2, state="normal")

            # Disabled entires
            p2_kwentry.config(borderwidth=2, state="disable")
            p2_kvaentry.config(borderwidth=2, state="disable")
            p2_kvarentry.config(borderwidth=2, state="disable")
            p2_IQentry.config(borderwidth=2, state="disable")

        elif p2_formula_combo.get() == "V, kW, IPF, TPF":
            clear_junk_entires()

            # Enabled entries
            p2_volentry.config(borderwidth=2, state="normal")
            p2_kwentry.config(borderwidth=2, state="normal")
            p2_IPFentry.config(borderwidth=2, state="normal")
            p2_TPFentry.config(borderwidth=2, state="normal")

            # Disabled entires
            p2_curentry.config(borderwidth=2, state="disable")
            p2_kvaentry.config(borderwidth=2, state="disable")
            p2_kvarentry.config(borderwidth=2, state="disable")
            p2_IQentry.config(borderwidth=2, state="disable")

        elif p2_formula_combo.get() == "V, kW, kVA, TPF":
            clear_junk_entires()

            # Enabled entries
            p2_volentry.config(borderwidth=2, state="normal")
            p2_kwentry.config(borderwidth=2, state="normal")
            p2_kvaentry.config(borderwidth=2, state="normal")
            p2_TPFentry.config(borderwidth=2, state="normal")

            # Disabled entires
            p2_curentry.config(borderwidth=2, state="disable")
            p2_kvarentry.config(borderwidth=2, state="disable")
            p2_IPFentry.config(borderwidth=2, state="disable")
            p2_IQentry.config(borderwidth=2, state="disable")

        elif p2_formula_combo.get() == "V, kVA, IPF, TPF":
            clear_junk_entires()

            # Enabled entries
            p2_volentry.config(borderwidth=2, state="normal")
            p2_kvaentry.config(borderwidth=2, state="normal")
            p2_IPFentry.config(borderwidth=2, state="normal")
            p2_TPFentry.config(borderwidth=2, state="normal")

            # Disabled entires
            p2_curentry.config(borderwidth=2, state="disable")
            p2_kwentry.config(borderwidth=2, state="disable")
            p2_kvarentry.config(borderwidth=2, state="disable")
            p2_IQentry.config(borderwidth=2, state="disable")

        elif p2_formula_combo.get() == "V, kVAR":
            clear_junk_entires()

            # Enabled entries
            p2_volentry.config(borderwidth=2, state="normal")
            p2_kvarentry.config(borderwidth=2, state="normal")

            # Disabled entires
            p2_curentry.config(borderwidth=2, state="disable")
            p2_kwentry.config(borderwidth=2, state="disable")
            p2_kvaentry.config(borderwidth=2, state="disable")
            p2_IPFentry.config(borderwidth=2, state="disable")
            p2_TPFentry.config(borderwidth=2, state="disable")
            p2_IQentry.config(borderwidth=2, state="disable")

        elif p2_formula_combo.get() == "IQ":
            clear_junk_entires()

            # Enabled entries
            p2_IQentry.config(borderwidth=2, state="normal")

            # Disabled entires
            p2_volentry.config(borderwidth=2, state="disable")
            p2_curentry.config(borderwidth=2, state="disable")
            p2_kwentry.config(borderwidth=2, state="disable")
            p2_kvaentry.config(borderwidth=2, state="disable")
            p2_kvarentry.config(borderwidth=2, state="disable")
            p2_IPFentry.config(borderwidth=2, state="disable")
            p2_TPFentry.config(borderwidth=2, state="disable")
        elif p2_formula_combo.get() == "select":
            print("inside page -2 select")

            clear_junk_entires()

            # Disabled entires
            p2_volentry.config(borderwidth=2, state="disable")
            p2_curentry.config(borderwidth=2, state="disable")
            p2_kwentry.config(borderwidth=2, state="disable")
            p2_kvaentry.config(borderwidth=2, state="disable")
            p2_kvarentry.config(borderwidth=2, state="disable")
            p2_IPFentry.config(borderwidth=2, state="disable")
            p2_TPFentry.config(borderwidth=2, state="disable")
            p2_IQentry.config(borderwidth=2, state="disable")
#----------------Selection code ended------------#

#------------------Update result-----------------#
def update_result(*args):
    if (astranotebook.index(astranotebook.select()) == 0):
        # page -1 formula dynamic calculation
        enable_modified_flag()
        if formula_combo.get() == "V, I, IPF, TPF":
            # calculating the results form the user input ( Voltage , current, IPF, TPF)
            try:
                voltage = int(volentry.get())
                vollable.config(fg="black")
            except:
                voltage = 1
                vollable.config(fg="red")
            try:
                current = int(curentry.get())
                curlable.config(fg="black")
            except:
                current = 1
                curlable.config(fg="red")
            try:
                ipf = float(IPFentry.get())
                IPFlable.config(fg="black")
            except:
                ipf = 1
                IPFlable.config(fg="red")
            try:
                tpf = float(TPFentry.get())
                TPFlable.config(fg="black")
            except:
                tpf = 1
                TPFlable.config(fg="red")

            global lab
            for widget in frame_4.winfo_children():
                widget.destroy()
            for widget_1 in frame_6.winfo_children():
                widget_1.destroy()
            pri_less_sec.config(text="")
            astranotebook.tab(3, text="True Power Factor*")
            head_notebook.config(text="True Power Factor Performance Calculator *")
            location_field.config(text="location :" + str(selected_directory))
            p2_location_field.config(text="location :" + str(selected_directory))
            p3_location_field.config(text="location :" + str(selected_directory))
            p4_location_field.config(text="location :" + str(selected_directory))
            print(modified_flag)

            kw = math.sqrt(3) * (voltage / 1000) * current * ipf
            kva = math.sqrt(3) * (voltage / 1000) * current
            kvar = kw * (math.tan(math.acos(ipf)) - math.tan(math.acos(tpf)))
            result = round((kvar) / ((math.sqrt(3) * voltage) / 1000))

            # inserting the results to the kva
            kvaentry.config(borderwidth=2, state=NORMAL,
                            disabledforeground="black")
            kvaentry.delete(0, "end")
            kvaentry.insert(0, (str(round(kva))))
            kvaentry.config(borderwidth=2, state=DISABLED, disabledforeground="black")

            # inserting the results to the kw
            kwentry.config(borderwidth=2, state=NORMAL,
                           disabledforeground="black")
            kwentry.delete(0, "end")
            kwentry.insert(0, (str(round(kw))))
            kwentry.config(borderwidth=2, state=DISABLED, disabledforeground="black")

            # inserting the results to the kvar
            kvarentry.config(borderwidth=2, state=NORMAL, disabledforeground="black")
            kvarentry.delete(0, "end")
            kvarentry.insert(0, (str(round(kvar))))
            kvarentry.config(borderwidth=2, state=DISABLED, disabledforeground="black")

            # inserting the results to the IQ
            IQentry.config(borderwidth=2, state=NORMAL, disabledforeground="black")
            IQentry.delete(0, "end")
            IQentry.insert(0, (str(result)))
            IQentry.config(borderwidth=2, state=DISABLED, disabledforeground="black")

            # Copying the IQ value to first current order
            harmonicsentries[10].config(borderwidth=2, state=NORMAL,
                                        disabledforeground="black")
            harmonicsentries[10].delete(0, "end")
            harmonicsentries[10].insert(0, (str(result)))
            harmonicsentries[10].config(borderwidth=2, state=DISABLED,
                                        disabledforeground="black")

        elif formula_combo.get() == "V, kW, IPF, TPF":
            # calculating the results form the user input ( Voltage , kw, IPF, TPF)
            try:
                voltage = int(volentry.get())
                vollable.config(fg="black")
            except:
                voltage = 1
                vollable.config(fg="red")
            try:
                kw = int(kwentry.get())
                kwlable.config(fg="black")
            except:
                kw = 1
                kwlable.config(fg="red")
            try:
                ipf = float(IPFentry.get())
                IPFlable.config(fg="black")
            except:
                ipf = 1
                IPFlable.config(fg="red")
            try:
                tpf = float(TPFentry.get())
                TPFlable.config(fg="black")
            except:
                tpf = 1
                TPFlable.config(fg="red")

            current = kw / (math.sqrt(3) * (voltage / 1000) * ipf)
            kva = kw / ipf
            kvar = kw * (math.tan(math.acos(ipf)) - math.tan(math.acos(tpf)))
            result = round((kvar) / ((math.sqrt(3) * voltage) / 1000))

            # inserting the results to the current
            curentry.config(borderwidth=2, state=NORMAL, disabledforeground="black")
            curentry.delete(0, "end")
            curentry.insert(0, (str(round(current))))
            curentry.config(borderwidth=2, state=DISABLED, disabledforeground="black")

            # inserting the results to the kva
            kvaentry.config(borderwidth=2, state=NORMAL, disabledforeground="black")
            kvaentry.delete(0, "end")
            kvaentry.insert(0, (str(round(kva))))
            kvaentry.config(borderwidth=2, state=DISABLED, disabledforeground="black")

            # inserting the results to the kvar
            kvarentry.config(borderwidth=2, state=NORMAL, disabledforeground="black")
            kvarentry.delete(0, "end")
            kvarentry.insert(0, (str(round(kvar))))
            kvarentry.config(borderwidth=2, state=DISABLED, disabledforeground="black")

            # inserting the results to the IQ
            IQentry.config(borderwidth=2, state=NORMAL, disabledforeground="black")
            IQentry.delete(0, "end")
            IQentry.insert(0, (str(result)))
            IQentry.config(borderwidth=2, state=DISABLED, disabledforeground="black")

            # Copying the IQ value to first current order
            harmonicsentries[10].config(borderwidth=2, state=NORMAL,
                                        disabledforeground="black")
            harmonicsentries[10].delete(0, "end")
            harmonicsentries[10].insert(0, (str(result)))
            harmonicsentries[10].config(borderwidth=2, state=DISABLED,
                                        disabledforeground="black")

        elif formula_combo.get() == "V, kW, kVA, TPF":
            # calculating the results form the user input ( Voltage , kw, kVA, TPF)
            try:
                voltage = int(volentry.get())
                vollable.config(fg="black")
            except:
                voltage = 1
                vollable.config(fg="red")
            try:
                kw = int(kwentry.get())
                kwlable.config(fg="black")
            except:
                kw = 1
                kwlable.config(fg="red")
            try:
                kva = float(kvaentry.get())
                kvalable.config(fg="black")
            except:
                kva = 1
                kvalable.config(fg="red")
            try:
                tpf = float(TPFentry.get())
                TPFlable.config(fg="black")
            except:
                tpf = 1
                TPFlable.config(fg="red")

            # kvar = math.sqrt((kva ** 2) - (kw ** 2))
            try:
                ipf = float(kw / kva)
            except:
                ipf = 1

            current = kva / (math.sqrt(3) * (voltage / 1000))
            kvar = kw * (math.tan(math.acos(ipf)) - math.tan(math.acos(tpf)))
            result = round((kvar) / ((math.sqrt(3) * voltage) / 1000))

            # inserting the results to the current
            curentry.config(borderwidth=2, state=NORMAL, disabledforeground="black")
            curentry.delete(0, "end")
            curentry.insert(0, (str(round(current))))
            curentry.config(borderwidth=2, state=DISABLED, disabledforeground="black")

            # inserting the results to the kvar
            kvarentry.config(borderwidth=2, state=NORMAL,
                             disabledforeground="black")
            kvarentry.delete(0, "end")
            kvarentry.insert(0, (str(round(kvar))))
            kvarentry.config(borderwidth=2, state=DISABLED, disabledforeground="black")

            # inserting the results to the IPF
            IPFentry.config(borderwidth=2, state=NORMAL, disabledforeground="black")
            IPFentry.delete(0, "end")
            IPFentry.insert(0, (str(round(ipf, 2))))
            IPFentry.config(borderwidth=2, state=DISABLED, disabledforeground="black")

            # inserting the results to the IQ
            IQentry.config(borderwidth=2, state=NORMAL, disabledforeground="black")
            IQentry.delete(0, "end")
            IQentry.insert(0, (str(result)))
            IQentry.config(borderwidth=2, state=DISABLED, disabledforeground="black")

            # Copying the IQ value to first current order
            harmonicsentries[10].config(borderwidth=2, state=NORMAL,
                                        disabledforeground="black")
            harmonicsentries[10].delete(0, "end")
            harmonicsentries[10].insert(0, (str(result)))
            harmonicsentries[10].config(borderwidth=2, state=DISABLED,
                                        disabledforeground="black")

        elif formula_combo.get() == "V, kVA, IPF, TPF":
            # calculating the results form the user input ( Voltage , kVA, IPF, TPF)
            try:
                voltage = int(volentry.get())
                vollable.config(fg="black")
            except:
                voltage = 1
                vollable.config(fg="red")
            try:
                kva = int(kvaentry.get())
                kvalable.config(fg="black")
            except:
                kva = 1
                kvalable.config(fg="red")
            try:
                ipf = float(IPFentry.get())
                IPFlable.config(fg="black")
            except:
                ipf = 1
                IPFlable.config(fg="red")
            try:
                tpf = float(TPFentry.get())
                TPFlable.config(fg="black")
            except:
                tpf = 1
                TPFlable.config(fg="red")

            current = kva / (math.sqrt(3) * (voltage / 1000))
            kw = kva * ipf
            kvar = kw * (math.tan(math.acos(ipf)) - math.tan(math.acos(tpf)))
            result = round((kvar) / ((math.sqrt(3) * voltage) / 1000))

            # inserting the results to the current
            curentry.config(borderwidth=2, state=NORMAL, disabledforeground="black")
            curentry.delete(0, "end")
            curentry.insert(0, (str(round(current))))
            curentry.config(borderwidth=2, state=DISABLED, disabledforeground="black")

            # inserting the results to the kw
            kwentry.config(borderwidth=2, state=NORMAL, disabledforeground="black")
            kwentry.delete(0, "end")
            kwentry.insert(0, (str(round(kw))))
            kwentry.config(borderwidth=2, state=DISABLED, disabledforeground="black")

            # inserting the results to the kvar
            kvarentry.config(borderwidth=2, state=NORMAL, disabledforeground="black")
            kvarentry.delete(0, "end")
            kvarentry.insert(0, (str(round(kvar))))
            kvarentry.config(borderwidth=2, state=DISABLED, disabledforeground="black")

            # inserting the results to the IQ
            IQentry.config(borderwidth=2, state=NORMAL, disabledforeground="black")
            IQentry.delete(0, "end")
            IQentry.insert(0, (str(result)))
            IQentry.config(borderwidth=2, state=DISABLED, disabledforeground="black")

            # Copying the IQ value to first current order
            harmonicsentries[10].config(borderwidth=2, state=NORMAL,
                                        disabledforeground="black")
            harmonicsentries[10].delete(0, "end")
            harmonicsentries[10].insert(0, (str(result)))
            harmonicsentries[10].config(borderwidth=2, state=DISABLED,
                                        disabledforeground="black")

        elif formula_combo.get() == "V, kVAR":
            # calculating the results form the user input ( Voltage , kVAR)
            try:
                voltage = int(volentry.get())
                vollable.config(fg="black")
            except:
                voltage = 1
                vollable.config(fg="red")
            try:
                kvar = int(kvarentry.get())
                kvarlable.config(fg="black")
            except:
                kvar = 1
                kvarlable.config(fg="red")

            result = round((kvar) / ((math.sqrt(3) * voltage) / 1000))

            # inserting the results to the IQ
            IQentry.config(borderwidth=2, state=NORMAL, disabledforeground="black")
            IQentry.delete(0, "end")
            IQentry.insert(0, (str(result)))
            IQentry.config(borderwidth=2, state=DISABLED, disabledforeground="black")

            # Copying the IQ value to first current order
            harmonicsentries[10].config(borderwidth=2, state=NORMAL, disabledbackground="white",
                                        disabledforeground="black")
            harmonicsentries[10].delete(0, "end")
            harmonicsentries[10].insert(0, (str(result)))
            harmonicsentries[10].config(borderwidth=2, state=DISABLED, disabledbackground="white",
                                        disabledforeground="black")

        elif formula_combo.get() == "IQ":
            harmonicsentries[10].config(borderwidth=2, state=NORMAL, disabledbackground="#F6F6F8",
                                        disabledforeground="black")
            harmonicsentries[10].delete(0, "end")
            harmonicsentries[10].insert(0, (str(IQentry.get())))
            harmonicsentries[10].config(borderwidth=2, state=DISABLED, disabledbackground="#F6F6F8",
                                        disabledforeground="black")

    elif (astranotebook.index(astranotebook.select()) == 1):
        # page -2 formula dynamic calculation
        enable_modified_flag()
        if p2_formula_combo.get() == "V, I, IPF, TPF":
            # calculating the results form the user input ( Voltage , current, IPF, TPF)
            try:
                p2_voltage = int(p2_volentry.get())
                p2_vollable.config(fg="black")
            except:
                p2_voltage = 1
                p2_vollable.config(fg="red")
            try:
                p2_current = int(p2_curentry.get())
                p2_curlable.config(fg="black")
            except:
                p2_current = 1
                p2_curlable.config(fg="red")
            try:
                p2_ipf = float(p2_IPFentry.get())
                p2_IPFlable.config(fg="black")
            except:
                p2_ipf = 1
                p2_IPFlable.config(fg="red")
            try:
                p2_tpf = float(p2_TPFentry.get())
                p2_TPFlable.config(fg="black")
            except:
                p2_tpf = 1
                p2_TPFlable.config(fg="red")

            p2_kw = math.sqrt(3) * (p2_voltage / 1000) * p2_current * p2_ipf
            p2_kva = math.sqrt(3) * (p2_voltage / 1000) * p2_current
            p2_kvar = p2_kw * (math.tan(math.acos(p2_ipf)) - math.tan(math.acos(p2_tpf)))
            p2_result = round((p2_kvar) / ((math.sqrt(3) * p2_voltage) / 1000))

            # inserting the results to the kva
            p2_kvaentry.config(borderwidth=2, state=NORMAL,
                               disabledforeground="black")
            p2_kvaentry.delete(0, "end")
            p2_kvaentry.insert(0, (str(round(p2_kva))))
            p2_kvaentry.config(borderwidth=2, state=DISABLED, disabledforeground="black")

            # inserting the results to the kw
            p2_kwentry.config(borderwidth=2, state=NORMAL,
                              disabledforeground="black")
            p2_kwentry.delete(0, "end")
            p2_kwentry.insert(0, (str(round(p2_kw))))
            p2_kwentry.config(borderwidth=2, state=DISABLED, disabledforeground="black")

            # inserting the results to the kvar
            p2_kvarentry.config(borderwidth=2, state=NORMAL, disabledforeground="black")
            p2_kvarentry.delete(0, "end")
            p2_kvarentry.insert(0, (str(round(p2_kvar))))
            p2_kvarentry.config(borderwidth=2, state=DISABLED, disabledforeground="black")

            # inserting the results to the IQ
            p2_IQentry.config(borderwidth=2, state=NORMAL, disabledforeground="black")
            p2_IQentry.delete(0, "end")
            p2_IQentry.insert(0, (str(p2_result)))
            p2_IQentry.config(borderwidth=2, state=DISABLED, disabledforeground="black")

            # Copying the IQ value to first current order
            p2_harmonicsentries[10].config(borderwidth=2, state=NORMAL,
                                           disabledforeground="black")
            p2_harmonicsentries[10].delete(0, "end")
            p2_harmonicsentries[10].insert(0, (str(p2_result)))
            p2_harmonicsentries[10].config(borderwidth=2, state=DISABLED,
                                           disabledforeground="black")

        elif p2_formula_combo.get() == "V, kW, IPF, TPF":
            # calculating the results form the user input ( Voltage , kw, IPF, TPF)
            try:
                p2_voltage = int(p2_volentry.get())
                p2_vollable.config(fg="black")
            except:
                p2_voltage = 1
                p2_vollable.config(fg="red")
            try:
                p2_kw = int(p2_kwentry.get())
                p2_kwlable.config(fg="black")
            except:
                p2_kw = 1
                p2_kwlable.config(fg="red")
            try:
                p2_ipf = float(p2_IPFentry.get())
                p2_IPFlable.config(fg="black")
            except:
                p2_ipf = 1
                p2_IPFlable.config(fg="red")
            try:
                p2_tpf = float(p2_TPFentry.get())
                p2_TPFlable.config(fg="black")
            except:
                p2_tpf = 1
                p2_TPFlable.config(fg="red")

            p2_current = p2_kw / (math.sqrt(3) * (p2_voltage / 1000) * p2_ipf)
            p2_kva = p2_kw / p2_ipf
            p2_kvar = p2_kw * (math.tan(math.acos(p2_ipf)) - math.tan(math.acos(p2_tpf)))
            p2_result = round((p2_kvar) / ((math.sqrt(3) * p2_voltage) / 1000))

            # inserting the results to the current
            p2_curentry.config(borderwidth=2, state=NORMAL, disabledforeground="black")
            p2_curentry.delete(0, "end")
            p2_curentry.insert(0, (str(round(p2_current))))
            p2_curentry.config(borderwidth=2, state=DISABLED, disabledforeground="black")

            # inserting the results to the kva
            p2_kvaentry.config(borderwidth=2, state=NORMAL, disabledforeground="black")
            p2_kvaentry.delete(0, "end")
            p2_kvaentry.insert(0, (str(round(p2_kva))))
            p2_kvaentry.config(borderwidth=2, state=DISABLED, disabledforeground="black")

            # inserting the results to the kvar
            p2_kvarentry.config(borderwidth=2, state=NORMAL, disabledforeground="black")
            p2_kvarentry.delete(0, "end")
            p2_kvarentry.insert(0, (str(round(p2_kvar))))
            p2_kvarentry.config(borderwidth=2, state=DISABLED, disabledforeground="black")

            # inserting the results to the IQ
            p2_IQentry.config(borderwidth=2, state=NORMAL, disabledforeground="black")
            p2_IQentry.delete(0, "end")
            p2_IQentry.insert(0, (str(p2_result)))
            p2_IQentry.config(borderwidth=2, state=DISABLED, disabledforeground="black")

            # Copying the IQ value to first current order
            p2_harmonicsentries[10].config(borderwidth=2, state=NORMAL,
                                           disabledforeground="black")
            p2_harmonicsentries[10].delete(0, "end")
            p2_harmonicsentries[10].insert(0, (str(p2_result)))
            p2_harmonicsentries[10].config(borderwidth=2, state=DISABLED,
                                           disabledforeground="black")

        elif p2_formula_combo.get() == "V, kW, kVA, TPF":
            # calculating the results form the user input ( Voltage , kw, kVA, TPF)
            try:
                p2_voltage = int(p2_volentry.get())
                p2_vollable.config(fg="black")
            except:
                p2_voltage = 1
                p2_vollable.config(fg="red")
            try:
                p2_kw = int(p2_kwentry.get())
                p2_kwlable.config(fg="black")
            except:
                p2_kw = 1
                p2_kwlable.config(fg="red")
            try:
                p2_kva = float(p2_kvaentry.get())
                p2_kvalable.config(fg="black")
            except:
                p2_kva = 1
                p2_kvalable.config(fg="red")
            try:
                p2_tpf = float(p2_TPFentry.get())
                p2_TPFlable.config(fg="black")
            except:
                p2_tpf = 1
                p2_TPFlable.config(fg="red")

            # kvar = math.sqrt((kva ** 2) - (kw ** 2))
            try:
                p2_ipf = float(p2_kw / p2_kva)
            except:
                p2_ipf = 1

            p2_current = p2_kva / (math.sqrt(3) * (p2_voltage / 1000))
            p2_kvar = p2_kw * (math.tan(math.acos(p2_ipf)) - math.tan(math.acos(p2_tpf)))
            p2_result = round((p2_kvar) / ((math.sqrt(3) * p2_voltage) / 1000))

            # inserting the results to the current
            p2_curentry.config(borderwidth=2, state=NORMAL, disabledforeground="black")
            p2_curentry.delete(0, "end")
            p2_curentry.insert(0, (str(round(p2_current))))
            p2_curentry.config(borderwidth=2, state=DISABLED, disabledforeground="black")

            # inserting the results to the kvar
            p2_kvarentry.config(borderwidth=2, state=NORMAL,
                                disabledforeground="black")
            p2_kvarentry.delete(0, "end")
            p2_kvarentry.insert(0, (str(round(p2_kvar))))
            p2_kvarentry.config(borderwidth=2, state=DISABLED, disabledforeground="black")

            # inserting the results to the IPF
            p2_IPFentry.config(borderwidth=2, state=NORMAL, disabledforeground="black")
            p2_IPFentry.delete(0, "end")
            p2_IPFentry.insert(0, (str(round(p2_ipf, 2))))
            p2_IPFentry.config(borderwidth=2, state=DISABLED, disabledforeground="black")

            # inserting the results to the IQ
            p2_IQentry.config(borderwidth=2, state=NORMAL, disabledforeground="black")
            p2_IQentry.delete(0, "end")
            p2_IQentry.insert(0, (str(p2_result)))
            p2_IQentry.config(borderwidth=2, state=DISABLED, disabledforeground="black")

            # Copying the IQ value to first current order
            p2_harmonicsentries[10].config(borderwidth=2, state=NORMAL,
                                           disabledforeground="black")
            p2_harmonicsentries[10].delete(0, "end")
            p2_harmonicsentries[10].insert(0, (str(p2_result)))
            p2_harmonicsentries[10].config(borderwidth=2, state=DISABLED,
                                           disabledforeground="black")

        elif p2_formula_combo.get() == "V, kVA, IPF, TPF":
            # calculating the results form the user input ( Voltage , kVA, IPF, TPF)
            try:
                p2_voltage = int(p2_volentry.get())
                p2_vollable.config(fg="black")
            except:
                p2_voltage = 1
                p2_vollable.config(fg="red")
            try:
                p2_kva = int(p2_kvaentry.get())
                p2_kvalable.config(fg="black")
            except:
                p2_kva = 1
                p2_kvalable.config(fg="red")
            try:
                p2_ipf = float(p2_IPFentry.get())
                p2_IPFlable.config(fg="black")
            except:
                p2_ipf = 1
                p2_IPFlable.config(fg="red")
            try:
                p2_tpf = float(p2_TPFentry.get())
                p2_TPFlable.config(fg="black")
            except:
                p2_tpf = 1
                p2_TPFlable.config(fg="red")

            p2_current = p2_kva / (math.sqrt(3) * (p2_voltage / 1000))
            p2_kw = p2_kva * p2_ipf
            p2_kvar = p2_kw * (math.tan(math.acos(p2_ipf)) - math.tan(math.acos(p2_tpf)))
            p2_result = round((p2_kvar) / ((math.sqrt(3) * p2_voltage) / 1000))

            # inserting the results to the current
            p2_curentry.config(borderwidth=2, state=NORMAL, disabledforeground="black")
            p2_curentry.delete(0, "end")
            p2_curentry.insert(0, (str(round(p2_current))))
            p2_curentry.config(borderwidth=2, state=DISABLED, disabledforeground="black")

            # inserting the results to the kw
            p2_kwentry.config(borderwidth=2, state=NORMAL, disabledforeground="black")
            p2_kwentry.delete(0, "end")
            p2_kwentry.insert(0, (str(round(p2_kw))))
            p2_kwentry.config(borderwidth=2, state=DISABLED, disabledforeground="black")

            # inserting the results to the kvar
            p2_kvarentry.config(borderwidth=2, state=NORMAL, disabledforeground="black")
            p2_kvarentry.delete(0, "end")
            p2_kvarentry.insert(0, (str(round(p2_kvar))))
            p2_kvarentry.config(borderwidth=2, state=DISABLED, disabledforeground="black")

            # inserting the results to the IQ
            p2_IQentry.config(borderwidth=2, state=NORMAL, disabledforeground="black")
            p2_IQentry.delete(0, "end")
            p2_IQentry.insert(0, (str(p2_result)))
            p2_IQentry.config(borderwidth=2, state=DISABLED, disabledforeground="black")

            # Copying the IQ value to first current order
            p2_harmonicsentries[10].config(borderwidth=2, state=NORMAL,
                                           disabledforeground="black")
            p2_harmonicsentries[10].delete(0, "end")
            p2_harmonicsentries[10].insert(0, (str(p2_result)))
            p2_harmonicsentries[10].config(borderwidth=2, state=DISABLED,
                                           disabledforeground="black")

        elif p2_formula_combo.get() == "V, kVAR":
            # calculating the results form the user input ( Voltage , kVAR)
            try:
                p2_voltage = int(p2_volentry.get())
                p2_vollable.config(fg="black")
            except:
                p2_voltage = 1
                p2_vollable.config(fg="red")
            try:
                p2_kvar = int(p2_kvarentry.get())
                p2_kvarlable.config(fg="black")
            except:
                p2_kvar = 1
                p2_kvarlable.config(fg="red")

            p2_result = round((p2_kvar) / ((math.sqrt(3) * p2_voltage) / 1000))

            # inserting the results to the IQ
            p2_IQentry.config(borderwidth=2, state=NORMAL, disabledforeground="black")
            p2_IQentry.delete(0, "end")
            p2_IQentry.insert(0, (str(p2_result)))
            p2_IQentry.config(borderwidth=2, state=DISABLED, disabledforeground="black")

            # Copying the IQ value to first current order
            p2_harmonicsentries[10].config(borderwidth=2, state=NORMAL, disabledbackground="white",
                                           disabledforeground="black")
            p2_harmonicsentries[10].delete(0, "end")
            p2_harmonicsentries[10].insert(0, (str(p2_result)))
            p2_harmonicsentries[10].config(borderwidth=2, state=DISABLED, disabledbackground="white",
                                           disabledforeground="black")

        elif p2_formula_combo.get() == "IQ":
            p2_harmonicsentries[10].config(borderwidth=2, state=NORMAL, disabledbackground="#F6F6F8",
                                           disabledforeground="black")
            p2_harmonicsentries[10].delete(0, "end")
            p2_harmonicsentries[10].insert(0, (str(p2_IQentry.get())))
            p2_harmonicsentries[10].config(borderwidth=2, state=DISABLED, disabledbackground="#F6F6F8",
                                           disabledforeground="black")
#----------------------Update result is ended----------------------#

# page -1 dynamic calculation bind function calls

volentry.bind('<KeyRelease>', update_result)
curentry.bind('<KeyRelease>', update_result)
kwentry.bind('<KeyRelease>', update_result)
kvaentry.bind('<KeyRelease>', update_result)
kvarentry.bind('<KeyRelease>', update_result)
IPFentry.bind('<KeyRelease>', update_result)
TPFentry.bind('<KeyRelease>', update_result)
IQentry.bind('<KeyRelease>', update_result)
basenentry.bind('<KeyRelease>', update_result)  # -->new

# Creating Combo Box

formula_combo = ttk.Combobox(astrap1_frame, font=('Verdana 12'), state="readonly", value=formula_title)
formula_combo.current(0)
formula_combo.place(x=810, y=130)

# Binding the combo box
formula_combo.bind("<<ComboboxSelected>>", selection_process)

#-----------------Ended--------------#

#---------------------Compile save data is started-----------------#
def compile_save_data():
    # collecting the all the intial data value

    disable_modified_flag()
    global sizing_data, p4_data_entries
    global rec_tab_val, minimum_kw
    p1_data_entries = []
    p2_data_entries = []
    p3_data_entries = []
    p4_data_entries = []
    p4_sugg_entries = []
    #rec_tab_val
    #minimum_kw


    for x in range(0, len(harmonicsentries)):
        p1_data_entries.append(harmonicsentries[x].get())
    for x in range(0, len(p2_harmonicsentries)):
        p2_data_entries.append(p2_harmonicsentries[x].get())
    for x in range(0, len(p3_harmonicsentries)):
        p3_data_entries.append(p3_harmonicsentries[x].get())

    try:
        for x in range(len(rec_tab_val)):
            for y in range(len(rec_tab_val[0])):
                p4_data_entries.append(rec_tab_val[x][y])
        for x in range(0, len(minimum_kw)):
            p4_sugg_entries.append(minimum_kw[x])
            print("Done..............................")
    except:
        p4_data_entries.append(None)
        p4_sugg_entries.append(None)

        # enable_at_save()
    sizing_data = {
        'p1_harmonicentries': p1_data_entries,
        'p1_astra_rating': AHFsizeentry.get(),
        'p1_amb_astra_rating': AHFsize1entry.get(),
        'p1_fhz': page1_frequency,
        'p1_notch': basenentry.get(),
        'p1_amb_temp': ambtempentry.get(),
        'p1_amp_fact': ambfactorentry.get(),
        'p1_reactive_selection': formula_combo.get(),
        'p1_voltage': volentry.get(),
        'p1_current': curentry.get(),
        'p1_active': kwentry.get(),
        'p1_apparent': kvaentry.get(),
        'p1_reactive': kvarentry.get(),
        'p1_ipf': IPFentry.get(),
        'p1_tpf': TPFentry.get(),
        'p1_iq': IQentry.get(),
        'p1_comments': comment_box_message.get(1.0, "end"),
        'p2_harmonicentries': p2_data_entries,
        'p2_astra_rating': p2_AHFsizeentry.get(),
        'p2_amb_astra_rating': p2_AHFsize1entry.get(),
        'p2_fhz': page2_frequency,
        'p2_notch': p2_basenentry.get(),
        'p2_amb_temp': p2_ambtempentry.get(),
        'p2_amp_fact': p2_ambfactorentry.get(),
        'p2_reactive_selection': p2_formula_combo.get(),
        'p2_voltage': p2_volentry.get(),
        'p2_current': p2_curentry.get(),
        'p2_active': p2_kwentry.get(),
        'p2_apparent': p2_kvaentry.get(),
        'p2_reactive': p2_kvarentry.get(),
        'p2_ipf': p2_IPFentry.get(),
        'p2_tpf': p2_TPFentry.get(),
        'p2_iq': p2_IQentry.get(),
        'p2_comments': p2_comment_box_message.get(1.0, "end"),
        'p3_harmonicentries': p3_data_entries,
        'p3_astra_rating': p3_AHFsizeentry.get(),
        'p3_amb_astra_rating': p3_AHFsize1entry.get(),
        'p3_fhz': page3_frequency,
        'p3_notch': p3_basenentry.get(),
        'p3_amb_temp': p3_ambtempentry.get(),
        'p3_amp_fact': p3_ambfactorentry.get(),
        'p3_comments': p3_comment_box_message.get(1.0, "end"),
        'p4_primary value': primary_entry.get(),
        'p4_secondary value': secondary_entry.get(),
        'p4_data values': p4_data_entries,
        'p4_suggestion values': p4_sugg_entries,
    }
    print("P4_Entries value : ", p4_sugg_entries)
    print("P4_Entries value : ", p4_data_entries)


#-------------------Compile save data is ended---------------#

#-------------------Clear contents started--------------------#
def clear_contents():
    # Clearing Page 1 contents
    if (astranotebook.index(astranotebook.select()) == 0):
        # Deleting the pervious loaded status
        status_p1_entry.config(borderwidth=2, state="normal")
        status_p1_entry.delete(0, "end")
        status_p1_entry.config(borderwidth=2, state="disable")

        # clearing the entry widget
        global page1_frequency
        for x in range(0, 20):
            if (x == 0):
                harmonicsentries[x].config(state='normal')
                harmonicsentries[x].delete(0, "end")
                harmonicsentries[x].insert(0, "1")
                harmonicsentries[x].config(state='disable')
            elif (x == 10):
                harmonicsentries[x].config(state='normal')
                harmonicsentries[x].delete(0, "end")
                harmonicsentries[x].config(state='disable')
            else:
                harmonicsentries[x].delete(0, "end")
        formula_combo.current(0)
        print("page-1")
        selection_process()
        print("PROCESS COMP")
        # enable_modified_flag()
        basenentry.delete(0, "end")
        ambtempentry.delete(0, "end")
        ambfactorentry.delete(0, "end")
        AHFsizeentry.delete(0, "end")
        AHFsize1entry.delete(0, "end")
        fhzbtn.config(bg="white", fg="black", font=font.Font(family='Calibri', size=10, weight=font.NORMAL))
        shzbtn.config(bg="white", fg="black", font=font.Font(family='Calibri', size=10, weight=font.NORMAL))
        highnotchbtn.config(bg="white", fg="black", font=font.Font(family='Calibri', size=10, weight=font.NORMAL))
        mildnotchbtn.config(bg="white", fg="black", font=font.Font(family='Calibri', size=10, weight=font.NORMAL))
        lownotchbtn.config(bg="white", fg="black", font=font.Font(family='Calibri', size=10, weight=font.NORMAL))
        vollable.config(fg="black")
        curlable.config(fg="black")
        kwlable.config(fg="black")
        kvalable.config(fg="black")
        kvarlable.config(fg="black")
        IPFlable.config(fg="black")
        TPFlable.config(fg="black")
        IQlable.config(fg="black")
        frequencylable.config(fg="black")
        ambtemplable.config(fg="black")
        ambfactorlable.config(fg="black")
        basenlable.config(fg="black")
        page1_frequency = 0
        comment_box_message.delete(1.0, "end")
        comment_box_message.insert(1.0, "Comments here...")
        comment_box_message.config(fg='gray')
    # Clearing Page 2 contents
    if (astranotebook.index(astranotebook.select()) == 1):
        # Deleting the pervious loaded status
        status_p2_entry.config(borderwidth=2, state="normal")
        status_p2_entry.delete(0, "end")
        status_p2_entry.config(borderwidth=2, state="disable")

        # clearing the entry widget
        global page2_frequency
        for x in range(0, 30):
            if (x == 0):
                p2_harmonicsentries[x].config(state='normal')
                p2_harmonicsentries[x].delete(0, "end")
                p2_harmonicsentries[x].insert(0, "1")
                p2_harmonicsentries[x].config(state='disable')
            elif (x == 10):
                p2_harmonicsentries[x].config(state='normal')
                p2_harmonicsentries[x].delete(0, "end")
                p2_harmonicsentries[x].config(state='disable')
            else:
                p2_harmonicsentries[x].delete(0, "end")
        p2_formula_combo.current(0)
        print("page-2")
        selection_process()
        # enable_modified_flag()
        p2_basenentry.delete(0, "end")
        p2_ambtempentry.delete(0, "end")
        p2_ambfactorentry.delete(0, "end")
        p2_AHFsizeentry.delete(0, "end")
        p2_AHFsize1entry.delete(0, "end")
        p2_fhzbtn.config(bg="white", fg="black", font=font.Font(family='Calibri', size=10, weight=font.NORMAL))
        p2_shzbtn.config(bg="white", fg="black", font=font.Font(family='Calibri', size=10, weight=font.NORMAL))
        p2_mildnotchbtn.config(bg="white", fg="black", font=font.Font(family='Calibri', size=10, weight=font.NORMAL))
        p2_lownotchbtn.config(bg="white", fg="black", font=font.Font(family='Calibri', size=10, weight=font.NORMAL))
        p2_vollable.config(fg="black")
        p2_curlable.config(fg="black")
        p2_kwlable.config(fg="black")
        p2_kvalable.config(fg="black")
        p2_kvarlable.config(fg="black")
        p2_IPFlable.config(fg="black")
        p2_TPFlable.config(fg="black")
        p2_IQlable.config(fg="black")
        p2_frequencylable.config(fg="black")
        p2_ambtemplable.config(fg="black")
        p2_ambfactorlable.config(fg="black")
        p2_basenlable.config(fg="black")
        page2_frequency = 0
        p2_comment_box_message.delete(1.0, "end")
        p2_comment_box_message.insert(1.0, "Comments here...")
        p2_comment_box_message.config(fg='gray')
    # Clearing Page 3 contents
    if (astranotebook.index(astranotebook.select()) == 2):
        # Deleting the pervious loaded status
        status_p3_entry.config(borderwidth=2, state="normal")
        status_p3_entry.delete(0, "end")
        status_p3_entry.config(borderwidth=2, state="disable")

        # clearing the entry widget
        global page3_frequency
        for x in range(0, 20):
            if (x == 0):
                p3_harmonicsentries[x].config(state='normal')
                p3_harmonicsentries[x].delete(0, "end")
                p3_harmonicsentries[x].insert(0, "1")
                p3_harmonicsentries[x].config(state='disable')
            else:
                p3_harmonicsentries[x].delete(0, "end")
        # enable_modified_flag()
        p3_basenentry.delete(0, "end")
        p3_ambtempentry.delete(0, "end")
        p3_ambfactorentry.delete(0, "end")
        p3_AHFsizeentry.delete(0, "end")
        p3_AHFsize1entry.delete(0, "end")
        p3_fhzbtn.config(bg="white", fg="black", font=font.Font(family='Calibri', size=10, weight=font.NORMAL))
        p3_shzbtn.config(bg="white", fg="black", font=font.Font(family='Calibri', size=10, weight=font.NORMAL))
        p3_mildnotchbtn.config(bg="white", fg="black", font=font.Font(family='Calibri', size=10, weight=font.NORMAL))
        p3_lownotchbtn.config(bg="white", fg="black", font=font.Font(family='Calibri', size=10, weight=font.NORMAL))
        p3_frequencylable.config(fg="black")
        p3_ambtemplable.config(fg="black")
        p3_ambfactorlable.config(fg="black")
        p3_basenlable.config(fg="black")
        page3_frequency = 0
        p3_comment_box_message.delete(1.0, "end")
        p3_comment_box_message.insert(1.0, "Comments here...")
        p3_comment_box_message.config(fg='gray')

    if (astranotebook.index(astranotebook.select()) == 3):
        # Deleting the pervious loaded status
        primary_entry.config(borderwidth=2, state="normal")
        primary_entry.delete(0, "end")
        secondary_entry.config(borderwidth=2, state="normal")
        secondary_entry.delete(0, "end")
#-----------------------Clear contents ended-------------------------#

#---------------------New Document is started----------------#
def new_document():
    # Define a function to clear all widgets in the window ( only page -1)

    global selected_directory, selected_file_name, export_dir_name, export_file_name, export_progress_flag, modified_flag
    if (modified_flag == 1):
        response = messagebox.askyesno("Astra Nova", "Do you want to save the changes?")

        if response == True:
            export_progress_flag = 1
            save_nfo()
            export_progress_flag = 0
        else:
            pass

    selected_directory = filedialog.asksaveasfilename(
        filetypes=(("NFO", "*.nfo"), ("All Files", "*.*")),
        initialfile="Design Document")  # asks user to choose a directory
    try:
        os.chdir(os.path.dirname(selected_directory))
        selected_file_name = os.path.basename(selected_directory).split('.', 1)[0]
        # copying the user file and directory to the save location
        export_dir_name = selected_directory
        export_file_name = selected_file_name

        # Select the second tab without changing the currently selected tab or showing/hiding any tabs
        # astranotebook.tab(astrap1_frame, state=NORMAL)
        astranotebook.select(astrap1_frame)

        # Deleting the pervious loaded status
        status_p1_entry.config(borderwidth=2, state="normal")
        status_p1_entry.delete(0, "end")
        status_p1_entry.config(borderwidth=2, state="disable")

        # clearing the entry widget
        global page1_frequency
        for x in range(0, 20):
            if (x == 0):
                harmonicsentries[x].config(state='normal')
                harmonicsentries[x].delete(0, "end")
                harmonicsentries[x].insert(0, "1")
                harmonicsentries[x].config(state='disable')
            elif (x == 10):
                harmonicsentries[x].config(state='normal')
                harmonicsentries[x].delete(0, "end")
                harmonicsentries[x].config(state='disable')
            else:
                harmonicsentries[x].delete(0, "end")
        formula_combo.current(0)
        print("page-1")
        selection_process()
        print("PROCESS COMP")
        # enable_modified_flag()
        basenentry.delete(0, "end")
        ambtempentry.delete(0, "end")
        ambfactorentry.delete(0, "end")
        AHFsizeentry.delete(0, "end")
        AHFsize1entry.delete(0, "end")
        fhzbtn.config(bg="white", fg="black", font=font.Font(family='Calibri', size=10, weight=font.NORMAL))
        shzbtn.config(bg="white", fg="black", font=font.Font(family='Calibri', size=10, weight=font.NORMAL))
        highnotchbtn.config(bg="white", fg="black", font=font.Font(family='Calibri', size=10, weight=font.NORMAL))
        mildnotchbtn.config(bg="white", fg="black", font=font.Font(family='Calibri', size=10, weight=font.NORMAL))
        lownotchbtn.config(bg="white", fg="black", font=font.Font(family='Calibri', size=10, weight=font.NORMAL))
        vollable.config(fg="black")
        curlable.config(fg="black")
        kwlable.config(fg="black")
        kvalable.config(fg="black")
        kvarlable.config(fg="black")
        IPFlable.config(fg="black")
        TPFlable.config(fg="black")
        IQlable.config(fg="black")
        frequencylable.config(fg="black")
        ambtemplable.config(fg="black")
        ambfactorlable.config(fg="black")
        basenlable.config(fg="black")
        page1_frequency = 0
        comment_box_message.delete(1.0, "end")
        comment_box_message.insert(1.0, "Comments here...")
        comment_box_message.config(fg='gray')
        disable_modified_flag()
        # Select the second tab without changing the currently selected tab or showing/hiding any tabs
        # astranotebook.tab(astrap2_frame, state=NORMAL)
        # print("CHANGING TO PAGE _2")

        astranotebook.select(astrap2_frame)

        # Deleting the pervious loaded status
        status_p2_entry.config(borderwidth=2, state="normal")
        status_p2_entry.delete(0, "end")
        status_p2_entry.config(borderwidth=2, state="disable")

        # clearing the entry widget
        global page2_frequency
        for x in range(0, 30):
            if (x == 0):
                p2_harmonicsentries[x].config(state='normal')
                p2_harmonicsentries[x].delete(0, "end")
                p2_harmonicsentries[x].insert(0, "1")
                p2_harmonicsentries[x].config(state='disable')
            elif (x == 10):
                p2_harmonicsentries[x].config(state='normal')
                p2_harmonicsentries[x].delete(0, "end")
                p2_harmonicsentries[x].config(state='disable')
            else:
                p2_harmonicsentries[x].delete(0, "end")
        p2_formula_combo.current(0)
        print("page-2")
        selection_process()
        # enable_modified_flag()
        p2_basenentry.delete(0, "end")
        p2_ambtempentry.delete(0, "end")
        p2_ambfactorentry.delete(0, "end")
        p2_AHFsizeentry.delete(0, "end")
        p2_AHFsize1entry.delete(0, "end")
        p2_fhzbtn.config(bg="white", fg="black", font=font.Font(family='Calibri', size=10, weight=font.NORMAL))
        p2_shzbtn.config(bg="white", fg="black", font=font.Font(family='Calibri', size=10, weight=font.NORMAL))
        p2_mildnotchbtn.config(bg="white", fg="black", font=font.Font(family='Calibri', size=10, weight=font.NORMAL))
        p2_lownotchbtn.config(bg="white", fg="black", font=font.Font(family='Calibri', size=10, weight=font.NORMAL))
        p2_vollable.config(fg="black")
        p2_curlable.config(fg="black")
        p2_kwlable.config(fg="black")
        p2_kvalable.config(fg="black")
        p2_kvarlable.config(fg="black")
        p2_IPFlable.config(fg="black")
        p2_TPFlable.config(fg="black")
        p2_IQlable.config(fg="black")
        p2_frequencylable.config(fg="black")
        p2_ambtemplable.config(fg="black")
        p2_ambfactorlable.config(fg="black")
        p2_basenlable.config(fg="black")
        page2_frequency = 0
        p2_comment_box_message.delete(1.0, "end")
        p2_comment_box_message.insert(1.0, "Comments here...")
        p2_comment_box_message.config(fg='gray')
        disable_modified_flag()

        astranotebook.select(astrap3_frame)

        # Deleting the pervious loaded status
        status_p3_entry.config(borderwidth=2, state="normal")
        status_p3_entry.delete(0, "end")
        status_p3_entry.config(borderwidth=2, state="disable")

        # clearing the entry widget
        global page3_frequency
        for x in range(0, 20):
            if (x == 0):
                p3_harmonicsentries[x].config(state='normal')
                p3_harmonicsentries[x].delete(0, "end")
                p3_harmonicsentries[x].insert(0, "1")
                p3_harmonicsentries[x].config(state='disable')
            else:
                p3_harmonicsentries[x].delete(0, "end")
        # enable_modified_flag()
        p3_basenentry.delete(0, "end")
        p3_ambtempentry.delete(0, "end")
        p3_ambfactorentry.delete(0, "end")
        p3_AHFsizeentry.delete(0, "end")
        p3_AHFsize1entry.delete(0, "end")
        p3_fhzbtn.config(bg="white", fg="black", font=font.Font(family='Calibri', size=10, weight=font.NORMAL))
        p3_shzbtn.config(bg="white", fg="black", font=font.Font(family='Calibri', size=10, weight=font.NORMAL))
        p3_mildnotchbtn.config(bg="white", fg="black", font=font.Font(family='Calibri', size=10, weight=font.NORMAL))
        p3_lownotchbtn.config(bg="white", fg="black", font=font.Font(family='Calibri', size=10, weight=font.NORMAL))
        p3_frequencylable.config(fg="black")
        p3_ambtemplable.config(fg="black")
        p3_ambfactorlable.config(fg="black")
        p3_basenlable.config(fg="black")
        page3_frequency = 0
        p3_comment_box_message.delete(1.0, "end")
        p3_comment_box_message.insert(1.0, "Comments here...")
        p3_comment_box_message.config(fg='gray')
        disable_modified_flag()

        save_nfo()
    except:
        pass
        print("Path not defined")
#----------------New documents is ended---------------#

#------------------Save as button code started-------------#
def saveas_nfo():
    # page - 1 saveas the all the data as nfo file
    disable_modified_flag()

    global selected_directory, selected_file_name, export_dir_name, export_file_name
    # Deleting the pervious loaded status
    status_p1_entry.config(borderwidth=2, state="normal")
    status_p1_entry.delete(0, "end")
    status_p1_entry.config(borderwidth=2, state="disable")

    selected_directory = filedialog.asksaveasfilename(
        filetypes=(("NFO", "*.nfo"), ("All Files", "*.*")),
        initialfile="Design Document")  # asks user to choose a directory
    try:
        os.chdir(os.path.dirname(selected_directory))
        selected_file_name = os.path.basename(selected_directory).split('.', 1)[0]
        # copying the user file and directory to the save location
        export_dir_name = selected_directory
        export_file_name = selected_file_name
    except:
        pass
        print("Path not Defined")
    save_nfo()
#---------------Save as button code ended---------------------#

#--------------------Save button code started---------------------#
def save_nfo():
    # page - 1 save the all the data as nfo file
    global save_flag, export_progress_flag
    save_flag = 1
    print("save_flag_S", save_flag)

    disable_modified_flag()

    global sizing_data, selected_directory, selected_file_name, export_dir_name, export_file_name

    # Deleting the pervious loaded status
    status_p1_entry.config(borderwidth=2, state="normal")
    status_p1_entry.delete(0, "end")
    status_p1_entry.config(borderwidth=2, state="disable")

    # Deleting the pervious loaded status
    status_p2_entry.config(borderwidth=2, state="normal")
    status_p2_entry.delete(0, "end")
    status_p2_entry.config(borderwidth=2, state="disable")

    if (selected_directory == ''):
        selected_directory = filedialog.asksaveasfilename(
            filetypes=(("NFO", "*.nfo"), ("All Files", "*.*")),
            initialfile="Design Document")  # asks user to choose a directory
        try:
            os.chdir(os.path.dirname(selected_directory))
            selected_file_name = os.path.basename(selected_directory).split('.', 1)[0]
            # copying the user file and directory to the save location
            export_dir_name = selected_directory
            export_file_name = selected_file_name
        except:
            pass
            print("Path not defined")

    compile_save_data()
    key = b'0HyXNkfi3STXD6Eo0BYYL6TxEMuh_DVoezb45sTDWeA='
    f = Fernet(key)
    complied_message = str(sizing_data)
    message = complied_message.encode()
    encrypted = f.encrypt(message)  # Encrypt the bytes. The returning object is of type bytes
    output_file_name = selected_file_name
    extension = '.NFO'
    output_file = output_file_name + extension
    with open(str(output_file), 'wb') as f:
        f.write(encrypted)  # Write the encrypted bytes to the output file
    if (export_progress_flag == 0): messagebox.showinfo("Save Information", output_file_name + ".nfo" + " saved")
    save_flag = 0
#---------------------Save button code ended-----------------#

#---------------------Import button code started-----------------#
def import_nfo_data():
    # importing page -1 data

    global modified_flag, save_flag, export_progress_flag, page1_frequency, selected_directory, export_dir_name, selected_file_name
    global page2_frequency, page3_frequency, import_progress_flag

    import_progress_flag = 1
    if (modified_flag == 1):
        response = messagebox.askyesno("Astra Nova", "Do you want to save the changes?")

        if response == True:
            export_progress_flag = 1
            save_nfo()
            export_progress_flag = 0
        else:
            pass

    save_flag = 1
    disable_modified_flag()
    # Deleting the pervious loaded status
    status_p1_entry.config(borderwidth=2, state="normal")
    status_p1_entry.delete(0, "end")
    status_p1_entry.config(borderwidth=2, state="disable")

    # Deleting the pervious loaded status
    status_p2_entry.config(borderwidth=2, state="normal")
    status_p2_entry.delete(0, "end")
    status_p2_entry.config(borderwidth=2, state="disable")

    # Deleting the pervious loaded status
    status_p3_entry.config(borderwidth=2, state="normal")
    status_p3_entry.delete(0, "end")
    status_p3_entry.config(borderwidth=2, state="disable")

    # Deleting the pervious loaded status
    primary_entry.config(borderwidth=2, state="normal")
    primary_entry.delete(0, "end")
    secondary_entry.config(borderwidth=2, state="normal")
    secondary_entry.delete(0, "end")

    retrive_file_path = filedialog.askopenfilename(title="Select a NFO file", filetypes=(("NFO files", "*.nfo"),))
    retrive_file_name = os.path.basename(retrive_file_path).split('.', 1)[0]

    selected_directory = str(os.path.dirname(retrive_file_path)) + '/' + retrive_file_name
    # os.chdir(os.path.splitext(os.path.basename(dir_name))[0])
    export_dir_name = retrive_file_path
    try:
        os.chdir(os.path.dirname(export_dir_name))
        # file_name = os.path.basename(dir_name).split('.', 1)[0]
        export_file_name = os.path.basename(export_dir_name).split('.', 1)[0]
        # copying the user file and directory to the save location
        selected_directory = export_dir_name
        selected_file_name = export_file_name
    except:
        pass
        print("Path Nor defined")

    # selected_directory = retrive_file_path
    input_file = retrive_file_path
    key = b'0HyXNkfi3STXD6Eo0BYYL6TxEMuh_DVoezb45sTDWeA='
    try:
        with open(input_file, 'rb') as f:
            data = f.read()  # Read the bytes of the encrypted file
        f = Fernet(key)
        try:
            decrypted = f.decrypt(data)  # Decrypt the bytes. The returning object is of type bytes
            imported_data = eval(decrypted.decode())

            astranotebook.select(astrap1_frame)
            # Select the second tab without changing the currently selected tab or showing/hiding any tabs
            # astranotebook.tab(astrap1_frame, state=NORMAL)

            # posting the retrieving data

            # inserting the page - 1 harmonic entries
            for x in range(0, 20):

                if (x == 0):
                    harmonicsentries[x].config(state='normal')
                    harmonicsentries[x].delete(0, "end")
                    harmonicsentries[x].insert(0, str(imported_data['p1_harmonicentries'][x]))
                    harmonicsentries[x].config(state='disable')
                elif (x == 10):
                    harmonicsentries[x].config(state='normal')
                    harmonicsentries[x].delete(0, "end")
                    harmonicsentries[x].insert(0, str(imported_data['p1_harmonicentries'][x]))
                    harmonicsentries[x].config(state='disable')
                else:
                    harmonicsentries[x].delete(0, "end")
                    harmonicsentries[x].insert(0, str(imported_data['p1_harmonicentries'][x]))

            # inserting page -1 other data
            # enable_at_save()
            AHFsizeentry.config(state='normal')
            AHFsizeentry.delete(0, "end")
            AHFsizeentry.insert(0, str(imported_data['p1_astra_rating']))
            AHFsizeentry.config(state='disable')
            AHFsize1entry.config(state='normal')
            AHFsize1entry.delete(0, "end")
            AHFsize1entry.insert(0, str(imported_data['p1_amb_astra_rating']))
            AHFsize1entry.config(state='disable')

            page1_frequency = int(imported_data['p1_fhz'])
            if (page1_frequency == 50):
                # page - 1 50 Hz frequency selection
                if (save_flag == 0): enable_modified_flag()
                print("save_flag", save_flag)
                fhzbtn.config(bg="#05c1fa", fg="white", font=font.Font(family='Calibri', size=10, weight=font.BOLD))
                shzbtn.config(bg="white", fg="black", font=font.Font(family='Calibri', size=10, weight=font.NORMAL))
            elif (page1_frequency == 60):
                # page - 2 50 Hz frequency selection
                if (save_flag == 0): enable_modified_flag()
                print("save_flag", save_flag)
                fhzbtn.config(bg="white", fg="black", font=font.Font(family='Calibri', size=10, weight=font.NORMAL))
                shzbtn.config(bg="#05c1fa", fg="white", font=font.Font(family='Calibri', size=10, weight=font.BOLD))
            basenentry.delete(0, "end")
            basenentry.insert(0, str(imported_data['p1_notch']))
            if (str(basenentry.get()) == "3"):
                if (save_flag == 0): enable_modified_flag()
                print("save_flag", save_flag)
                highnotchbtn.config(bg="#05c1fa", fg="white",
                                    font=font.Font(family='Calibri', size=10, weight=font.BOLD))
                mildnotchbtn.config(bg="white", fg="black",
                                    font=font.Font(family='Calibri', size=10, weight=font.NORMAL))
                lownotchbtn.config(bg="white", fg="black",
                                   font=font.Font(family='Calibri', size=10, weight=font.NORMAL))
                basenentry.delete(0, "end")
                basenentry.insert(0, ("3"))
            elif (str(basenentry.get()) == "5"):
                # page - 1 notch profile medium
                if (save_flag == 0): enable_modified_flag()
                print("save_flag", save_flag)
                highnotchbtn.config(bg="white", fg="black",
                                    font=font.Font(family='Calibri', size=10, weight=font.NORMAL))
                mildnotchbtn.config(bg="#05c1fa", fg="white",
                                    font=font.Font(family='Calibri', size=10, weight=font.BOLD))
                lownotchbtn.config(bg="white", fg="black",
                                   font=font.Font(family='Calibri', size=10, weight=font.NORMAL))
                basenentry.delete(0, "end")
                basenentry.insert(0, ("5"))
            elif (str(basenentry.get()) == "10"):
                if (save_flag == 0): enable_modified_flag()
                print("save_flag", save_flag)
                highnotchbtn.config(bg="white", fg="black",
                                    font=font.Font(family='Calibri', size=10, weight=font.NORMAL))
                mildnotchbtn.config(bg="white", fg="black",
                                    font=font.Font(family='Calibri', size=10, weight=font.NORMAL))
                lownotchbtn.config(bg="#05c1fa", fg="white",
                                   font=font.Font(family='Calibri', size=10, weight=font.BOLD))
                basenentry.delete(0, "end")
                basenentry.insert(0, ("10"))
            ambtempentry.delete(0, "end")
            ambtempentry.insert(0, str(imported_data['p1_amb_temp']))
            ambfactorentry.delete(0, "end")
            ambfactorentry.insert(0, str(imported_data['p1_amp_fact']))

            if (str(imported_data['p1_reactive_selection']) == "select"):
                formula_combo.current(0)
                selection_process()
            elif (str(imported_data['p1_reactive_selection']) == "V, I, IPF, TPF"):
                formula_combo.current(1)
                selection_process()
            elif (str(imported_data['p1_reactive_selection']) == "V, kW, IPF, TPF"):
                formula_combo.current(2)
                selection_process()
            elif (str(imported_data['p1_reactive_selection']) == "V, kW, kVA, TPF"):
                formula_combo.current(3)
                selection_process()
            elif (str(imported_data['p1_reactive_selection']) == "V, kVA, IPF, TPF"):
                formula_combo.current(4)
                selection_process()
            elif (str(imported_data['p1_reactive_selection']) == "V, kVAR"):
                formula_combo.current(5)
                selection_process()
            elif (str(imported_data['p1_reactive_selection']) == "IQ"):
                formula_combo.current(6)
                selection_process()

            volentry.delete(0, "end")
            volentry.insert(0, str(imported_data['p1_voltage']))
            curentry.delete(0, "end")
            curentry.insert(0, str(imported_data['p1_current']))
            kwentry.delete(0, "end")
            kwentry.insert(0, str(imported_data['p1_active']))
            kvaentry.delete(0, "end")
            kvaentry.insert(0, str(imported_data['p1_apparent']))
            kvarentry.delete(0, "end")
            kvarentry.insert(0, str(imported_data['p1_reactive']))
            IPFentry.delete(0, "end")
            IPFentry.insert(0, str(imported_data['p1_ipf']))
            TPFentry.delete(0, "end")
            TPFentry.insert(0, str(imported_data['p1_tpf']))
            IQentry.delete(0, "end")
            IQentry.insert(0, str(imported_data['p1_iq']))
            if (str(imported_data['p1_reactive_selection']) == "IQ"):
                IQentry.delete(0, "end")
                IQentry.insert(0, str(imported_data['p1_iq']))
            else:
                IQentry.config(borderwidth=2, state="normal")
                IQentry.delete(0, "end")
                IQentry.insert(0, str(imported_data['p1_iq']))
                IQentry.config(borderwidth=2, state="disable", disabledforeground="black")

            # disable_at_save()
            if (str(imported_data['p1_comments']) == "Comments here...\n"):
                if not comment_box_message.get('1.0', 'end-1c'):
                    comment_box_message.insert('1.0', placeholder_text)
                    comment_box_message.config(fg='gray')
            else:
                comment_box_message.delete(1.0, "end")
                comment_box_message.insert(1.0, str(imported_data['p1_comments']))
                comment_box_message.config(fg='black')

            # posting the retrieving data in page -2

            astranotebook.select(astrap2_frame)
            # Select the second tab without changing the currently selected tab or showing/hiding any tabs
            # astranotebook.tab(astrap2_frame, state=NORMAL)
            # inserting the page - 2 harmonic entries
            for x in range(0, 30):

                if (x == 0):
                    p2_harmonicsentries[x].config(state='normal')
                    p2_harmonicsentries[x].delete(0, "end")
                    p2_harmonicsentries[x].insert(0, str(imported_data['p2_harmonicentries'][x]))
                    p2_harmonicsentries[x].config(state='disable')
                elif (x == 10):
                    p2_harmonicsentries[x].config(state='normal')
                    p2_harmonicsentries[x].delete(0, "end")
                    p2_harmonicsentries[x].insert(0, str(imported_data['p2_harmonicentries'][x]))
                    p2_harmonicsentries[x].config(state='disable')
                else:
                    p2_harmonicsentries[x].delete(0, "end")
                    p2_harmonicsentries[x].insert(0, str(imported_data['p2_harmonicentries'][x]))

            # inserting page -1 other data
            # enable_at_save()
            p2_AHFsizeentry.config(state='normal')
            p2_AHFsizeentry.delete(0, "end")
            p2_AHFsizeentry.insert(0, str(imported_data['p2_astra_rating']))
            p2_AHFsizeentry.config(state='disable')
            p2_AHFsize1entry.config(state='normal')
            p2_AHFsize1entry.delete(0, "end")
            p2_AHFsize1entry.insert(0, str(imported_data['p2_amb_astra_rating']))
            p2_AHFsize1entry.config(state='disable')

            page2_frequency = int(imported_data['p2_fhz'])
            if (page2_frequency == 50):
                # page - 2 50 Hz frequency selection
                if (save_flag == 0): enable_modified_flag()
                print("save_flag", save_flag)
                p2_fhzbtn.config(bg="#05c1fa", fg="white",
                                 font=font.Font(family='Calibri', size=10, weight=font.BOLD))
                p2_shzbtn.config(bg="white", fg="black",
                                 font=font.Font(family='Calibri', size=10, weight=font.NORMAL))
            elif (page2_frequency == 60):
                # page - 1 60 Hz frequency selection
                if (save_flag == 0): enable_modified_flag()
                print("save_flag", save_flag)
                p2_fhzbtn.config(bg="white", fg="black",
                                 font=font.Font(family='Calibri', size=10, weight=font.NORMAL))
                p2_shzbtn.config(bg="#05c1fa", fg="white",
                                 font=font.Font(family='Calibri', size=10, weight=font.BOLD))
                page2_frequency = 60
            p2_basenentry.delete(0, "end")
            p2_basenentry.insert(0, str(imported_data['p2_notch']))
            if (str(p2_basenentry.get()) == "3"):
                # page - 2 notch profile medium
                if (save_flag == 0): enable_modified_flag()
                print("save_flag", save_flag)
                p2_lownotchbtn.config(bg="white", fg="black",
                                      font=font.Font(family='Calibri', size=10, weight=font.NORMAL))
                p2_mildnotchbtn.config(bg="#05c1fa", fg="white",
                                       font=font.Font(family='Calibri', size=10, weight=font.BOLD))
                p2_basenentry.delete(0, "end")
                p2_basenentry.insert(0, ("3"))
            elif (str(p2_basenentry.get()) == "6"):
                # page - 2 notch profile high
                if (save_flag == 0): enable_modified_flag()
                print("save_flag", save_flag)
                p2_lownotchbtn.config(bg="#05c1fa", fg="white",
                                      font=font.Font(family='Calibri', size=10, weight=font.BOLD))
                p2_mildnotchbtn.config(bg="white", fg="black",
                                       font=font.Font(family='Calibri', size=10, weight=font.NORMAL))
                p2_basenentry.delete(0, "end")
                p2_basenentry.insert(0, ("6"))
            p2_ambtempentry.delete(0, "end")
            p2_ambtempentry.insert(0, str(imported_data['p2_amb_temp']))
            p2_ambfactorentry.delete(0, "end")
            p2_ambfactorentry.insert(0, str(imported_data['p2_amp_fact']))

            if (str(imported_data['p2_reactive_selection']) == "select"):
                p2_formula_combo.current(0)
                selection_process()
            elif (str(imported_data['p2_reactive_selection']) == "V, I, IPF, TPF"):
                p2_formula_combo.current(1)
                selection_process()
            elif (str(imported_data['p2_reactive_selection']) == "V, kW, IPF, TPF"):
                p2_formula_combo.current(2)
                selection_process()
            elif (str(imported_data['p2_reactive_selection']) == "V, kW, kVA, TPF"):
                p2_formula_combo.current(3)
                selection_process()
            elif (str(imported_data['p2_reactive_selection']) == "V, kVA, IPF, TPF"):
                p2_formula_combo.current(4)
                selection_process()
            elif (str(imported_data['p2_reactive_selection']) == "V, kVAR"):
                p2_formula_combo.current(5)
                selection_process()
            elif (str(imported_data['p2_reactive_selection']) == "IQ"):
                p2_formula_combo.current(6)
                selection_process()

            p2_volentry.delete(0, "end")
            p2_volentry.insert(0, str(imported_data['p2_voltage']))
            p2_curentry.delete(0, "end")
            p2_curentry.insert(0, str(imported_data['p2_current']))
            p2_kwentry.delete(0, "end")
            p2_kwentry.insert(0, str(imported_data['p2_active']))
            p2_kvaentry.delete(0, "end")
            p2_kvaentry.insert(0, str(imported_data['p2_apparent']))
            p2_kvarentry.delete(0, "end")
            p2_kvarentry.insert(0, str(imported_data['p2_reactive']))
            p2_IPFentry.delete(0, "end")
            p2_IPFentry.insert(0, str(imported_data['p2_ipf']))
            p2_TPFentry.delete(0, "end")
            p2_TPFentry.insert(0, str(imported_data['p2_tpf']))
            if (str(imported_data['p2_reactive_selection']) == "IQ"):
                p2_IQentry.delete(0, "end")
                p2_IQentry.insert(0, str(imported_data['p2_iq']))
            else:
                p2_IQentry.config(borderwidth=2, state="normal")
                p2_IQentry.delete(0, "end")
                p2_IQentry.insert(0, str(imported_data['p2_iq']))
                p2_IQentry.config(borderwidth=2, state="disable", disabledforeground="black")

            # disable_at_save()
            if (str(imported_data['p2_comments']) == "Comments here...\n"):
                if not p2_comment_box_message.get('1.0', 'end-1c'):
                    p2_comment_box_message.insert('1.0', placeholder_text)
                    p2_comment_box_message.config(fg='gray')
            else:
                p2_comment_box_message.delete(1.0, "end")
                p2_comment_box_message.insert(1.0, str(imported_data['p2_comments']))
                p2_comment_box_message.config(fg='black')

            # posting the retrieving data in page -3

            astranotebook.select(astrap3_frame)
            # Select the second tab without changing the currently selected tab or showing/hiding any tabs
            # astranotebook.tab(astrap3_frame, state=NORMAL)
            # inserting the page - 2 harmonic entries
            for x in range(0, 20):

                if (x == 0):
                    p3_harmonicsentries[x].config(state='normal')
                    p3_harmonicsentries[x].delete(0, "end")
                    p3_harmonicsentries[x].insert(0, str(imported_data['p3_harmonicentries'][x]))
                    p3_harmonicsentries[x].config(state='disable')
                else:
                    p3_harmonicsentries[x].delete(0, "end")
                    p3_harmonicsentries[x].insert(0, str(imported_data['p3_harmonicentries'][x]))

            # inserting page -1 other data
            # enable_at_save()
            p3_AHFsizeentry.config(state='normal')
            p3_AHFsizeentry.delete(0, "end")
            p3_AHFsizeentry.insert(0, str(imported_data['p3_astra_rating']))
            p3_AHFsizeentry.config(state='disable')
            p3_AHFsize1entry.config(state='normal')
            p3_AHFsize1entry.delete(0, "end")
            p3_AHFsize1entry.insert(0, str(imported_data['p3_amb_astra_rating']))
            p3_AHFsize1entry.config(state='disable')

            page3_frequency = int(imported_data['p3_fhz'])
            if (page3_frequency == 50):
                # page - 2 50 Hz frequency selection
                if (save_flag == 0): enable_modified_flag()
                print("save_flag", save_flag)
                p3_fhzbtn.config(bg="#05c1fa", fg="white",
                                 font=font.Font(family='Calibri', size=10, weight=font.BOLD))
                p3_shzbtn.config(bg="white", fg="black",
                                 font=font.Font(family='Calibri', size=10, weight=font.NORMAL))
            elif (page3_frequency == 60):
                # page - 1 60 Hz frequency selection
                if (save_flag == 0): enable_modified_flag()
                print("save_flag", save_flag)
                p3_fhzbtn.config(bg="white", fg="black",
                                 font=font.Font(family='Calibri', size=10, weight=font.NORMAL))
                p3_shzbtn.config(bg="#05c1fa", fg="white",
                                 font=font.Font(family='Calibri', size=10, weight=font.BOLD))
                page3_frequency = 60
            p3_basenentry.delete(0, "end")
            p3_basenentry.insert(0, str(imported_data['p3_notch']))
            if (str(p3_basenentry.get()) == "3"):
                # page - 2 notch profile medium
                if (save_flag == 0): enable_modified_flag()
                print("save_flag", save_flag)
                p3_lownotchbtn.config(bg="white", fg="black",
                                      font=font.Font(family='Calibri', size=10, weight=font.NORMAL))
                p3_mildnotchbtn.config(bg="#05c1fa", fg="white",
                                       font=font.Font(family='Calibri', size=10, weight=font.BOLD))
                p3_basenentry.delete(0, "end")
                p3_basenentry.insert(0, ("3"))
            elif (str(p3_basenentry.get()) == "6"):
                # page - 2 notch profile high
                if (save_flag == 0): enable_modified_flag()
                print("save_flag", save_flag)
                p3_lownotchbtn.config(bg="#05c1fa", fg="white",
                                      font=font.Font(family='Calibri', size=10, weight=font.BOLD))
                p3_mildnotchbtn.config(bg="white", fg="black",
                                       font=font.Font(family='Calibri', size=10, weight=font.NORMAL))
                p3_basenentry.delete(0, "end")
                p3_basenentry.insert(0, ("6"))
            p3_ambtempentry.delete(0, "end")
            p3_ambtempentry.insert(0, str(imported_data['p3_amb_temp']))
            p3_ambfactorentry.delete(0, "end")
            p3_ambfactorentry.insert(0, str(imported_data['p3_amp_fact']))

            if (str(imported_data['p3_comments']) == "Comments here...\n"):
                if not p3_comment_box_message.get('1.0', 'end-1c'):
                    p3_comment_box_message.insert('1.0', placeholder_text)
                    p3_comment_box_message.config(fg='gray')
            else:
                p3_comment_box_message.delete(1.0, "end")
                p3_comment_box_message.insert(1.0, str(imported_data['p3_comments']))
                p3_comment_box_message.config(fg='black')
                # posting the retrieving data in page -3
            global p4_data_entries
            astranotebook.select(astrap4_frame)
            print("importing...")
            #print(p4_data_entries)
                # Select the second tab without changing the currently selected tab or showing/hiding any tabs
                # astranotebook.tab(astrap3_frame, state=NORMAL)
                # inserting the page - 2 harmonic entries
            for x in range(0, 2):
                print("Done - 1")
                if (x == 0):
                    primary_entry.config(state='normal')
                    primary_entry.delete(0, "end")
                    primary_entry.insert(0, str(imported_data['p4_primary value']))
                    primary_entry.config(state='disable')
                    print("Done - 2")
                else:
                    primary_entry.delete(0, "end")
                    primary_entry.insert(0, str(imported_data['p4_primary value']))
                    print("Done - 2A")

            # enable_at_save()
            primary_entry.config(state='normal')
            primary_entry.delete(0, "end")
            primary_entry.insert(0, str(imported_data['p4_primary value']))
            primary_entry.config(state='normal')
            secondary_entry.config(state='normal')
            secondary_entry.delete(0, "end")
            secondary_entry.insert(0, str(imported_data['p4_secondary value']))
            secondary_entry.config(state='normal')

            # Get the data from 'imported_data'
            imported_opt = imported_data['p4_data values']

            # Determine the appropriate shape based on the length of data and the desired rows/columns
            num_rows = 13  # Adjust this based on your requirement
            #num_cols = len(hello_data) // num_rows  # Calculate the number of columns
            num_cols = 6  # Calculate the number of columns
            # Reshape the data into the desired shape
            try:
                opt_table = np.array(imported_opt).reshape(num_rows, num_cols)
            except:
                opt_table = np.array(imported_opt).reshape(0,0)

            # Iterate over the reshaped data to populate GUI elements
            for x, row in enumerate(opt_table):
                for y, item in enumerate(row):
                    if isinstance(item, int) or isinstance(item, str):
                        font_style = ("Arial", 10, "bold") if x == 0 else ("Arial", 10)

                        rec_table = Label(frame_4, text=str(item), width=20, anchor="center", font=font_style, borderwidth=1, relief="solid", bg=background_color)
                        rec_table.grid(row=x, column=y, padx=0, pady=0, sticky="nsew")

                        # Get the data from 'imported_data'
            imported_sugg = imported_data['p4_suggestion values']

            # Determine the appropriate shape based on the length of data and the desired rows/columns
            num_rows_1 = 3  # Adjust this based on your requirement
            num_cols_1 = 3  # Calculate the number of columns

            # Reshape the data into the desired shape
            try:
                sugg_table = np.array(imported_sugg).reshape(num_rows_1, num_cols_1)
                try_1 = 1
            except Exception as e:
                print("Error in reshaping:", e)
                sugg_table = np.array([]).reshape(0, 0)
                try_1 = 0  # Set try_1 to 0 in case of error

            # Setup GUI based on try_1 value
            if try_1 == 1:
                # Add the label "Optimal Panel Rating Based on Optimum kW"
                Label(frame_6, text="Optimal Panel Rating Based on Optimum kW", font=("Arial", 11, "bold"),
                      bg=background_color).grid(row=0, column=0, columnspan=num_cols_1, padx=9, pady=1)
                headings = ["Panel ID", "Panel Rating", "Minimum kW"]
            else:
                # Add an empty label since reshaping failed
                Label(frame_6, text="", font=("Arial", 10, "bold"), bg=background_color).grid(row=0, column=0, columnspan=num_cols_1, padx=9, pady=1)
                headings = ["", "", ""]  # Empty headings

            # Create and place the headings labels in the GUI
            for col, heading in enumerate(headings):
                Label(frame_6, text=heading, width=15, anchor="center", font=("Arial", 11, "bold"), bg=background_color).grid(row=1, column=col, padx=3, pady=5, sticky="nsew")

            # Iterate over the reshaped data to populate GUI elements
            for a, row_1 in enumerate(sugg_table, start=2):  # Start at row 2 to skip the headings row and the label row
                for b, item_1 in enumerate(row_1):
                    if isinstance(item_1, int) or isinstance(item_1, str):
                        font_style = ("Arial", 10)
                        Label(frame_6, text=str(item_1), width=15, anchor="center", font=font_style, bg=background_color).grid(row=a, column=b, padx=3, pady=5, sticky="nsew")

            location_field.config(text="location :" + str(selected_directory))
            p2_location_field.config(text="location :" + str(selected_directory))
            p3_location_field.config(text="location :" + str(selected_directory))
            p4_location_field.config(text="location :" + str(selected_directory))
            save_flag = 0
            print("save_flag_r", save_flag)
            messagebox.showinfo("Import Information", retrive_file_name + ".nfo" + " Imported")
            import_progress_flag = 0

        except Exception as e:
            print(e)
            messagebox.showerror("Import Information", "Invalid File")
    except:
        print("Path Not Defined")
#----------------Import button code ended-----------------#

#--------------------- validate_licensing code started-----------------#
def validate_licensing():
    # validating the license at the time of opening the software

    global activation_window, lic_sts, lic_validity

    validate_license_flag = 0
    # Get the MAC address of the user computer
    try:
        user_mac_address = ':'.join(format(s, '02x') for s in uuid.getnode().to_bytes(6, 'big'))
    except:
        msgbox_result = messagebox.showerror("Authentication", "Admin ERROR")
        validate_license_flag = 405

    print(user_mac_address)
    # Connecting to the Server
    try:
        url = "https://oncloud.inphase.in/php/astrasizing/validatelicense.php"
        data = {"macid": str(user_mac_address)}
        server_response = requests.post(url, data=data)
    except:
        msgbox_result = messagebox.showerror("Authentication", "Server ERROR")
        validate_license_flag = 404
    response_data = server_response.json()
    print("Response data", response_data)
    # Assuming response_data is a dictionary with a key 'status'
    lic_sts = response_data['status']
    lic_validity = response_data['validity']

    if (response_data['status'] == "Valid"):
        validate_license_flag = 1

    else:
        validate_license_flag = 0


    if (validate_license_flag == 1):
        # if the software is Registered
        # Enabling the New
        astramenu.entryconfig("New", state="normal")
        # Enabling the save
        astramenu.entryconfig("Save", state="normal")
        # Enabling the save As
        astramenu.entryconfig("Save As...", state="normal")
        # Enabling the Import
        astramenu.entryconfig("Import", state="normal")
        # Enabling the Export
        astramenu.entryconfig("Export", state="normal")
        # Enabling the Clear Content
        astramenu.entryconfig("Clear", state="normal")
        # Enabling the login
        astramenu.entryconfig("Login", state="normal")
        # Enabling the process button
        calculate.config(state="normal")
        # Enabling the clear button
        clear_comment_box.config(state="normal")
        # Enabling the formula drop down list box
        formula_combo.config(state="normal")
        status_p1_entry.config(borderwidth=2, state="normal")
        status_p1_entry.delete(0, "end")
        status_p1_entry.config(borderwidth=2, state="disable")

        # Enabling the page -2
        # Enabling the process button
        p2_calculate.config(state="normal")
        # Enabling the clear button
        p2_clear_comment_box.config(state="normal")
        # Enabling the formula drop down list box
        p2_formula_combo.config(state="normal")
        status_p2_entry.config(borderwidth=2, state="normal")
        status_p2_entry.delete(0, "end")
        status_p2_entry.config(borderwidth=2, state="disable")

        # Enabling the page -2
        # Enabling the process button
        p3_calculate.config(state="normal")
        # Enabling the clear button
        p3_clear_comment_box.config(state="normal")
        status_p3_entry.config(borderwidth=2, state="normal")
        status_p3_entry.delete(0, "end")
        status_p3_entry.config(borderwidth=2, state="disable")


    elif (validate_license_flag == 0):
        print("DISABLEING ALL ")
        # if the software is not Registered
        # disabling the New
        astramenu.entryconfig("New", state="disable")
        # disabling the save
        astramenu.entryconfig("Save", state="disable")
        # disabling the save As
        astramenu.entryconfig("Save As...", state="disable")
        # disabling the Import
        astramenu.entryconfig("Import", state="disable")
        # disabling the login
        astramenu.entryconfig("Login", state="disable")
        # disabling the Export
        astramenu.entryconfig("Export", state="disable")
        # Enabling the Clear Content
        astramenu.entryconfig("Clear", state="disable")
        # disabling the Calculation button
        calculate.config(state="disable")
        # disabling the clear button
        clear_comment_box.config(state="disable")
        # disabling the formula drop down list box
        formula_combo.config(state="disable")
        status_p1_entry.config(borderwidth=2, state="normal")
        status_p1_entry.delete(0, "end")
        status_p1_entry.insert(0, "Activate to proceed")
        status_p1_entry.config(borderwidth=2, state="disable")

        # disabling the page -2 entries

        # disabling the Calculation button
        p2_calculate.config(state="disable")
        # disabling the clear button
        p2_clear_comment_box.config(state="disable")
        # disabling the formula drop down list box
        p2_formula_combo.config(state="disable")
        # msgbox_result = messagebox.showerror("License", "Vist licensing tab to proceed")
        status_p2_entry.config(borderwidth=2, state="normal")
        status_p2_entry.delete(0, "end")
        status_p2_entry.insert(0, "Activate to proceed")
        status_p2_entry.config(borderwidth=2, state="disable")

        # disabling the page -3 entries

        # disabling the Calculation button
        p3_calculate.config(state="disable")
        # disabling the clear button
        p3_clear_comment_box.config(state="disable")
        # disabling the formula drop down list box
        p3_formula_combo.config(state="disable")
        # msgbox_result = messagebox.showerror("License", "Vist licensing tab to proceed")
        status_p3_entry.config(borderwidth=2, state="normal")
        status_p3_entry.delete(0, "end")
        status_p3_entry.insert(0, "Vist licensing to proceed")
        status_p3_entry.config(borderwidth=2, state="disable")

        msgbox_result = messagebox.showerror("License", "Activate to proceed_1")

        status_p4_entry.config(text="Activate to proceed")

    elif (validate_license_flag == 404):
        # if the software could not able to reach the data base for validation
        # disabling the New
        astramenu.entryconfig("New", state="disable")
        # disabling the save
        astramenu.entryconfig("Save", state="disable")
        # disabling the save As
        astramenu.entryconfig("Save As...", state="disable")
        # disabling the Import
        astramenu.entryconfig("Import", state="disable")
        # disabling the login
        astramenu.entryconfig("Login", state="disable")
        # disabling the Export
        astramenu.entryconfig("Export", state="disable")
        # disabling the page -1 entries
        # Enabling the Clear Content
        astramenu.entryconfig("Clear", state="disable")

        # disabling the Calculation button
        calculate.config(state="disable")
        # disabling the clear button
        clear_comment_box.config(state="disable")
        # disabling the formula drop down list box
        formula_combo.config(state="disable")
        # msgbox_result = messagebox.showerror("License", "Vist licensing tab to proceed")
        status_p1_entry.config(borderwidth=2, state="normal")
        status_p1_entry.delete(0, "end")
        status_p1_entry.insert(0, "Database not connected")
        status_p1_entry.config(borderwidth=2, state="disable")

        # disabling the page -2 entries

        # disabling the Calculation button
        p2_calculate.config(state="disable")
        # disabling the clear button
        p2_clear_comment_box.config(state="disable")
        # disabling the formula drop down list box
        p2_formula_combo.config(state="disable")
        # msgbox_result = messagebox.showerror("License", "Vist licensing tab to proceed")
        status_p2_entry.config(borderwidth=2, state="normal")
        status_p2_entry.delete(0, "end")
        status_p2_entry.insert(0, "Database not connected")
        status_p2_entry.config(borderwidth=2, state="disable")

        # disabling the page -3 entries

        # disabling the Calculation button
        p3_calculate.config(state="disable")
        # disabling the clear button
        p3_clear_comment_box.config(state="disable")
        # disabling the formula drop down list box
        p3_formula_combo.config(state="disable")
        # msgbox_result = messagebox.showerror("License", "Vist licensing tab to proceed")
        status_p3_entry.config(borderwidth=2, state="normal")
        status_p3_entry.delete(0, "end")
        status_p3_entry.insert(0, "Database not connected")
        status_p3_entry.config(borderwidth=2, state="disable")
    elif (validate_license_flag == 405):
        # if the software could not able to get the mac ID
        # disabling the New
        astramenu.entryconfig("New", state="disable")
        # disabling the save
        astramenu.entryconfig("Save", state="disable")
        # disabling the save As
        astramenu.entryconfig("Save As...", state="disable")
        # disabling the Import
        astramenu.entryconfig("Import", state="disable")
        # disabling the login
        astramenu.entryconfig("Login", state="disable")
        # disabling the Export
        astramenu.entryconfig("Export", state="disable")
        # disabling the page -1 entries
        # Enabling the Clear Content
        astramenu.entryconfig("Clear", state="disable")

        # disabling the Calculation button
        calculate.config(state="disable")
        # disabling the clear button
        clear_comment_box.config(state="disable")
        # disabling the formula drop down list box
        formula_combo.config(state="disable")
        # msgbox_result = messagebox.showerror("License", "Vist licensing tab to proceed")
        status_p1_entry.config(borderwidth=2, state="normal")
        status_p1_entry.delete(0, "end")
        status_p1_entry.insert(0, "Admin Acess Denined")
        status_p1_entry.config(borderwidth=2, state="disable")

        # disabling the page -2 entries

        # disabling the Calculation button
        p2_calculate.config(state="disable")
        # disabling the clear button
        p2_clear_comment_box.config(state="disable")
        # disabling the formula drop down list box
        p2_formula_combo.config(state="disable")
        # msgbox_result = messagebox.showerror("License", "Vist licensing tab to proceed")
        status_p2_entry.config(borderwidth=2, state="normal")
        status_p2_entry.delete(0, "end")
        status_p2_entry.insert(0, "Admin Acess Denined")
        status_p2_entry.config(borderwidth=2, state="disable")

        # disabling the page -3 entries

        # disabling the Calculation button
        p3_calculate.config(state="disable")
        # disabling the clear button
        p3_clear_comment_box.config(state="disable")
        # disabling the formula drop down list box
        p3_formula_combo.config(state="disable")
        # msgbox_result = messagebox.showerror("License", "Vist licensing tab to proceed")
        status_p3_entry.config(borderwidth=2, state="normal")
        status_p3_entry.delete(0, "end")
        status_p3_entry.insert(0, "Admin Acess Denined")
        status_p3_entry.config(borderwidth=2, state="disable")
#------------------------- validate_licensing code ended-------------------#

#---------------------Activate license code started-----------------------#
def activate_licensing():
    # activating the license which will be triggered after the submit
    global activation_window
    activate_license_flag = 0

    # user_key = str(lickeyentry[0].get()) + str(lickeyentry[1].get()) + str(lickeyentry[2].get()) + str(
    #    lickeyentry[3].get()+str(lickeyentry[4].get()))

    user_key = str(lickeyentry.get())

    # Get the MAC address of the user computer
    try:
        user_mac_address = ':'.join(format(s, '02x') for s in uuid.getnode().to_bytes(6, 'big'))
    except:
        msgbox_result = messagebox.showerror("Authentication", "Admin ERROR")
        validate_license_flag = 405

    # Connecting to the Server
    try:
        url = "https://oncloud.inphase.in/php/astrasizing/registerlicense.php"
        data = {"macid": str(user_mac_address), "licenseid": str(user_key)}
        server_response = requests.post(url, data=data)
    except:
        msgbox_result = messagebox.showerror("Authentication", "Server ERROR")
        validate_license_flag = 404
    response_data = server_response.json()
    print("Response data", response_data, response_data['status'])
    # or response_data['status'] == "Pre-Registered"
    if (response_data['status'] == "Registered"):
        activate_license_flag = 1
    elif (response_data['status'] == "Pre-Registered"):
        activate_license_flag = 2
    else:
        activate_license_flag = 0

    if (activate_license_flag == 1):
        status_p1_entry.config(borderwidth=2, state="normal")
        status_p1_entry.delete(0, "end")
        status_p1_entry.config(borderwidth=2, state="disable")
        status_p2_entry.config(borderwidth=2, state="normal")
        status_p2_entry.delete(0, "end")
        status_p2_entry.config(borderwidth=2, state="disable")
        status_p3_entry.config(borderwidth=2, state="normal")
        status_p3_entry.delete(0, "end")
        status_p3_entry.config(borderwidth=2, state="disable")
        status_p4_entry.config(text="")

        msgbox_result = messagebox.showinfo("Authentication", "Activation Successful")
        if (msgbox_result == "ok"): activation_window.destroy()
        validate_licensing()  # -- remove this
    elif (activate_license_flag == 2):
        status_p1_entry.config(borderwidth=2, state="normal")
        status_p1_entry.delete(0, "end")
        status_p1_entry.config(borderwidth=2, state="disable")
        status_p2_entry.config(borderwidth=2, state="normal")
        status_p2_entry.delete(0, "end")
        status_p2_entry.config(borderwidth=2, state="disable")
        status_p3_entry.config(borderwidth=2, state="normal")
        status_p3_entry.delete(0, "end")
        status_p3_entry.config(borderwidth=2, state="disable")
        status_p4_entry.config(text="")

        msgbox_result = messagebox.showinfo("Authentication", "Already Registered")
        if (msgbox_result == "ok"): activation_window.destroy()
        validate_licensing()  # -- remove this
    else:
        msgbox_result = messagebox.showerror("Authentication", "Activation Failed")
        if (msgbox_result == "ok"): activation_window.destroy()
#-------------------------Activate license code ended----------------#

#-----------------------Fetch license validity code started------------------#
def fetch_licens_validity():
    # fetching the license from the database to validate

    global activation_window, validity_info_lable
    validate_license_flag = 0

    # Get the MAC address of the user computer
    try:
        user_mac_address = ':'.join(format(s, '02x') for s in uuid.getnode().to_bytes(6, 'big'))
    except:
        msgbox_result = messagebox.showerror("Authentication", "Admin ERROR")
        validate_license_flag = 405

    # Connecting to the Server
    try:
        url = "https://oncloud.inphase.in/php/astrasizing/validatelicense.php"
        data = {"macid": str(user_mac_address)}
        server_response = requests.post(url, data=data)
    except:
        msgbox_result = messagebox.showerror("Authentication", "Server ERROR")
        validate_license_flag = 404
    response_data = server_response.json()
    print("Response data", response_data)

    if (response_data['status'] == "Valid"):
        validate_license_flag = 1
        date_str = str(response_data['validity'])
        date_obj = datetime.strptime(date_str, "%Y-%m-%d")

        new_date_str = date_obj.strftime("%d-%b-%Y")
    else:
        validate_license_flag = 0
        new_date_str = "Not Activated"

    if (validate_license_flag == 1):
        validity_info_lable.config(text="Your license is Valid upto : " + new_date_str)

#_--------------------------Fetch license validity code ended---------------#

#---------------------License tab code started---------------------------#
def licensing_tab():
    # opening the licensing dialogue box after clicking the licensing tab
    global lickeyentry, indian_date_str, updated_date_str, activation_window, validity_info_lable
    activation_window = Toplevel(root)
    activation_window.title("License Information")
    # child_window.geometry("200x100")

    # set the size of the child window
    child_window_width = 275
    child_window_height = 115

    # get the screen dimensions
    screen_width = root.winfo_screenwidth()
    screen_height = root.winfo_screenheight()

    # calculate the position of the child window
    x = (screen_width / 2) - (child_window_width / 2)
    y = (screen_height / 2) - (child_window_height / 2)

    # set the position of the child window
    activation_window.geometry('%dx%d+%d+%d' % (child_window_width, child_window_height, x, y))

    # Send request to Google Time API
    response = requests.get('https://www.googleapis.com/oauth2/v1/certs')

    # Extract current time from response
    current_time_str = response.headers['date']
    current_time = datetime.strptime(current_time_str, '%a, %d %b %Y %H:%M:%S %Z')

    # Convert to Indian time (UTC+05:30)
    indian_time = current_time + timedelta(hours=5, minutes=30)

    # updating 45 days from the current date
    updated_time = indian_time + timedelta(days=180)

    # Format Indian time as a string
    indian_date_str = indian_time.strftime('%d-%m-%Y')
    indian_time_str = indian_time.strftime('%H:%M:%S')

    # Format Updated time as a string
    updated_date_str = updated_time.strftime('%d-%m-%Y')
    updated_time_str = updated_time.strftime('%H:%M:%S')

    # Print Indian time
    print('Indian Time:', updated_date_str)
    print('Indian date:', updated_time_str)

    Label(activation_window, text="").grid(row=0)
    Label(activation_window, text="Enter the Key", fg="black",
          font=font.Font(family='Calibri', size=9, weight=font.NORMAL)).place(x=90, y=1)
    # creating entries for key entry

    # lickeyentry = []
    # for x in range(0, 5):
    '''
    lic_key_entry = Entry(activation_window)
    lic_key_entry.config(width=7)
    lic_key_entry.delete(0, "end")
    lic_key_entry.grid(row=1, column=x, padx=5, pady=5)
    lickeyentry.append(lic_key_entry)
    '''

    lickeyentry = Entry(activation_window)
    lickeyentry.config(width=40)
    lickeyentry.delete(0, "end")
    lickeyentry.grid(row=1, column=1, padx=5, pady=5)
    lickeyentry.focus_set()
    # lickeyentry.append(lic_key_entry)

    key_btn = Button(activation_window, text="Activate", command=activate_licensing)
    key_btn.config(height=1, width=7, bg="white", fg="black",
                   font=font.Font(family='Calibri', size=9, weight=font.NORMAL))
    key_btn.place(x=105, y=70)
    validity_info_lable = Label(activation_window, text="Your license is Valid upto :")
    validity_info_lable.place(x=1, y=95)
    fetch_licens_validity()  # -->valid function need to invoked
#-----------------------License tab code ended------------------------------#

#---------------------- License Status ------------------#
def license_sts():
    global lic_sts, lic_validity
    messagebox.showinfo("License Status", f"License Status     : {lic_sts}\nLicense Validity   : {lic_validity}")
#------------------- License Status Ended ---------------#

#-------------------------Login code started----------------------#
def open_child_window_login():
    # Login in from the manin menu
    global password_entry, adminflag, active_user

    child_window = Toplevel(root)
    child_window.title("Authentication")
    # child_window.geometry("200x100")

    # set the size of the child window
    child_window_width = 200
    child_window_height = 100

    # get the screen dimensions
    screen_width = root.winfo_screenwidth()
    screen_height = root.winfo_screenheight()

    # calculate the position of the child window
    x = (screen_width / 2) - (child_window_width / 2)
    y = (screen_height / 2) - (child_window_height / 2)

    # set the position of the child window
    child_window.geometry('%dx%d+%d+%d' % (child_window_width, child_window_height, x, y))
    # create label and entry widget for password input
    Label(child_window, text="Password:").pack()
    password_entry = Entry(child_window, show="•")
    password_entry.pack()
    password_entry.focus_set()
    adminflag = 1

    def page1_enable_admin_details():
        # print("p-1 - enable")
        global adminflag, active_user

        if (str(password_entry.get()) == "sathya"):
            basenlable.configure(text="Base (n)")
            basenentry.grid()
            basenentry.delete(0, "end")
            highnotchbtn.place_forget()
            mildnotchbtn.place_forget()
            lownotchbtn.place_forget()
            active_user = 1
            msgbox_result = messagebox.showinfo("Authentication", "Login Successful")
            if (msgbox_result == "ok"): child_window.destroy()
            # enabling the logout
            astramenu.entryconfig("Logout", state="normal")
            # disabling the login
            astramenu.entryconfig("Login", state="disable")

        else:
            adminflag = 0
            basenlable.configure(text="Notch Profile")
            password_entry.delete(0, "end")
            active_user = 0
            msgbox_result = messagebox.showerror("Authentication", "Login Failed")
            if (msgbox_result == "ok"): child_window.destroy()
            # disabling the logout
            astramenu.entryconfig("Logout", state="disable")
            # enabling the login
            astramenu.entryconfig("Login", state="normal")

    def page2_enable_admin_details():
        # print("p-2 - enable")
        global adminflag, active_user

        # if (str(password_entry.get()) == "sathya"):
        if (active_user == 1):
            p2_basenlable.configure(text="Base (n)")
            p2_basenentry.grid()
            p2_basenentry.delete(0, "end")
            p2_highnotchbtn.place_forget()
            p2_mildnotchbtn.place_forget()
            p2_lownotchbtn.place_forget()
            # active_user = 1
            # msgbox_result = messagebox.showinfo("Authentication", "Correct password!")
            # if (msgbox_result == "ok"): child_window.destroy()

        else:
            adminflag = 0
            p2_basenlable.configure(text="Notch Profile")
            active_user = 0

    def page3_enable_admin_details():
        # print("p-2 - enable")
        global adminflag, active_user

        # if (str(password_entry.get()) == "sathya"):
        if (active_user == 1):
            p3_basenlable.configure(text="Base (n)")
            p3_basenentry.grid()
            p3_basenentry.delete(0, "end")
            p3_highnotchbtn.place_forget()
            p3_mildnotchbtn.place_forget()
            p3_lownotchbtn.place_forget()
            # active_user = 1
            # msgbox_result = messagebox.showinfo("Authentication", "Correct password!")
            # if (msgbox_result == "ok"): child_window.destroy()

        else:
            adminflag = 0
            p2_basenlable.configure(text="Notch Profile")
            active_user = 0

    def validate_credentials(*args):
        page1_enable_admin_details()
        page2_enable_admin_details()
        page3_enable_admin_details()

    Label(child_window, text=" ").pack()
    Button(child_window, text="Login", command=validate_credentials).pack()
    password_entry.bind('<Return>', validate_credentials)

#----------------------Login code ended-----------------#

#---------------------Logout code started--------------------#
def open_child_window_logout():
    # Logout from the main menu
    global password_entry, adminflag, active_user

    adminflag = 0

    def page1_disable_admin_details():
        # print("p-1 - disable")
        global adminflag, active_user

        basenlable.configure(text="Notch Profile")
        basenentry.grid_remove()
        highnotchbtn.grid()
        highnotchbtn.place(x=512, y=215)
        mildnotchbtn.grid()
        mildnotchbtn.place(x=607, y=215)
        lownotchbtn.grid()
        lownotchbtn.place(x=417, y=215)
        active_user = 0
        messagebox.showinfo("Authentication", "Loged Out")
        # disabling the logout
        astramenu.entryconfig("Logout", state="disable")
        # enabling the login
        astramenu.entryconfig("Login", state="normal")

    def page2_disable_admin_details():
        # print("p-2 - disable")
        global adminflag, active_user

        p2_basenlable.configure(text="Notch Profile")
        p2_basenentry.grid_remove()
        # p2_highnotchbtn.grid()
        p2_mildnotchbtn.grid()
        p2_mildnotchbtn.place(x=683, y=215)
        p2_lownotchbtn.grid()
        p2_lownotchbtn.place(x=588, y=215)
        active_user = 0

    def page3_disable_admin_details():
        # print("p-2 - disable")
        global adminflag, active_user

        p3_basenlable.configure(text="Notch Profile")
        p3_basenentry.grid_remove()
        # p2_highnotchbtn.grid()
        p3_mildnotchbtn.grid()
        p3_mildnotchbtn.place(x=570, y=215)
        p3_lownotchbtn.grid()
        p3_lownotchbtn.place(x=475, y=215)
        active_user = 0

    def validate_credentials():
        page1_disable_admin_details()
        page2_disable_admin_details()
        page3_disable_admin_details()

    validate_credentials()
#---------------------Logout code started--------------------#

#-------------------------Menu bar code started----------------------#
menu_bar = Menu(root)
# Create file menu
astramenu = Menu(menu_bar, tearoff=0)
astramenu.add_command(label="New", command=new_document)
astramenu.add_separator()
astramenu.add_command(label="Save", command=save_nfo)
astramenu.add_command(label="Save As...", command=saveas_nfo)
astramenu.add_separator()
astramenu.add_command(label="Login", command=open_child_window_login)
astramenu.add_command(label="Logout", command=open_child_window_logout)
astramenu.add_separator()
astramenu.add_command(label="Import", command=import_nfo_data)

#------------------ Creating submenu for exporting PDF --------------------#
export_sub_menu = Menu(astramenu, tearoff=0)
export_sub_menu.add_command(label="Export", command=export_to_pdf)
export_sub_menu.add_command(label="Export all data")
astramenu.add_cascade(label="Export", menu=export_sub_menu)
#------------------------ submenu ended ----------------------#

astramenu.add_separator()
astramenu.add_command(label="Clear", command=clear_contents)
astramenu.add_separator()
astramenu.add_command(label="Exit", command=root.quit)
# disabling the logout menu at the starting
astramenu.entryconfig("Logout", state="disabled")

# Adding File option to the menu bar
menu_bar.add_cascade(label="File", menu=astramenu)

astralicense = Menu(menu_bar, tearoff=0)
astralicense.add_command(label="Activate", command=licensing_tab)
astralicense.add_separator()
astralicense.add_command(label="Status", command=license_sts)
astralicense.add_separator()

# Adding Licensing option to the menu bar
menu_bar.add_cascade(label="Licensing", menu=astralicense)

# Set menu bar as application menu
root.config(menu=menu_bar)

# Positioning on Screen

# Page Title
astralable.grid(row=0, column=2, columnspan=7, padx=10, pady=10)
# astralable.place(x=650, y=10)
# Harmonics(n) column Lable
harmonicslable.grid(row=1, column=0, padx=5, pady=5)

# I column Lable
currentlable.grid(row=1, column=1, padx=5, pady=5)

# 50 Hz Button
# fhzbtn.grid(row=2, column=3, padx=10, pady=10, columnspan=5, rowspan=3, sticky="E")
fhzbtn.place(x=465, y=140)
# 60 Hz Button
# shzbtn.grid(row=2, column=5, padx=10, pady=10, columnspan=6, rowspan=3, sticky="E")
shzbtn.place(x=560, y=140)

# Low Notch Button
lownotchbtn.place(x=417, y=215)
# High Notch Button
highnotchbtn.place(x=512, y=215)
# mild Notch Button
mildnotchbtn.place(x=607, y=215)

# AHF size (A)	lable
AHFsizelable.grid(row=12, column=0, padx=5, pady=5, columnspan=2, sticky="w")

# AHF size (A) Entry
AHFsizeentry.grid(row=12, column=1, padx=5, pady=5, columnspan=2, sticky="S")

# AHF size (A, @Ta)	lable
AHFsize1lable.grid(row=13, column=0, padx=5, pady=5, columnspan=2, sticky="w")

# AHF size (A, @Ta)	Entry
AHFsize1entry.grid(row=13, column=1, padx=5, pady=5, columnspan=2, sticky="S")

# Base n lable
basenlable.grid(row=4, column=3, columnspan=6, padx=10, pady=10, sticky="n")

# frequency lable
frequencylable.grid(row=2, column=3, padx=5, pady=5, columnspan=6)

# Base n Entry
basenentry.grid(row=5, column=3, padx=5, pady=5, columnspan=6)

# Ambient temperature  lable
ambtemplable.grid(row=6, column=3, columnspan=6, padx=10, pady=10)

# Ambient temperature Entry
ambtempentry.grid(row=7, column=3, padx=5, pady=5, columnspan=6)

# Amplification Factor
ambfactorlable.grid(row=8, column=3, columnspan=6, padx=10, pady=10)

# Ambient temperature Entry
ambfactorentry.grid(row=9, column=3, padx=5, pady=5, columnspan=6)

# Empty lable
empty_entry.grid(row=0, column=0)

#----------------------Menu bar code ended---------------------#

#-----------------------Create entries page -1 code started----------------#
def create_entries():
    # ceating the page - 1 entries

    # Harmonics entries
    xposition = 2
    global manual_entry
    for y in range(0, 2):
        for x in range(0, 10):
            if (x == 0 and y == 0):  # disabling the first entry of harmonics
                manual_entry = Entry(astrap1_frame, width=15, font=('Verdana 12'), justify='center')
                manual_entry.delete(0, "end")
                manual_entry.insert(0, (str(1)))
                manual_entry.config(borderwidth=2, state=DISABLED, disabledbackground="#F6F6F8",
                                    disabledforeground="black")
                manual_entry.grid(row=x + xposition, column=y, padx=5, pady=5)
                harmonicsentries.append(manual_entry)
            elif (x == 0 and y == 1):  # disabling the first entry of Current
                manual_entry = Entry(astrap1_frame, width=15, font=('Verdana 12'), justify='center')
                manual_entry.delete(0, "end")
                # manual_entry.insert(0, (str(100)))
                manual_entry.config(borderwidth=2, state=DISABLED, disabledbackground="#F6F6F8",
                                    disabledforeground="black")
                manual_entry.grid(row=x + xposition, column=y, padx=5, pady=5)
                harmonicsentries.append(manual_entry)
            else:
                manual_entry = Entry(astrap1_frame, width=15, font=('Verdana 12'), justify='center')
                manual_entry.config(borderwidth=2)
                manual_entry.grid(row=x + xposition, column=y, padx=5, pady=5)
                harmonicsentries.append(manual_entry)
    global genrated_data
    genrated_data = []
    # Entry widgets for Isqure
    for x in range(0, 10):
        genrated_data.append(Entry(astrap1_frame, width=15, font=('Verdana 12'), justify='center', relief=FLAT))
        genrated_data[x].config(borderwidth=1, state=DISABLED, disabledbackground="#F6F6F8", disabledforeground="black")
        # genrated_data[x].grid(row=x + xposition, column=0 + xposition, padx=5, pady=5, sticky="E")
    # Entry widgets for Ilinear
    for x in range(0, 10):
        genrated_data.append(Entry(astrap1_frame, width=15, font=('Verdana 12'), justify='center', relief=FLAT))
        genrated_data[10 + x].config(borderwidth=1, state=DISABLED, disabledbackground="#F6F6F8",
                                     disabledforeground="black")
#-------------------------Create entries page -1 code ended----------------------#

#-------------------------Create entries page -2 code started----------------------#
# Astra Lable Text
p2_astralable = Label(astrap2_frame, text="Astra Rating Calculator - 3P4W")
p2_astralable.configure(font=('Verdana', 16), bg="#F6F6F8")

# Harmonics Column
p2_harmonicslable = Label(astrap2_frame, text="Harmonics")
p2_harmonicslable.configure(font=('Verdana', 14), bg="#F6F6F8")

# Current Column
p2_currentlable = Label(astrap2_frame, text="Current-P")
p2_currentlable.configure(font=('Verdana', 14), bg="#F6F6F8")

# Current neutral Column
p2_currentnlable = Label(astrap2_frame, text="Current-N")
p2_currentnlable.configure(font=('Verdana', 14), bg="#F6F6F8")

# I*I Column
p2_currentsqrlable = Label(astrap2_frame, text="I*I")
p2_currentsqrlable.configure(font=('Verdana', 14), bg="#F6F6F8")

# I*I total Entry
p2_currentsqrentry = Entry(astrap2_frame, width=15, font=('Verdana 12'), justify='center', relief=FLAT)
# currentsqrentry = Entry(root, width=15, font=('Verdana 12'), justify='center')
p2_currentsqrentry.config(borderwidth=1, state=DISABLED, disabledbackground="#F6F6F8", disabledforeground="black")
# currentsqrentry.config(borderwidth=4, disabledbackground="white", disabledforeground="black")

# Ilin Cloumn
p2_currentlinlable = Label(astrap2_frame, text="Ilin")
p2_currentlinlable.configure(font=('Verdana', 14), bg="#F6F6F8")

# Ilin total Entry
p2_currentlinentry = Entry(astrap2_frame, width=15, font=('Verdana 12'), justify='center', relief=FLAT)
p2_currentlinentry.config(borderwidth=1, state=DISABLED, disabledbackground="#F6F6F8", disabledforeground="black")

# Status Entry
status_p2_entry = Entry(astrap2_frame, width=20, font=('Verdana 12'), justify='center', relief=FLAT)
status_p2_entry.config(borderwidth=1, state=DISABLED, disabledbackground="#F6F6F8", disabledforeground="red")

# calculate  button
p2_calculate = Button(astrap2_frame, text="Process", command=process)
p2_calculate.config(height=2, width=10, bg="white", fg="black",
                    font=font.Font(family='Calibri', size=9, weight=font.NORMAL))
# Frequency button
# 50 Hz Button
# bg="SystemButtonFace"
p2_fhzbtn = Button(astrap2_frame, text="50 Hz", command=fhz)
p2_fhzbtn.config(height=1, width=10, bg="white", fg="black",
                 font=font.Font(family='Calibri', size=10, weight=font.NORMAL))

# 60 Hz Button
# bg="SystemButtonFace"
p2_shzbtn = Button(astrap2_frame, text="60 Hz", command=shz)
p2_shzbtn.config(height=1, width=10, bg="white", fg="black",
                 font=font.Font(family='Calibri', size=10, weight=font.NORMAL))

# Notch button
# High notch Button
p2_highnotchbtn = Button(astrap2_frame, text="HIGH", command=high)
p2_highnotchbtn.config(height=1, width=10, bg="white", fg="black",
                       font=font.Font(family='Calibri', size=10, weight=font.NORMAL))

# Mild notch Button
p2_mildnotchbtn = Button(astrap2_frame, text="MEDIUM", command=mild)
p2_mildnotchbtn.config(height=1, width=10, bg="white", fg="black",
                       font=font.Font(family='Calibri', size=10, weight=font.NORMAL))

# Low notch Button
p2_lownotchbtn = Button(astrap2_frame, text="LOW", command=nonotch)
p2_lownotchbtn.config(height=1, width=10, bg="white", fg="black",
                      font=font.Font(family='Calibri', size=10, weight=font.NORMAL))

# AHf Size
p2_AHFsizelable = Label(astrap2_frame, text="AHF size (A)")
p2_AHFsizelable.configure(font=('Verdana', 14), bg="#F6F6F8")

# AHF size (A) Entry
p2_AHFsizeentry = Entry(astrap2_frame, width=15, font=('Verdana 12'), justify='center', relief=FLAT)
p2_AHFsizeentry.config(borderwidth=1, state=DISABLED, disabledbackground="#F6F6F8", disabledforeground="black")

# AHF size (A, @Ta)	lable
p2_AHFsize1lable = Label(astrap2_frame, text="AHF size (A, @Ta)")
p2_AHFsize1lable.configure(font=('Verdana', 14), bg="#F6F6F8")

# AHF size (A, @Ta)	Entry
p2_AHFsize1entry = Entry(astrap2_frame, width=15, font=('Verdana 12'), justify='center', relief=FLAT)
p2_AHFsize1entry.config(borderwidth=1, state=DISABLED, disabledbackground="#F6F6F8", disabledforeground="black")

# Frequency
p2_frequencylable = Label(astrap2_frame, text="Frequency")
p2_frequencylable.configure(font=('Verdana', 14), bg="#F6F6F8")

# Base n
p2_basenlable = Label(astrap2_frame, text="Notch Profile")
p2_basenlable.configure(font=('Verdana', 14), bg="#F6F6F8")

# Base n Entry
p2_basenentry = Entry(astrap2_frame, width=20, font=('Verdana 12'), justify='center')
p2_basenentry.config(borderwidth=2)

# Ambient temperature
p2_ambtemplable = Label(astrap2_frame, text="Ambient temperature")
p2_ambtemplable.configure(font=('Verdana', 14), bg="#F6F6F8")

# Ambient temperature Entry
p2_ambtempentry = Entry(astrap2_frame, width=20, font=('Verdana 12'), justify='center')
p2_ambtempentry.config(borderwidth=2)

# Amplification Factor
p2_ambfactorlable = Label(astrap2_frame, text="Amplification Factor")
p2_ambfactorlable.configure(font=('Verdana', 14), bg="#F6F6F8")

# Amplification Factor Entry
p2_ambfactorentry = Entry(astrap2_frame, width=20, font=('Verdana 12'), justify='center')
p2_ambfactorentry.config(borderwidth=2)

# empty position
p2_empty_entry = Text(astrap2_frame, height=1, width=1)
p2_empty_entry.config(borderwidth=0, bg="#F6F6F8", state="disabled")

# clear comment box
# bg="SystemButtonFace"
p2_clear_comment_box = Button(astrap2_frame, text="Clear", command=clear_comment_box_message)
p2_clear_comment_box.config(height=1, width=10, bg="white", fg="black",
                            font=font.Font(family='Calibri', size=9, weight=font.NORMAL))
p2_clear_comment_box.place(x=640, y=580)
# x=573
# comment box
p2_comment_box_message = Text(astrap2_frame, height=5, width=48)
p2_comment_box_message.config(borderwidth=2)
p2_comment_box_message.place(x=500, y=490)
# Set the placeholder text
p2_placeholder_text = 'Comments here...'
p2_comment_box_message.insert('1.0', p2_placeholder_text)
p2_comment_box_message.config(fg='gray')

p2_comment_box_message.bind('<FocusIn>', on_focus_in)
p2_comment_box_message.bind('<FocusOut>', on_focus_out)

# Fundamental current lable
p2_FIlable = Label(astrap2_frame, text="Reactive current (IQ)")
p2_FIlable.configure(font=('Verdana', 14), bg="#F6F6F8")
p2_FIlable.place(x=915, y=100)

# ----
# V lable
p2_vollable = Label(astrap2_frame, text="V")
p2_vollable.configure(font=('Verdana', 14), bg="#F6F6F8")
p2_vollable.place(x=910, y=170)

# V Entry
p2_volentry = Entry(astrap2_frame, width=10, font=('Verdana 12'), justify='center')
p2_volentry.config(borderwidth=2)
p2_volentry.place(x=985, y=170)

# I lable
p2_curlable = Label(astrap2_frame, text="I")
p2_curlable.configure(font=('Verdana', 14), bg="#F6F6F8")
p2_curlable.place(x=910, y=210)

# I Entry
p2_curentry = Entry(astrap2_frame, width=10, font=('Verdana 12'), justify='center')
p2_curentry.config(borderwidth=2)
p2_curentry.place(x=985, y=210)

# -----
# kw lable
p2_kwlable = Label(astrap2_frame, text="kW")
p2_kwlable.configure(font=('Verdana', 14), bg="#F6F6F8")
p2_kwlable.place(x=910, y=250)

# kw Entry
p2_kwentry = Entry(astrap2_frame, width=10, font=('Verdana 12'), justify='center')
p2_kwentry.config(borderwidth=2)
p2_kwentry.place(x=985, y=250)

# kVA lable
p2_kvalable = Label(astrap2_frame, text="kVA")
p2_kvalable.configure(font=('Verdana', 14), bg="#F6F6F8")
p2_kvalable.place(x=910, y=290)

# kva Entry
p2_kvaentry = Entry(astrap2_frame, width=10, font=('Verdana 12'), justify='center')
p2_kvaentry.config(borderwidth=2)
p2_kvaentry.place(x=985, y=290)

# kVAR lable
p2_kvarlable = Label(astrap2_frame, text="kVAR")
p2_kvarlable.configure(font=('Verdana', 14), bg="#F6F6F8")
p2_kvarlable.place(x=910, y=330)

# kVAR Entry
p2_kvarentry = Entry(astrap2_frame, width=10, font=('Verdana 12'), justify='center')
p2_kvarentry.config(borderwidth=2)
p2_kvarentry.place(x=985, y=330)

# ---
# IPF lable
p2_IPFlable = Label(astrap2_frame, text="IPF")
p2_IPFlable.configure(font=('Verdana', 14), bg="#F6F6F8")
p2_IPFlable.place(x=910, y=370)

# IPF Entry
p2_IPFentry = Entry(astrap2_frame, width=10, font=('Verdana 12'), justify='center')
p2_IPFentry.config(borderwidth=2)
p2_IPFentry.place(x=985, y=370)

# TPF lable
p2_TPFlable = Label(astrap2_frame, text="TPF")
p2_TPFlable.configure(font=('Verdana', 14), bg="#F6F6F8")
p2_TPFlable.place(x=910, y=410)

# TPF Entry
p2_TPFentry = Entry(astrap2_frame, width=10, font=('Verdana 12'), justify='center')
p2_TPFentry.config(borderwidth=2)
p2_TPFentry.place(x=985, y=410)

# ---

# IQ lable
p2_IQlable = Label(astrap2_frame, text="IQ")
p2_IQlable.configure(font=('Verdana', 14), bg="#F6F6F8")
p2_IQlable.place(x=910, y=450)

# IQ Entry
p2_IQentry = Entry(astrap2_frame, width=10, font=('Verdana 12'), justify='center')
p2_IQentry.config(borderwidth=2)
p2_IQentry.place(x=985, y=450)

# ---

#####
# Formula sections for page -1

p2_formula_title = [
    "select",
    "V, I, IPF, TPF",
    "V, kW, IPF, TPF",
    "V, kW, kVA, TPF",
    "V, kVA, IPF, TPF",
    "V, kVAR",
    "IQ",
]

# page -2 dynamic calculation bind function calls

p2_volentry.bind('<KeyRelease>', update_result)
p2_curentry.bind('<KeyRelease>', update_result)
p2_kwentry.bind('<KeyRelease>', update_result)
p2_kvaentry.bind('<KeyRelease>', update_result)
p2_kvarentry.bind('<KeyRelease>', update_result)
p2_IPFentry.bind('<KeyRelease>', update_result)
p2_TPFentry.bind('<KeyRelease>', update_result)
p2_IQentry.bind('<KeyRelease>', update_result)
p2_basenentry.bind('<KeyRelease>', update_result)  # -->new

# Creating Combo Box

p2_formula_combo = ttk.Combobox(astrap2_frame, font=('Verdana 12'), state="readonly", value=p2_formula_title)
p2_formula_combo.current(0)
p2_formula_combo.place(x=910, y=135)

# Binding the combo box
p2_formula_combo.bind("<<ComboboxSelected>>", selection_process)

# Positioning on Screen
# Page Title
p2_astralable.grid(row=0, column=3, columnspan=4, padx=10, pady=10)

# Harmonics(n) column Lable
p2_harmonicslable.grid(row=1, column=0, padx=5, pady=5)

# I column Lable
p2_currentlable.grid(row=1, column=1, padx=5, pady=5)

# I column Lable
p2_currentnlable.grid(row=1, column=2, padx=5, pady=5)

# Status Entry
# status_p2_entry.grid(row=11, column=3, columnspan=6, padx=5, pady=6, sticky="n")
status_p2_entry.place(x=565, y=415)
# Process Button
p2_calculate.place(x=710, y=445, anchor=NE)

# p2_calculate.grid(row=8, column=2, padx=10, pady=10, columnspan=3, sticky="e")

# 50 Hz Button
# fhzbtn.grid(row=2, column=3, padx=10, pady=10, columnspan=5, rowspan=3, sticky="E")
p2_fhzbtn.place(x=588, y=135)
# 60 Hz Button
# shzbtn.grid(row=2, column=5, padx=10, pady=10, columnspan=6, rowspan=3, sticky="E")
p2_shzbtn.place(x=683, y=135)

# Low Notch Button
p2_lownotchbtn.place(x=588, y=215)

# medium Notch Button

p2_mildnotchbtn.place(x=683, y=215)
# AHF size (A)	lable
p2_AHFsizelable.grid(row=12, column=0, padx=5, pady=5, columnspan=2, sticky="w")

# AHF size (A) Entry
p2_AHFsizeentry.grid(row=12, column=1, padx=5, pady=5, columnspan=2, sticky="S")

# AHF size (A, @Ta)	lable
p2_AHFsize1lable.grid(row=13, column=0, padx=5, pady=5, columnspan=2, sticky="w")

# AHF size (A, @Ta)	Entry
p2_AHFsize1entry.grid(row=13, column=1, padx=5, pady=5, columnspan=2, sticky="S")

# frequency lable
p2_frequencylable.grid(row=2, column=3, padx=5, pady=5, columnspan=6)

# Base n lable
p2_basenlable.grid(row=4, column=3, columnspan=6, padx=10, pady=10, sticky="n")

# Base n Entry
p2_basenentry.grid(row=5, column=3, padx=5, pady=5, columnspan=6)

# Ambient temperature  lable
p2_ambtemplable.grid(row=6, column=3, columnspan=6, padx=10, pady=10)

# Ambient temperature Entry
p2_ambtempentry.grid(row=7, column=3, padx=5, pady=5, columnspan=6)

# Amplification Factor
p2_ambfactorlable.grid(row=8, column=3, columnspan=6, padx=10, pady=10)

# Ambient temperature Entry
p2_ambfactorentry.grid(row=9, column=3, padx=5, pady=5, columnspan=6)

# Empty lable
p2_empty_entry.grid(row=0, column=0)


def create_entries_p2():
    # Harmonics entries
    p2_xposition = 2
    global p2_manual_entry
    for y in range(0, 3):
        for x in range(0, 10):
            if (x == 0 and y == 0):
                p2_manual_entry = Entry(astrap2_frame, width=15, font=('Verdana 12'), justify='center')
                p2_manual_entry.delete(0, "end")
                p2_manual_entry.insert(0, (str(1)))
                p2_manual_entry.config(borderwidth=2, state=DISABLED, disabledbackground="#F6F6F8",
                                       disabledforeground="black")
                p2_manual_entry.grid(row=x + p2_xposition, column=y, padx=5, pady=5)
                p2_harmonicsentries.append(p2_manual_entry)
            elif (x == 0 and y == 1):  # disabling the first entry of Current
                p2_manual_entry = Entry(astrap2_frame, width=15, font=('Verdana 12'), justify='center')
                p2_manual_entry.delete(0, "end")
                # manual_entry.insert(0, (str(100)))
                p2_manual_entry.config(borderwidth=2, state=DISABLED, disabledbackground="#F6F6F8",
                                       disabledforeground="black")
                p2_manual_entry.grid(row=x + p2_xposition, column=y, padx=5, pady=5)
                p2_harmonicsentries.append(p2_manual_entry)
            else:
                p2_manual_entry = Entry(astrap2_frame, width=15, font=('Verdana 12'), justify='center')
                p2_manual_entry.config(borderwidth=2)
                p2_manual_entry.grid(row=x + p2_xposition, column=y, padx=5, pady=5)
                p2_harmonicsentries.append(p2_manual_entry)
    global p2_genrated_data
    p2_genrated_data = []
    # Entry widgets for Isqure
    for x in range(0, 40):
        p2_genrated_data.append(Entry(astrap2_frame, width=15, font=('Verdana 12'), justify='center', relief=FLAT))
        p2_genrated_data[x].config(borderwidth=1, state=DISABLED, disabledbackground="white",
                                   disabledforeground="black")

#-------------------------Create entries page -2 code ended----------------------#

#-------------------------Create entries page -3 code stared----------------------#
# Astra Lable Text
p3_astralable = Label(astrap3_frame, text="Astra Rating Calculator - 3P(Netrural)")
p3_astralable.configure(font=('Verdana', 16), bg="#F6F6F8")

# Harmonics Column
p3_harmonicslable = Label(astrap3_frame, text="Harmonics")
p3_harmonicslable.configure(font=('Verdana', 14), bg="#F6F6F8")

# Current neutral Column
p3_currentnlable = Label(astrap3_frame, text="Current-N")
p3_currentnlable.configure(font=('Verdana', 14), bg="#F6F6F8")

# I*I Column
p3_currentsqrlable = Label(astrap3_frame, text="I*I")
p3_currentsqrlable.configure(font=('Verdana', 14), bg="#F6F6F8")

# I*I total Entry
p3_currentsqrentry = Entry(astrap3_frame, width=15, font=('Verdana 12'), justify='center', relief=FLAT)
p3_currentsqrentry.config(borderwidth=1, state=DISABLED, disabledbackground="#F6F6F8", disabledforeground="black")

# Ilin Cloumn
p3_currentlinlable = Label(astrap3_frame, text="Ilin")
p3_currentlinlable.configure(font=('Verdana', 14), bg="#F6F6F8")

# Ilin total Entry
p3_currentlinentry = Entry(astrap3_frame, width=15, font=('Verdana 12'), justify='center', relief=FLAT)
p3_currentlinentry.config(borderwidth=1, state=DISABLED, disabledbackground="#F6F6F8", disabledforeground="black")

# Status Entry
status_p3_entry = Entry(astrap3_frame, width=20, font=('Verdana 12'), justify='center', relief=FLAT)
status_p3_entry.config(borderwidth=1, state=DISABLED, disabledbackground="#F6F6F8", disabledforeground="red")

# calculate  button
p3_calculate = Button(astrap3_frame, text="Process", command=process)
p3_calculate.config(height=2, width=10, bg="white", fg="black",
                    font=font.Font(family='Calibri', size=9, weight=font.NORMAL))
# Frequency button
# 50 Hz Button
# bg="SystemButtonFace"
p3_fhzbtn = Button(astrap3_frame, text="50 Hz", command=fhz)
p3_fhzbtn.config(height=1, width=10, bg="white", fg="black",
                 font=font.Font(family='Calibri', size=10, weight=font.NORMAL))

# 60 Hz Button
# bg="SystemButtonFace"
p3_shzbtn = Button(astrap3_frame, text="60 Hz", command=shz)
p3_shzbtn.config(height=1, width=10, bg="white", fg="black",
                 font=font.Font(family='Calibri', size=10, weight=font.NORMAL))

# Notch button
# High notch Button
p3_highnotchbtn = Button(astrap3_frame, text="HIGH", command=high)
p3_highnotchbtn.config(height=1, width=10, bg="white", fg="black",
                       font=font.Font(family='Calibri', size=10, weight=font.NORMAL))

# Mild notch Button
p3_mildnotchbtn = Button(astrap3_frame, text="MEDIUM", command=mild)
p3_mildnotchbtn.config(height=1, width=10, bg="white", fg="black",
                       font=font.Font(family='Calibri', size=10, weight=font.NORMAL))

# Low notch Button
p3_lownotchbtn = Button(astrap3_frame, text="LOW", command=nonotch)
p3_lownotchbtn.config(height=1, width=10, bg="white", fg="black",
                      font=font.Font(family='Calibri', size=10, weight=font.NORMAL))

# AHf Size
p3_AHFsizelable = Label(astrap3_frame, text="AHF size (A)")
p3_AHFsizelable.configure(font=('Verdana', 14), bg="#F6F6F8")

# AHF size (A) Entry
p3_AHFsizeentry = Entry(astrap3_frame, width=15, font=('Verdana 12'), justify='center', relief=FLAT)
p3_AHFsizeentry.config(borderwidth=1, state=DISABLED, disabledbackground="#F6F6F8", disabledforeground="black")

# AHF size (A, @Ta)	lable
p3_AHFsize1lable = Label(astrap3_frame, text="AHF size (A,@Ta)")
p3_AHFsize1lable.configure(font=('Verdana', 14), bg="#F6F6F8")

# AHF size (A, @Ta)	Entry
p3_AHFsize1entry = Entry(astrap3_frame, width=15, font=('Verdana 12'), justify='center', relief=FLAT)
p3_AHFsize1entry.config(borderwidth=1, state=DISABLED, disabledbackground="#F6F6F8", disabledforeground="black")

# Frequency
p3_frequencylable = Label(astrap3_frame, text="Frequency")
p3_frequencylable.configure(font=('Verdana', 14), bg="#F6F6F8")

# Base n
p3_basenlable = Label(astrap3_frame, text="Notch Profile")
p3_basenlable.configure(font=('Verdana', 14), bg="#F6F6F8")

# Base n Entry
p3_basenentry = Entry(astrap3_frame, width=20, font=('Verdana 12'), justify='center')
p3_basenentry.config(borderwidth=2)

# Ambient temperature
p3_ambtemplable = Label(astrap3_frame, text="Ambient temperature")
p3_ambtemplable.configure(font=('Verdana', 14), bg="#F6F6F8")

# Ambient temperature Entry
p3_ambtempentry = Entry(astrap3_frame, width=20, font=('Verdana 12'), justify='center')
p3_ambtempentry.config(borderwidth=2)

# Amplification Factor
p3_ambfactorlable = Label(astrap3_frame, text="Amplification Factor")
p3_ambfactorlable.configure(font=('Verdana', 14), bg="#F6F6F8")

# Amplification Factor Entry
p3_ambfactorentry = Entry(astrap3_frame, width=20, font=('Verdana 12'), justify='center')
p3_ambfactorentry.config(borderwidth=2)

# empty position
p3_empty_entry = Text(astrap3_frame, height=1, width=1)
p3_empty_entry.config(borderwidth=0, bg="#F6F6F8", state="disabled")

# clear comment box
# bg="SystemButtonFace"
p3_clear_comment_box = Button(astrap3_frame, text="Clear", command=clear_comment_box_message)
p3_clear_comment_box.config(height=1, width=10, bg="white", fg="black",
                            font=font.Font(family='Calibri', size=9, weight=font.NORMAL))
p3_clear_comment_box.place(x=520, y=580)
# x=573
# comment box
p3_comment_box_message = Text(astrap3_frame, height=5, width=48)
p3_comment_box_message.config(borderwidth=2)
p3_comment_box_message.place(x=385, y=490)
# Set the placeholder text
p3_placeholder_text = 'Comments here...'
p3_comment_box_message.insert('1.0', p2_placeholder_text)
p3_comment_box_message.config(fg='gray')

p3_comment_box_message.bind('<FocusIn>', on_focus_in)
p3_comment_box_message.bind('<FocusOut>', on_focus_out)

# Fundamental current lable
p3_FIlable = Label(astrap3_frame, text="Reactive current (IQ)")
p3_FIlable.configure(font=('Verdana', 14), bg="#F6F6F8")

# V lable
p3_vollable = Label(astrap3_frame, text="V")
p3_vollable.configure(font=('Verdana', 14), bg="#F6F6F8")

# V Entry
p3_volentry = Entry(astrap3_frame, width=10, font=('Verdana 12'), justify='center')
p3_volentry.config(borderwidth=2)

# I lable
p3_curlable = Label(astrap3_frame, text="I")
p3_curlable.configure(font=('Verdana', 14), bg="#F6F6F8")

# I Entry
p3_curentry = Entry(astrap3_frame, width=10, font=('Verdana 12'), justify='center')
p3_curentry.config(borderwidth=2)

# -----
# kw lable
p3_kwlable = Label(astrap3_frame, text="kW")
p3_kwlable.configure(font=('Verdana', 14), bg="#F6F6F8")

# kw Entry
p3_kwentry = Entry(astrap3_frame, width=10, font=('Verdana 12'), justify='center')
p3_kwentry.config(borderwidth=2)
# p3_kwentry.place(x=985, y=250)

# kVA lable
p3_kvalable = Label(astrap3_frame, text="kVA")
p3_kvalable.configure(font=('Verdana', 14), bg="#F6F6F8")

# kva Entry
p3_kvaentry = Entry(astrap3_frame, width=10, font=('Verdana 12'), justify='center')
p3_kvaentry.config(borderwidth=2)

# kVAR lable
p3_kvarlable = Label(astrap3_frame, text="kVAR")
p3_kvarlable.configure(font=('Verdana', 14), bg="#F6F6F8")

# kVAR Entry
p3_kvarentry = Entry(astrap3_frame, width=10, font=('Verdana 12'), justify='center')
p3_kvarentry.config(borderwidth=2)
# p3_kvarentry.place(x=985, y=330)

# ---
# IPF lable
p3_IPFlable = Label(astrap3_frame, text="IPF")
p3_IPFlable.configure(font=('Verdana', 14), bg="#F6F6F8")

# IPF Entry
p3_IPFentry = Entry(astrap3_frame, width=10, font=('Verdana 12'), justify='center')
p3_IPFentry.config(borderwidth=2)

# TPF lable
p3_TPFlable = Label(astrap3_frame, text="TPF")
p3_TPFlable.configure(font=('Verdana', 14), bg="#F6F6F8")
# p3_TPFlable.place(x=910, y=410)

# TPF Entry
p3_TPFentry = Entry(astrap3_frame, width=10, font=('Verdana 12'), justify='center')
p3_TPFentry.config(borderwidth=2)

# ---

# IQ lable
p3_IQlable = Label(astrap3_frame, text="IQ")
p3_IQlable.configure(font=('Verdana', 14), bg="#F6F6F8")

# IQ Entry
p3_IQentry = Entry(astrap3_frame, width=10, font=('Verdana 12'), justify='center')
p3_IQentry.config(borderwidth=2)

# ---

#####
# Formula sections for page -1

p3_formula_title = [
    "select",
    "V, I, IPF, TPF",
    "V, kW, IPF, TPF",
    "V, kW, kVA, TPF",
    "V, kVA, IPF, TPF",
    "V, kVAR",
    "IQ",
]

# page -2 dynamic calculation bind function calls

p3_volentry.bind('<KeyRelease>', update_result)
p3_curentry.bind('<KeyRelease>', update_result)
p3_kwentry.bind('<KeyRelease>', update_result)
p3_kvaentry.bind('<KeyRelease>', update_result)
p3_kvarentry.bind('<KeyRelease>', update_result)
p3_IPFentry.bind('<KeyRelease>', update_result)
p3_TPFentry.bind('<KeyRelease>', update_result)
p3_IQentry.bind('<KeyRelease>', update_result)
p3_basenentry.bind('<KeyRelease>', update_result)  # -->new

# Creating Combo Box

p3_formula_combo = ttk.Combobox(astrap3_frame, font=('Verdana 12'), state="readonly", value=p2_formula_title)
p3_formula_combo.current(0)

# Binding the combo box
p3_formula_combo.bind("<<ComboboxSelected>>", selection_process)

# Positioning on Screen
# Page Title
p3_astralable.grid(row=0, column=3, columnspan=4, padx=10, pady=10)

# Harmonics(n) column Lable
p3_harmonicslable.grid(row=1, column=0, padx=5, pady=5)

# I column Lable
p3_currentnlable.grid(row=1, column=1, padx=5, pady=5)

# I*I column Lable
# currentsqrlable.grid(row=1, column=2, padx=5, pady=5)

# I*I total Entry
# currentsqrentry.grid(row=12, column=2, padx=5, pady=5, sticky="E")

# Ilin column Lable
# currentlinlable.grid(row=1, column=3, padx=5, pady=5)

# Ilin total Entry
# currentlinentry.grid(row=12, column=3, padx=5, pady=5, sticky="E")

# Status Entry
# status_p2_entry.grid(row=11, column=3, columnspan=6, padx=5, pady=6, sticky="n")
status_p3_entry.place(x=455, y=415)
# Process Button
# p2_calculate.grid(row=12, column=3, padx=10, pady=10, columnspan=4, sticky="NE")
p3_calculate.place(x=590, y=445, anchor=NE)  # 445

# 50 Hz Button
# fhzbtn.grid(row=2, column=3, padx=10, pady=10, columnspan=5, rowspan=3, sticky="E")
p3_fhzbtn.place(x=475, y=135)
# 60 Hz Button
p3_shzbtn.place(x=570, y=135)

# Low Notch Button
p3_lownotchbtn.place(x=475, y=215)

# mild Notch Button
p3_mildnotchbtn.place(x=570, y=215)

# AHF size (A)	lable
p3_AHFsizelable.grid(row=12, column=0, padx=5, pady=5, columnspan=2, sticky="w")

# AHF size (A) Entry
p3_AHFsizeentry.grid(row=12, column=1, padx=5, pady=5, columnspan=2, sticky="S")

# AHF size (A, @Ta)	lable
p3_AHFsize1lable.grid(row=13, column=0, padx=5, pady=5, columnspan=1, sticky="w")

# AHF size (A, @Ta)	Entry
p3_AHFsize1entry.grid(row=13, column=1, padx=5, pady=5, columnspan=2, sticky="S")

# frequency lable
p3_frequencylable.grid(row=2, column=3, padx=5, pady=5, columnspan=6)

# Base n lable
p3_basenlable.grid(row=4, column=3, columnspan=6, padx=10, pady=10, sticky="n")

# Base n Entry
p3_basenentry.grid(row=5, column=3, padx=5, pady=5, columnspan=6)

# Ambient temperature  lable
p3_ambtemplable.grid(row=6, column=3, columnspan=6, padx=10, pady=10)

# Ambient temperature Entry
p3_ambtempentry.grid(row=7, column=3, padx=5, pady=5, columnspan=6)

# Amplification Factor
p3_ambfactorlable.grid(row=8, column=3, columnspan=6, padx=10, pady=10)

# Ambient temperature Entry
p3_ambfactorentry.grid(row=9, column=3, padx=5, pady=5, columnspan=6)

# Empty lable
p3_empty_entry.grid(row=0, column=0)


def create_entries_p3():
    # Harmonics entries
    p3_xposition = 2
    global p3_manual_entry
    for y in range(0, 2):
        for x in range(0, 10):
            if (x == 0 and y == 0):
                p3_manual_entry = Entry(astrap3_frame, width=15, font=('Verdana 12'), justify='center')
                p3_manual_entry.delete(0, "end")
                p3_manual_entry.insert(0, (str(1)))
                p3_manual_entry.config(borderwidth=2, state=DISABLED, disabledbackground="#F6F6F8",
                                       disabledforeground="black")
                p3_manual_entry.grid(row=x + p3_xposition, column=y, padx=5, pady=5)
                p3_harmonicsentries.append(p3_manual_entry)
            else:
                p3_manual_entry = Entry(astrap3_frame, width=15, font=('Verdana 12'), justify='center')
                p3_manual_entry.config(borderwidth=2)
                p3_manual_entry.grid(row=x + p3_xposition, column=y, padx=5, pady=5)
                p3_harmonicsentries.append(p3_manual_entry)
    global p3_genrated_data
    p3_genrated_data = []
    # Entry widgets for Isqure
    for x in range(0, 40):
        p3_genrated_data.append(Entry(astrap3_frame, width=15, font=('Verdana 12'), justify='center', relief=FLAT))
        p3_genrated_data[x].config(borderwidth=1, state=DISABLED, disabledbackground="white",
                                   disabledforeground="black")
#-------------------------Create entries page -3 code endeed----------------------#

basenentry.grid_remove()
p2_basenentry.grid_remove()
status_p2_entry.config(borderwidth=2, state="normal")
status_p2_entry.delete(0, "end")
# status_p2_entry.insert(0, ("password matched"))
status_p2_entry.config(borderwidth=2, state="disable")

p3_basenentry.grid_remove()
status_p3_entry.config(borderwidth=2, state="normal")
status_p3_entry.delete(0, "end")
# status_p3_entry.insert(0, ("password matched"))
status_p3_entry.config(borderwidth=2, state="disable")

# Creating page - 1 entries
create_entries()
# Creating page - 2 entries
create_entries_p2()
# Creating page - 3 entries
create_entries_p3()

# Tab Change events
# root.bind("<<NotebookTabChanged>>", tab_changed)

# binding all page -1 harmonics events to modified_flag
for x in range(0, 20): harmonicsentries[x].bind('<KeyRelease>', enable_modified_flag)
ambtempentry.bind('<KeyRelease>', enable_modified_flag)
ambfactorentry.bind('<KeyRelease>', enable_modified_flag)

# binding all page -2 harmonics events to modified_flag
for x in range(0, 30): p2_harmonicsentries[x].bind('<KeyRelease>', enable_modified_flag)
p2_ambtempentry.bind('<KeyRelease>', enable_modified_flag)
p2_ambfactorentry.bind('<KeyRelease>', enable_modified_flag)

# binding all page -3 harmonics events to modified_flag
for x in range(0, 20): p3_harmonicsentries[x].bind('<KeyRelease>', enable_modified_flag)

p3_ambtempentry.bind('<KeyRelease>', enable_modified_flag)
p3_ambfactorentry.bind('<KeyRelease>', enable_modified_flag)

# Modification indication for page 1
modified_indication = Label(astrap1_frame, text="")
modified_indication.configure(font=('Verdana', 14), bg="#F6F6F8")
modified_indication.place(x=721, y=10)

# Modification indication for page 2
p2_modified_indication = Label(astrap2_frame, text="")
p2_modified_indication.configure(font=('Verdana', 14), bg="#F6F6F8")
p2_modified_indication.place(x=838, y=10)

# Modification indication for page 3
p3_modified_indication = Label(astrap3_frame, text="")
p3_modified_indication.configure(font=('Verdana', 14), bg="#F6F6F8")
p3_modified_indication.place(x=763, y=10)

# location field - page - 1
location_field = Label(astrap1_frame, text="location :" + str(selected_directory))
location_field.configure(font=('Verdana', 9), bg="#F6F6F8")
# location_field.place(x=5, y=580)
location_field.place(relx=0, rely=1.0, anchor="sw")

# location field - page - 2
p2_location_field = Label(astrap2_frame, text="location :" + str(selected_directory))
p2_location_field.configure(font=('Verdana', 9), bg="#F6F6F8")
# p2_location_field.place(x=5, y=580)
p2_location_field.place(relx=0, rely=1.0, anchor="sw")

# location field - page - 3
p3_location_field = Label(astrap3_frame, text="location :" + str(selected_directory))
p3_location_field.configure(font=('Verdana', 9), bg="#F6F6F8")
# p3_location_field.place(x=5, y=580)
p3_location_field.place(relx=0, rely=1.0, anchor="sw")

#-----------------------------------------------True Power Factor------------------------------------------------------#

#--------------This frame for entries and submit button----------------#

frame_1 = Frame(astrap4_frame, bg=background_color, height=300)
frame_1.pack(expand=True, fill=BOTH)

# Modification indication for page 4
p4_modified_indication = Label(frame_1, text="", font=('Verdana', 14, 'bold'), bg=background_color, fg="black")
p4_modified_indication.place(relx=0.651, rely=0.117, anchor=CENTER)
#------------------this frame for label the condition-1 name---------------------#
frame_2 = Frame(astrap4_frame, height=200, bg=background_color, highlightbackground="Black", highlightthickness=2)
frame_2.pack(expand=True, fill=BOTH)
rec_lab = Label(frame_2, text="Recommendation", font=("Arial", 11), bg=background_color).pack(anchor=N)

#Label(frame_2, text="", bg=background_color).pack(anchor=N)
empty_lab = Label(astrap4_frame, text="", bg=background_color, font=("Arial", 5)).pack()

#----------------- This frame for output of the condition-1 ------------------#
frame_4 = Frame(frame_2, height=260,width=1200, bg=background_color)
frame_4.pack(expand=True, anchor=CENTER)

'''for x in range(0, 13):
    for y in range(0, 6):
        font_style = ("Arial", 15, "bold")
        true_pf = Label(frame_4, text="", width=20, anchor="center", bg=background_color, font=font_style)
        true_pf.grid(row=x, column=y, padx=0, pady=0, sticky="nsew")
        #print(data_2[x][y])'''

#------------------this frame for label the suggestion name---------------------#
frame_5 = Frame(astrap4_frame, height=150, bg=background_color, highlightbackground="Black", highlightthickness=2)
frame_5.pack(expand=True, fill=BOTH)
sug_label = Label(frame_5, text="Suggestion", font=("Arial", 11), bg=background_color).pack(anchor=N)

#--------------------This frame for location label--------------------------#
frame_3 = Frame(astrap4_frame, height=90, bg=background_color)
frame_3.pack(expand=True, fill=BOTH)

#----------------- This frame for output of the suggestion ------------------#
frame_6 = Frame(frame_5, height=155, bg=background_color)
frame_6.pack(expand=True, anchor=CENTER)

pri_less_sec = Label(frame_1, text="", bg=background_color, fg="Red", font=("Arial", 10))
pri_less_sec.place(relx=0.98, rely=1.0,anchor=SE)

status_p4_entry = Label(frame_3, text="", font=("Arial", 10, "bold"),bg=background_color, fg="Red")
status_p4_entry.place(relx=0.5, rely=0.5, anchor=CENTER)

# location field - page - 4
p4_location_field = Label(frame_3, text="location :" + str(selected_directory))
p4_location_field.configure(font=('Verdana', 9), bg="#F6F6F8")
# p3_location_field.place(x=5, y=580)
p4_location_field.place(relx=0, rely=1.0, anchor="sw")

# Create a Frame widget and place it in the bottom-right corner of the root window
logoframe4 = Frame(frame_3, bg="#F6F6F8")
logoframe4.place(relx=1.0, rely=1.0, anchor="se")

# Create a text widget with some text
astramakep4 = Text(logoframe4, height=1, width=20, highlightthickness=0, relief="flat")
astramakep4.insert("end", "  Made with ❤ in India")
astramakep4.configure(font=('Verdana', 8), bg="#F6F6F8", fg="gray")
astramakep4.grid(row=0, column=0)

# Get the position of the 'W' character
posp4 = astramakep3.search("❤", "1.0")

# Add a tag to the 'W' character
astramakep4.tag_add("red", posp3, f"{posp3}+1c")

# Configure the tag to use a different color
astramakep4.tag_config("red", foreground="red", background="#F6F6F8", font=('Verdana', 8))

# Disable the text widget so it's read-only and non-editable
astramakep4.configure(state="disabled")

# Copy Right Lable - page - 4
astracoprightp4 = Label(logoframe4, text="© 2024, InPhase | All Rights Reserved")
astracoprightp4.configure(font=('Verdana', 8), bg="#F6F6F8")
astracoprightp4.grid(row=1, column=0)

primary = StringVar()
secondary = StringVar()

def submit():
    global lab
    for widget_1 in frame_6.winfo_children():
        widget_1.destroy()
    try:
        if not primary_entry.get():
            messagebox.showerror("Error!", "Enter primary values")
        elif not secondary_entry.get():
            messagebox.showerror("Error!", "Enter secondary values")
        elif not primary_entry.get() and not secondary_entry.get():
            messagebox.showerror("Error!", "Enter both values")

        pri_val = float(primary_entry.get())
        sec_val = float(secondary_entry.get())
        if pri_val > 0 and sec_val > 0:
            if pri_val > sec_val:
                cal_opt_val = round(pri_val / sec_val)
                condition_1(pri_val, sec_val, cal_opt_val)
            else:
                pri_less_sec.config(text="Warning: primary value is less than secondary value")
                swap_proceed = messagebox.askyesnocancel('Conformation', "Warning: primary value is less than secondary value.\n\tDo you want to swap the values?")
                if swap_proceed == True:
                    temp = float(primary_entry.get())
                    pri_val = float(secondary_entry.get())
                    sec_val = temp
                    messagebox.showinfo('Processed', "Values are swapped and processed")
                    cal_opt_val = round(pri_val / sec_val)
                    condition_1(pri_val, sec_val, cal_opt_val)

                elif swap_proceed == False:
                    messagebox.showinfo('Processed', "You entered value is processed")
                    cal_opt_val = round(pri_val / sec_val)
                    condition_1(pri_val, sec_val, cal_opt_val)

        else:
            messagebox.showerror('Error!', "Enter positive integers and greater than zero values")
    except :
        pass
        '''messagebox.showerror("Error!", "Enter both values as numbers")
            status_p4_entry.config(text="Enter the both values...")'''

#------------------- Submit button command function --------------------#
def condition_1(pri_val, sec_val, cal_opt_val):
    global hello, rows, columns
    global minimum_kw
    global rec_tab_val

    kva = math.floor(1.732 * 415.0 * (pri_val / 1000))
    kvar = math.ceil(kva * 0.03)
    kw = math.ceil(kva * 0.05)
    opt_ct_ratio = [
            ["ID", "Panel Rating", "Optimal CT Ratio", "Condition", "Minimum kVAr", "Minimum kW"],
            [1, "630A ASTRA - B03", 1252, "", "", ""],
            [2, "630A ASTRA - B05", 1809, "", "", ""],
            [3, "630A ASTRA - B10", 2782, "", "", ""],
            [4, "315A ASTRA - B03", 835, "", "", ""],
            [5, "315A ASTRA - B05", 1252, "", "", ""],
            [6, "315A ASTRA - B10", 1530, "", "", ""],
            [7, "420A ASTRA - B10", 1669, "", "", ""],
            [8, "420A ASTRA - B05", 1252, "", "", ""],
            [9, "210A ASTRA - B10", 1530, "", "", ""],
            [10, "210A ASTRA - B05", 696, "", "", ""],
            [11, "150A ASTRA - B05", 835, "", "", ""],
            [12, "150A ASTRA - B10", 1669, "", "", ""]
        ]
    opt_kw = [
            ["Optimal kW"],
            [225],
            [325],
            [500],
            [150],
            [225],
            [275],
            [300],
            [225],
            [275],
            [125],
            [150],
            [300]
        ]
    rows = len(opt_ct_ratio)
    columns = len(opt_ct_ratio[0])

    rec_tab_val = [row[:] for row in opt_ct_ratio]
    #print("Data 2 : ", data_2)
    for row_idx, row in enumerate(rec_tab_val[1:], start=1):
        data_val = row[2]
        if data_val < cal_opt_val:
            row[3:5] = "Condition - 1", kvar, kw
        else:
            if row_idx < len(opt_kw):
                row[3:5] = "Condition - 2", "-", opt_kw[row_idx][0]

    for widget in frame_4.winfo_children():
        widget.destroy()

    #WATCH
    for x in range(len(rec_tab_val)):
        for y in range(len(rec_tab_val[0])):
            font_style = ("Arial", 10, "bold") if x == 0 else ("Arial", 10)
            #------------------------------------- To generate the Labels in table format -------------------------------#
            rec_table = Label(frame_4, text=rec_tab_val[x][y], width=20, anchor="center",bg=background_color, borderwidth=1, relief="solid", font=font_style)
            rec_table.grid(row=x, column=y, padx=0, pady=0, sticky="nsew")
            print(rec_tab_val[x][y])

    minimum_kw = []
    sort_column = sorted(rec_tab_val, key=lambda x: str(x[2]))
    first = sort_column[:3]
    for row in sort_column:
        if str(row[5]) in [str(item[5]) for item in first]:
            minimum_kw.append((row[0], row[1], row[5]))

    table_frame = Frame(frame_6, bg=background_color)
    table_frame.grid(row=len(opt_ct_ratio) + 1, column=0, columnspan=1, pady=3)

    sugg_lable_head_1 = Label(table_frame, text="Optimal Panel Rating Based on Optimum kW", font=("Arial", 11, "bold"), bg=background_color)
    sugg_lable_head_1.grid(row=0, column=0, columnspan=5, padx=9, pady=0)
    sugg_lable_head_2 = Label(table_frame, text="Panel ID ", font=("Arial", 10, "bold"), bg=background_color, width=15)
    sugg_lable_head_2.grid(row=1, column=0, padx=5, pady=3)
    sugg_lable_head_3 = Label(table_frame, text="Panel Rating", font=("Arial", 10, "bold"), bg=background_color, width=15)
    sugg_lable_head_3.grid(row=1, column=1, padx=5, pady=3)
    sugg_lable_head_4 = Label(table_frame, text="Minimum kW", font=("Arial", 10, "bold"), bg=background_color, width=15)
    sugg_lable_head_4.grid(row=1, column=2, padx=5, pady=3)

    row_index = 2
    for i in range(min(3, len(minimum_kw))):
        panel_id = minimum_kw[i][0]
        panel_rating = minimum_kw[i][1]
        panel_kw = minimum_kw[i][2]
        pan_rat = (panel_id, panel_rating, panel_kw)

        category, data, opt_ct_ratio = pan_rat

        sugg_pan_id = Label(table_frame, text=category, font=("Arial", 10), bg=background_color, width=15).grid(row=row_index, column=0, padx=5, pady=5)
        sugg_pan_rat = Label(table_frame, text=str(data), font=("Arial", 10), bg=background_color, width=15).grid(row=row_index, column=1, padx=5, pady=5)
        sugg_pan_kw = Label(table_frame, text=str(opt_ct_ratio), font=("Arial", 10), bg=background_color, width=15).grid(row=row_index, column=2, padx=5, pady=5)
        row_index += 1
    print("minimum_kw : ", minimum_kw)
#-------------------------- reset the entry and frames ---------------------#
def update_results(*args):
    global pri_less_sec
    global head_notebook
    global status_p4_entry
    print(pri_less_sec)
    for widget in frame_4.winfo_children():
        widget.destroy()
    for widget_1 in frame_6.winfo_children():
        widget_1.destroy()
    pri_less_sec.config(text="")
    status_p4_entry.config(text="")
    astranotebook.tab(3, text="True Power Factor*")
    print(head_notebook)
    head_notebook.config(text="True Power Factor Performance Calculator *")
    enable_modified_flag()
    # p4_modified_indication.config(text="*")
    location_field.config(text="location :" + str(selected_directory))
    p2_location_field.config(text="location :" + str(selected_directory))
    p3_location_field.config(text="location :" + str(selected_directory))
    p4_location_field.config(text="location :" + str(selected_directory))
    print(modified_flag)

# Labels
head_notebook=Label(frame_1, text="True Power Factor Performance Calculator  ", bg=background_color, fg="Black", font=("Arial", 17))
head_notebook.pack(padx=10, pady=10)
pri_ent_label = Label(frame_1, text="Primary value of transformer", bg=background_color, fg="Black", font=("Arial", 11)).place(relx=0.35, rely=0.48, anchor=W)
primary_entry = Entry(frame_1, textvariable=primary, font=("Arial", 11))
primary_entry.place(relx=0.52, rely=0.48, anchor=W)

sec_ent_label = Label(frame_1, text="Secondary value of transformer", bg=background_color, fg="Black", font=("Arial", 11)).place(relx=0.35, rely=0.75, anchor=W)
secondary_entry = Entry(frame_1, textvariable=secondary, font=("Arial", 11))
secondary_entry.place(relx=0.52, rely=0.75, anchor=W)

primary.trace_add('write', update_results)
secondary.trace_add('write', update_results)

Button(frame_1, text="Submit", command=submit, font=("Arial", 12, "bold")).pack(side=RIGHT, padx=350, pady=18)
#-----------------------------------------------------------------------------------------------------------------------------------------------------------#

#-------------------------Access the buttons by keys--------------------------#
def on_ctrl_1(event):
    if event.state & 0x4 and event.keysym == "1":
        astranotebook.select(0)
def on_ctrl_2(event):
    if event.state & 0x4 and event.keysym == "2":
        astranotebook.select(1)
def on_ctrl_3(event):
    if event.state & 0x4 and event.keysym == "3":
        astranotebook.select(2)
def on_ctrl_4(event):
    if event.state & 0x4 and event.keysym == "4":
        astranotebook.select(3)

root.bind("<Control-KeyPress-1>", on_ctrl_1)
root.bind("<Control-KeyPress-2>", on_ctrl_2)
root.bind("<Control-KeyPress-3>", on_ctrl_3)
root.bind("<Control-KeyPress-4>", on_ctrl_4)


#-------------------------------------------------------#

# maximize the window
root.state('zoomed')

# root.resizable(0, 0)  # can not maximize

import_progress_flag = 1
astranotebook.select(astrap3_frame)
selection_process()
astranotebook.select(astrap4_frame)
astranotebook.select(astrap2_frame)
selection_process()
astranotebook.select(astrap1_frame)
selection_process()
import_progress_flag = 0
validate_licensing() #remove this code 29-APR-2024
root.mainloop()
